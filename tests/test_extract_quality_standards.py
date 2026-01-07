"""
Unit tests for extract_quality_standards.py
"""

import pytest
import os
import sys
import logging
from docx import Document

# Add scripts directory to path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
scripts_dir = os.path.join(project_root, '.claude', 'skills', 'ra-doc-assit', 'scripts')
sys.path.insert(0, scripts_dir)

from extract_quality_standards import (
    extract_text_with_formatting,
    extract_quality_standards_table,
    format_as_markdown_table,
    extract_quality_standards_table_from_docx
)

# Configure logging for tests
logging.basicConfig(level=logging.DEBUG)


class TestExtractTextWithFormatting:
    """Test extract_text_with_formatting function"""

    def test_extract_plain_text(self):
        """Test extraction of plain text without formatting"""
        # Create a mock cell with plain text
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "Plain Text"

        result = extract_text_with_formatting(cell)
        assert result == "Plain Text"

    def test_extract_superscript_numbers(self):
        """Test extraction of superscript numbers"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        # Add text with superscript
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run("H")
        run2 = paragraph.add_run("2")
        run2.font.superscript = True

        result = extract_text_with_formatting(cell)
        assert "²" in result or "^2" in result

    def test_extract_subscript_numbers(self):
        """Test extraction of subscript numbers"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        # Add text with subscript
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run("H")
        run2 = paragraph.add_run("2")
        run2.font.subscript = True

        result = extract_text_with_formatting(cell)
        assert "₂" in result or "_2" in result

    def test_extract_mixed_formatting(self):
        """Test extraction of mixed superscript and subscript"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        # Add text with mixed formatting
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run("H")
        run2 = paragraph.add_run("2")
        run2.font.subscript = True
        run3 = paragraph.add_run("SO")
        run4 = paragraph.add_run("4")
        run4.font.superscript = True

        result = extract_text_with_formatting(cell)
        assert len(result) > 0

    def test_extract_empty_cell(self):
        """Test extraction of empty cell"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = ""

        result = extract_text_with_formatting(cell)
        assert result == ""

    def test_extract_multiple_paragraphs(self):
        """Test extraction of cell with multiple paragraphs"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        # Add multiple paragraphs
        cell.paragraphs[0].text = "First"
        cell.add_paragraph("Second")

        result = extract_text_with_formatting(cell)
        assert "First" in result
        assert "Second" in result


class TestFormatAsMarkdownTable:
    """Test format_as_markdown_table function"""

    def test_format_empty_table(self):
        """Test formatting of empty table"""
        result = format_as_markdown_table([])
        assert result == "No table data found"

    def test_format_simple_table(self):
        """Test formatting of simple table"""
        table_data = [
            ['类型', '检验项目', '检验方法', '质量标准'],
            ['鉴别', '外观', '目视', '澄清']
        ]
        result = format_as_markdown_table(table_data)

        assert '| 类型 |' in result
        assert '| --- |' in result
        assert '| 鉴别 |' in result

    def test_format_with_custom_columns(self):
        """Test formatting with custom column names"""
        table_data = [
            ['A', 'B'],
            ['data1', 'data2']
        ]
        result = format_as_markdown_table(table_data, target_columns=['Col1', 'Col2'])

        assert '| Col1 | Col2 |' in result

    def test_format_uneven_rows(self):
        """Test formatting with uneven row lengths"""
        table_data = [
            ['A', 'B', 'C'],
            ['1', '2'],
            ['3', '4', '5', '6']
        ]
        result = format_as_markdown_table(table_data, target_columns=['X', 'Y', 'Z'])

        # Should pad and truncate to match target columns
        assert '| X | Y | Z |' in result
        assert '---' in result

    def test_format_with_special_characters(self):
        """Test formatting with special characters"""
        table_data = [
            ['类型', '检验项目'],
            ['鉴别', 'pH值³⁄₄']
        ]
        result = format_as_markdown_table(table_data)

        assert 'pH值' in result


class TestExtractQualityStandardsTable:
    """Test extract_quality_standards_table function"""

    @pytest.fixture
    def sample_doc_path(self, tmp_path):
        """Create a sample Word document for testing"""
        doc = Document()
        doc.add_paragraph("4.3 检验项目、方法和标准")

        # Add quality standards table
        table = doc.add_table(rows=2, cols=4)
        table.rows[0].cells[0].text = '类型'
        table.rows[0].cells[1].text = '检验项目'
        table.rows[0].cells[2].text = '检验方法'
        table.rows[0].cells[3].text = '质量标准'
        table.rows[1].cells[0].text = '鉴别'
        table.rows[1].cells[1].text = '外观'
        table.rows[1].cells[2].text = '目视'
        table.rows[1].cells[3].text = '澄清'

        doc_path = tmp_path / "test.docx"
        doc.save(str(doc_path))
        return str(doc_path)

    def test_extract_with_section_43(self, sample_doc_path):
        """Test extraction with section 4.3 present"""
        result = extract_quality_standards_table(sample_doc_path)

        assert len(result) > 0
        assert len(result[0]) == 4  # Should have 4 columns
        assert '类型' in result[0]

    def test_extract_without_quality_table(self, tmp_path):
        """Test extraction when no quality table exists"""
        doc = Document()
        doc.add_paragraph("4.3 检验项目、方法和标准")
        # Add a non-quality table
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = 'A'
        table.rows[0].cells[1].text = 'B'

        doc_path = tmp_path / "no_quality.docx"
        doc.save(str(doc_path))

        result = extract_quality_standards_table(str(doc_path))
        assert result == []

    def test_extract_table_with_keywords(self, tmp_path):
        """Test extraction with table containing quality keywords"""
        doc = Document()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = '检验项目'
        table.rows[0].cells[1].text = '检验方法'
        table.rows[0].cells[2].text = '质量标准'

        doc_path = tmp_path / "keywords.docx"
        doc.save(str(doc_path))

        result = extract_quality_standards_table(str(doc_path))
        assert len(result) > 0

    def test_extract_nonexistent_file(self):
        """Test extraction with non-existent file"""
        with pytest.raises(Exception):
            extract_quality_standards_table("nonexistent.docx")


class TestExtractQualityStandardsTableFromDocx:
    """Test extract_quality_standards_table_from_docx function"""

    def test_extract_from_nonexistent_file(self):
        """Test extraction from non-existent file"""
        result = extract_quality_standards_table_from_docx("nonexistent.docx")
        assert "Error" in result or "File not found" in result

    def test_extract_from_invalid_file(self, tmp_path):
        """Test extraction from invalid file"""
        # Create a text file instead of docx
        invalid_path = tmp_path / "invalid.txt"
        invalid_path.write_text("Not a docx file")

        result = extract_quality_standards_table_from_docx(str(invalid_path))
        assert "Error" in result

    @pytest.fixture
    def valid_doc_path(self, tmp_path):
        """Create a valid document for testing"""
        doc = Document()
        doc.add_paragraph("4.3 检验项目、方法和标准")

        table = doc.add_table(rows=2, cols=4)
        table.rows[0].cells[0].text = '类型'
        table.rows[0].cells[1].text = '检验项目'
        table.rows[0].cells[2].text = '检验方法'
        table.rows[0].cells[3].text = '质量标准'
        table.rows[1].cells[0].text = '鉴别'
        table.rows[1].cells[1].text = '外观'

        doc_path = tmp_path / "valid.docx"
        doc.save(str(doc_path))
        return str(doc_path)

    def test_successful_extraction(self, valid_doc_path):
        """Test successful extraction"""
        result = extract_quality_standards_table_from_docx(valid_doc_path)

        assert "| 类型 |" in result
        assert "| --- |" in result
        assert "| 鉴别 |" in result


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
