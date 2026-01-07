"""
Unit tests for fill_quality_standards.py
"""

import pytest
import os
import sys
import tempfile
import logging
from docx import Document

# Add scripts directory to path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
scripts_dir = os.path.join(project_root, '.claude', 'skills', 'ra-doc-assit', 'scripts')
sys.path.insert(0, scripts_dir)

from fill_quality_standards import (
    parse_markdown_table_from_string,
    parse_markdown_table_from_file,
    restore_formatting_to_cell,
    clear_table_content,
    insert_table_rows,
    find_quality_standards_table,
    merge_cells_in_column,
    auto_merge_duplicate_cells,
    fill_word_document_table,
    fill_quality_standards_from_markdown,
    fill_quality_standards_from_file,
    fill_quality_standards_inplace
)

# Configure logging for tests
logging.basicConfig(level=logging.DEBUG)


class TestParseMarkdownTableFromString:
    """Test parse_markdown_table_from_string function"""

    def test_parse_simple_table(self):
        """Test parsing of simple markdown table"""
        markdown = """| 类型 | 检验项目 | 检验方法 | 质量标准 |
| --- | --- | --- | --- |
| 鉴别 | 外观 | 目视 | 澄清 |
"""
        result = parse_markdown_table_from_string(markdown)
        assert len(result) == 1
        assert result[0] == ['鉴别', '外观', '目视', '澄清']

    def test_parse_with_header_only(self):
        """Test parsing table with header only"""
        markdown = """| 类型 | 检验项目 |
| --- | --- |
"""
        result = parse_markdown_table_from_string(markdown)
        assert len(result) == 0

    def test_parse_with_multiple_rows(self):
        """Test parsing table with multiple rows"""
        markdown = """| 类型 | 检验项目 |
| --- | --- |
| 鉴别 | 外观 |
| 鉴别 | pH值 |
"""
        result = parse_markdown_table_from_string(markdown)
        assert len(result) == 2
        assert result[0][0] == '鉴别'
        assert result[1][0] == '鉴别'

    def test_parse_with_empty_rows(self):
        """Test parsing table with empty rows"""
        markdown = """| 类型 | 检验项目 |
| --- | --- |
| 鉴别 | 外观 |
| | |
"""
        result = parse_markdown_table_from_string(markdown)
        # Should skip empty rows
        assert len(result) == 1

    def test_parse_with_unicode_formatting(self):
        """Test parsing table with Unicode super/subscript characters"""
        markdown = """| 检验项目 | 质量标准 |
| --- | --- |
| 含量 | ≥95.0% |
| pH值 | 6.0-8.0 |
"""
        result = parse_markdown_table_from_string(markdown)
        assert len(result) == 2

    def test_parse_with_caret_notation(self):
        """Test parsing table with ^ notation for superscript"""
        markdown = """| 检验项目 | 公式 |
| --- | --- |
| 分子式 | H^2O |
"""
        result = parse_markdown_table_from_string(markdown)
        assert len(result) == 1
        assert 'H' in result[0][1]

    def test_parse_invalid_markdown(self):
        """Test parsing invalid markdown"""
        markdown = "Not a table"
        result = parse_markdown_table_from_string(markdown)
        assert len(result) == 0


class TestParseMarkdownTableFromFile:
    """Test parse_markdown_table_from_file function"""

    def test_parse_from_valid_file(self, tmp_path):
        """Test parsing from valid markdown file"""
        md_file = tmp_path / "test.md"
        md_file.write_text("""| 类型 | 检验项目 |
| --- | --- |
| 鉴别 | 外观 |
""", encoding='utf-8')

        result = parse_markdown_table_from_file(str(md_file))
        assert len(result) == 1
        assert result[0] == ['鉴别', '外观']

    def test_parse_from_nonexistent_file(self):
        """Test parsing from non-existent file"""
        with pytest.raises(Exception):
            parse_markdown_table_from_file("nonexistent.md")

    def test_parse_from_empty_file(self, tmp_path):
        """Test parsing from empty file"""
        md_file = tmp_path / "empty.md"
        md_file.write_text("", encoding='utf-8')

        result = parse_markdown_table_from_file(str(md_file))
        assert len(result) == 0


class TestRestoreFormattingToCell:
    """Test restore_formatting_to_cell function"""

    def test_restore_plain_text(self):
        """Test restoring plain text without formatting"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        restore_formatting_to_cell(cell, "Plain Text")
        assert cell.text == "Plain Text"

    def test_restore_superscript_unicode(self):
        """Test restoring Unicode superscript characters"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        restore_formatting_to_cell(cell, "H²O")
        # Check that text is restored
        assert "H" in cell.text
        assert "O" in cell.text

    def test_restore_subscript_unicode(self):
        """Test restoring Unicode subscript characters"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        restore_formatting_to_cell(cell, "H₂O")
        assert "H" in cell.text
        assert "O" in cell.text

    def test_restore_caret_notation(self):
        """Test restoring ^ notation"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        restore_formatting_to_cell(cell, "m^3")
        assert "m" in cell.text
        assert "3" in cell.text

    def test_restore_underscore_notation(self):
        """Test restoring _ notation"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        restore_formatting_to_cell(cell, "H_2O")
        assert "H" in cell.text
        assert "2" in cell.text

    def test_restore_mixed_formatting(self):
        """Test restoring mixed superscript and subscript"""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)

        restore_formatting_to_cell(cell, "H²SO₄")
        assert len(cell.text) > 0


class TestClearTableContent:
    """Test clear_table_content function"""

    def test_clear_with_header(self):
        """Test clearing table while keeping header"""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = 'Header1'
        table.rows[1].cells[0].text = 'Data1'
        table.rows[2].cells[0].text = 'Data2'

        clear_table_content(table, keep_header=True)

        assert len(table.rows) == 1
        assert table.rows[0].cells[0].text == 'Header1'

    def test_clear_without_header(self):
        """Test clearing entire table"""
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = 'Header1'
        table.rows[1].cells[0].text = 'Data1'

        clear_table_content(table, keep_header=False)

        assert len(table.rows) == 0


class TestInsertTableRows:
    """Test insert_table_rows function"""

    def test_insert_single_row(self):
        """Test inserting a single row"""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)

        insert_table_rows(table, 1)

        assert len(table.rows) == 2

    def test_insert_multiple_rows(self):
        """Test inserting multiple rows"""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)

        insert_table_rows(table, 3)

        assert len(table.rows) == 4

    def test_insert_zero_rows(self):
        """Test inserting zero rows"""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)

        insert_table_rows(table, 0)

        assert len(table.rows) == 1


class TestFindQualityStandardsTable:
    """Test find_quality_standards_table function"""

    def test_find_quality_table(self):
        """Test finding quality standards table"""
        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        table.rows[0].cells[0].text = '类型'
        table.rows[0].cells[1].text = '检验项目'
        table.rows[0].cells[2].text = '检验方法'
        table.rows[0].cells[3].text = '质量标准'

        result = find_quality_standards_table(doc)
        assert result is not None
        assert result == 0

    def test_find_no_quality_table(self):
        """Test when no quality table exists"""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = 'A'
        table.rows[0].cells[1].text = 'B'

        result = find_quality_standards_table(doc)
        assert result is None

    def test_find_multiple_tables(self):
        """Test finding quality table among multiple tables"""
        doc = Document()
        # Add non-quality table
        table1 = doc.add_table(rows=1, cols=2)
        table1.rows[0].cells[0].text = 'A'
        table1.rows[0].cells[1].text = 'B'

        # Add quality table
        table2 = doc.add_table(rows=2, cols=4)
        table2.rows[0].cells[0].text = '检验项目'
        table2.rows[0].cells[1].text = '检验方法'
        table2.rows[0].cells[2].text = '质量标准'

        result = find_quality_standards_table(doc)
        assert result == 1


class TestMergeCellsInColumn:
    """Test merge_cells_in_column function"""

    def test_merge_valid_range(self):
        """Test merging cells in valid range"""
        doc = Document()
        table = doc.add_table(rows=4, cols=2)
        for i in range(4):
            table.rows[i].cells[0].text = f'Text{i}'

        merge_cells_in_column(table, 0, 0, 2)

        # First cell should be merged (strip trailing whitespace)
        assert table.rows[0].cells[0].text.strip() == 'Text0'

    def test_merge_invalid_range(self):
        """Test merging with invalid range"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # Should not raise error, just log
        merge_cells_in_column(table, 0, 1, 0)  # start > end

    def test_merge_out_of_bounds(self):
        """Test merging with out of bounds indices"""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # Should handle gracefully
        merge_cells_in_column(table, 0, 0, 10)


class TestAutoMergeDuplicateCells:
    """Test auto_merge_duplicate_cells function"""

    def test_auto_merge_type_column(self):
        """Test auto-merging in 类型 column"""
        doc = Document()
        table = doc.add_table(rows=4, cols=4)
        # Header
        table.rows[0].cells[0].text = '类型'
        table.rows[0].cells[1].text = '检验项目'
        # Data
        table.rows[1].cells[0].text = '鉴别'
        table.rows[2].cells[0].text = '鉴别'
        table.rows[3].cells[0].text = '检查'

        auto_merge_duplicate_cells(table, ['类型'])

        # Should have merged 鉴别 cells (strip trailing whitespace)
        assert table.rows[1].cells[0].text.strip() == '鉴别'

    def test_auto_merge_item_column(self):
        """Test auto-merging in 检验项目 column"""
        doc = Document()
        table = doc.add_table(rows=4, cols=4)
        # Header
        table.rows[0].cells[0].text = '类型'
        table.rows[0].cells[1].text = '检验项目'
        # Data
        table.rows[1].cells[0].text = '鉴别'
        table.rows[1].cells[1].text = '外观'
        table.rows[2].cells[0].text = '鉴别'
        table.rows[2].cells[1].text = '外观'
        table.rows[3].cells[0].text = '鉴别'
        table.rows[3].cells[1].text = 'pH值'

        auto_merge_duplicate_cells(table, ['类型', '检验项目'])

        # Should have merged 外观 cells within same 类型 (strip trailing whitespace)
        assert table.rows[1].cells[1].text.strip() == '外观'

    def test_auto_merge_empty_table(self):
        """Test auto-merging with empty table"""
        doc = Document()
        table = doc.add_table(rows=1, cols=2)

        # Should handle gracefully
        auto_merge_duplicate_cells(table, ['类型'])


class TestFillWordDocumentTable:
    """Test fill_word_document_table function"""

    @pytest.fixture
    def template_doc(self, tmp_path):
        """Create a template document with quality table"""
        doc = Document()
        table = doc.add_table(rows=2, cols=4)
        table.rows[0].cells[0].text = '类型'
        table.rows[0].cells[1].text = '检验项目'
        table.rows[0].cells[2].text = '检验方法'
        table.rows[0].cells[3].text = '质量标准'
        table.rows[1].cells[0].text = 'Old Data'

        doc_path = tmp_path / "template.docx"
        doc.save(str(doc_path))
        return str(doc_path)

    def test_fill_table_successfully(self, template_doc, tmp_path):
        """Test successfully filling table"""
        output_path = tmp_path / "output.docx"
        table_data = [
            ['鉴别', '外观', '目视', '澄清'],
            ['鉴别', 'pH值', 'pH计', '6.0-8.0']
        ]

        result = fill_word_document_table(
            template_doc,
            str(output_path),
            table_data,
            table_index=0,
            auto_merge=False
        )

        assert "Successfully" in result
        assert os.path.exists(str(output_path))

    def test_fill_with_auto_merge(self, template_doc, tmp_path):
        """Test filling with auto-merge enabled"""
        output_path = tmp_path / "output_merged.docx"
        table_data = [
            ['鉴别', '外观', '目视', '澄清'],
            ['鉴别', 'pH值', 'pH计', '6.0-8.0']
        ]

        result = fill_word_document_table(
            template_doc,
            str(output_path),
            table_data,
            table_index=0,
            auto_merge=True
        )

        assert "Successfully" in result

    def test_fill_with_invalid_table_index(self, template_doc, tmp_path):
        """Test filling with invalid table index"""
        output_path = tmp_path / "output.docx"
        table_data = [['Data']]

        result = fill_word_document_table(
            template_doc,
            str(output_path),
            table_data,
            table_index=10
        )

        assert "Error" in result

    def test_fill_with_empty_data(self, template_doc, tmp_path):
        """Test filling with empty data"""
        output_path = tmp_path / "output.docx"
        table_data = []

        result = fill_word_document_table(
            template_doc,
            str(output_path),
            table_data
        )

        assert "Error" in result


class TestFillQualityStandardsFromMarkdown:
    """Test fill_quality_standards_from_markdown function"""

    @pytest.fixture
    def template_doc(self, tmp_path):
        """Create a template document"""
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        table.rows[0].cells[0].text = '类型'
        table.rows[0].cells[1].text = '检验项目'
        table.rows[0].cells[2].text = '检验方法'
        table.rows[0].cells[3].text = '质量标准'

        doc_path = tmp_path / "template.docx"
        doc.save(str(doc_path))
        return str(doc_path)

    def test_fill_from_valid_markdown(self, template_doc, tmp_path):
        """Test filling from valid markdown"""
        output_path = tmp_path / "output.docx"
        markdown = """| 类型 | 检验项目 | 检验方法 | 质量标准 |
| --- | --- | --- | --- |
| 鉴别 | 外观 | 目视 | 澄清 |
"""

        result = fill_quality_standards_from_markdown(
            template_doc,
            str(output_path),
            markdown
        )

        assert "Successfully" in result

    def test_fill_from_invalid_markdown(self, template_doc, tmp_path):
        """Test filling from invalid markdown"""
        output_path = tmp_path / "output.docx"
        markdown = "Not a table"

        result = fill_quality_standards_from_markdown(
            template_doc,
            str(output_path),
            markdown
        )

        assert "Error" in result


class TestFillQualityStandardsFromFile:
    """Test fill_quality_standards_from_file function"""

    @pytest.fixture
    def template_doc(self, tmp_path):
        """Create a template document"""
        doc = Document()
        table = doc.add_table(rows=1, cols=4)
        table.rows[0].cells[0].text = '类型'
        table.rows[0].cells[1].text = '检验项目'

        doc_path = tmp_path / "template.docx"
        doc.save(str(doc_path))
        return str(doc_path)

    def test_fill_from_valid_file(self, template_doc, tmp_path):
        """Test filling from valid markdown file"""
        output_path = tmp_path / "output.docx"
        md_file = tmp_path / "data.md"
        md_file.write_text("""| 类型 | 检验项目 |
| --- | --- |
| 鉴别 | 外观 |
""", encoding='utf-8')

        result = fill_quality_standards_from_file(
            template_doc,
            str(output_path),
            str(md_file)
        )

        assert "Successfully" in result

    def test_fill_from_nonexistent_file(self, template_doc, tmp_path):
        """Test filling from non-existent file"""
        output_path = tmp_path / "output.docx"

        result = fill_quality_standards_from_file(
            template_doc,
            str(output_path),
            "nonexistent.md"
        )

        assert "Error" in result


class TestFillQualityStandardsInplace:
    """Test fill_quality_standards_inplace function"""

    def test_fill_inplace(self, tmp_path):
        """Test in-place modification"""
        # Create template
        doc = Document()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = '类型'
        table.rows[0].cells[1].text = '检验项目'

        doc_path = tmp_path / "template.docx"
        doc.save(str(doc_path))

        markdown = """| 类型 | 检验项目 |
| --- | --- |
| 鉴别 | 外观 |
"""

        result = fill_quality_standards_inplace(str(doc_path), markdown)

        assert "Successfully" in result
        # File should still exist
        assert os.path.exists(str(doc_path))


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
