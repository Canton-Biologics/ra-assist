"""
使用 Markdown 格式的数据填充 Word 文档表格中的质量标准
基于fill_word_table.py和merge_word_cells.py的功能，实现将markdown格式的质量标准表格填写到指定docx文件中
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import logging
import traceback
from typing import List, Tuple, Optional, Dict

# 获取此模块的日志记录器
logger = logging.getLogger(__name__)

def parse_markdown_table_from_string(markdown_content: str) -> Tuple[List[List[str]], List[str]]:
    """
    从字符串内容中解析 Markdown 表格以提取数据

    Args:
        markdown_content: Markdown 表格内容字符串

    Returns:
        元组：(数据行列表, 列头列表)
        数据行：每行是单元格值列表
        列头：列名列表
    """
    logger.debug(f"从字符串内容解析 Markdown 表格")

    data = []
    headers = []
    lines = markdown_content.strip().split('\n')

    # 查找表格行（跳过带有 --- 的标题分隔符行）
    table_started = False
    header_found = False

    for line_num, line in enumerate(lines):
        line = line.strip()
        logger.debug(f"处理第 {line_num} 行: '{line}'")

        if line.startswith('|') and line.endswith('|'):
            if '---' in line:
                # 这是分隔符行，跳过
                logger.debug(f"跳过分隔符行: {line}")
                continue

            # 解析表格行
            cells = [cell.strip() for cell in line.split('|')[1:-1]]  # 移除首尾空元素

            if not header_found:
                # 这是标题行
                logger.debug(f"找到标题行: {cells}")
                headers = cells
                header_found = True
                # 存储标题以供参考，但不包含在数据中
                continue
            else:
                # 这是数据行
                if cells and any(cell.strip() for cell in cells):  # 跳过空行
                    logger.debug(f"找到数据行: {cells}")
                    data.append(cells)

    logger.info(f"从 Markdown 表格解析了 {len(data)} 行数据，列头: {headers}")
    return data, headers

def parse_markdown_table_from_file(md_file_path: str) -> Tuple[List[List[str]], List[str]]:
    """
    从文件中解析 Markdown 表格以提取数据

    Args:
        md_file_path: Markdown 文件路径

    Returns:
        元组：(数据行列表, 列头列表)
    """
    logger.info(f"从文件读取 Markdown 表格: {md_file_path}")

    try:
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return parse_markdown_table_from_string(content)
    except Exception as e:
        logger.error(f"读取 Markdown 文件 {md_file_path} 时出错: {str(e)}")
        raise

def restore_formatting_to_cell(cell, text: str):
    """
    将 Unicode 字符的上标和下标格式恢复为 Word 格式

    Args:
        cell: Word 表格单元格
        text: 包含 Unicode 上标/下标字符的文本
    """
    try:
        # 清除现有内容
        cell.text = ""
        paragraph = cell.paragraphs[0]

        # 用于将 Unicode 转换回普通字符的映射表
        superscript_map = {
            '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4', '⁵': '5',
            '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9', '⁺': '+', '⁻': '-',
            '⁼': '=', '⁽': '(', '⁾': ')', 'ⁿ': 'n'
        }

        subscript_map = {
            '₀': '0', '₁': '1', '₂': '2', '₃': '3', '₄': '4', '₅': '5',
            '₆': '6', '₇': '7', '₈': '8', '₉': '9', '₊': '+', '₋': '-',
            '₌': '=', '₍': '(', '₎': ')', 'ₐ': 'a', 'ₑ': 'e', 'ᵢ': 'i',
            'ₒ': 'o', 'ᵤ': 'u', 'ₓ': 'x', 'ₕ': 'h', 'ₖ': 'k', 'ₗ': 'l',
            'ₘ': 'm', 'ₙ': 'n', 'ₚ': 'p', 'ₛ': 's', 'ₜ': 't'
        }

        # 同时处理 ^{} 和 _{} 符号
        # ^{text} 和 _{text} 的模式
        super_pattern = r'\^\{([^}]+)\}'
        sub_pattern = r'_\{([^}]+)\}'

        # 首先替换 ^{text} 和 _{text} 模式
        def replace_super(match):
            return ''.join(superscript_map.get(c, f'^{c}') for c in match.group(1))

        def replace_sub(match):
            return ''.join(subscript_map.get(c, f'_{c}') for c in match.group(1))

        text = re.sub(super_pattern, replace_super, text)
        text = re.sub(sub_pattern, replace_sub, text)

        # 逐个字符处理以应用格式
        i = 0
        while i < len(text):
            char = text[i]

            if char in superscript_map:
                # 添加上标字符
                run = paragraph.add_run(superscript_map[char])
                run.font.superscript = True
            elif char in subscript_map:
                # 添加下标字符
                run = paragraph.add_run(subscript_map[char])
                run.font.subscript = True
            elif char == '^' and i + 1 < len(text):
                # 处理 ^字符 符号
                i += 1
                next_char = text[i]
                run = paragraph.add_run(next_char)
                run.font.superscript = True
            elif char == '_' and i + 1 < len(text):
                # 处理 _字符 符号
                i += 1
                next_char = text[i]
                run = paragraph.add_run(next_char)
                run.font.subscript = True
            else:
                # 普通字符
                run = paragraph.add_run(char)

            i += 1

    except Exception as e:
        logger.warning(f"恢复单元格格式时出错，使用纯文本: {str(e)}")
        # 回退到纯文本
        cell.text = text

def clear_table_content(table, keep_header: bool = True):
    """
    清除表格中的所有行（可选保留标题行）

    Args:
        table: Word 表格对象
        keep_header: 是否保留第一行（标题）
    """
    start_row = 1 if keep_header else 0
    rows_to_remove = len(table.rows) - start_row

    logger.debug(f"清除 {rows_to_remove} 行现有数据（保留标题: {keep_header}）...")

    # 从末尾开始删除行以避免索引问题
    for _ in range(rows_to_remove):
        if len(table.rows) > start_row:
            # 删除最后一行
            table._tbl.remove(table.rows[-1]._tr)

    logger.debug(f"表现在有 {len(table.rows)} 行")

def insert_table_rows(table, num_rows: int):
    """
    在表格中插入指定数量的空行

    Args:
        table: Word 表格对象
        num_rows: 要插入的行数
    """
    logger.debug(f"插入 {num_rows} 个新空行...")

    for _ in range(num_rows):
        # 向表格添加新行
        new_row = table.add_row()
        # 用空字符串初始化单元格
        for cell in new_row.cells:
            cell.text = ""

    logger.debug(f"表现在共有 {len(table.rows)} 行")

def create_column_mapping(source_headers: List[str], target_headers: List[str]) -> Dict[int, int]:
    """
    创建源列到目标列的映射关系

    Args:
        source_headers: 源表格（Markdown）的列头
        target_headers: 目标表格（Word）的列头

    Returns:
        字典：{源列索引: 目标列索引}
    """
    mapping = {}

    for source_idx, source_col in enumerate(source_headers):
        # 在目标列头中查找匹配的列
        if source_col in target_headers:
            target_idx = target_headers.index(source_col)
            mapping[source_idx] = target_idx
            logger.debug(f"列映射: '{source_col}' (源索引 {source_idx}) -> 目标索引 {target_idx}")
        else:
            logger.warning(f"源列 '{source_col}' (索引 {source_idx}) 在目标表格中未找到匹配")

    logger.info(f"创建了 {len(mapping)} 个列映射: {mapping}")
    return mapping

def find_quality_standards_table(doc: Document) -> Optional[int]:
    """
    在文档中查找质量标准表格（增强版）

    检测策略：
    1. 必须是4列表格
    2. 标题行至少匹配3个关键词
    3. 优先选择包含完整4个核心列名的表格

    Args:
        doc: Word 文档对象

    Returns:
        如果找到返回表格索引，否则返回 None
    """
    logger.debug(f"在 {len(doc.tables)} 个表格中搜索质量标准表格")

    # 核心列名（完整匹配）
    core_columns = ['类型', '检验项目', '检验方法', '质量标准']
    quality_keywords = ['检验项目', '检验方法', '质量标准', '类型', '项目', '方法', '标准']

    best_match_idx = None
    best_score = 0

    for table_idx, table in enumerate(doc.tables):
        try:
            # 验证1：必须有标题行
            if len(table.rows) == 0:
                continue

            # 验证2：必须是4列（关键！）
            if len(table.columns) != 4:
                logger.debug(f"表格 {table_idx}: 列数为 {len(table.columns)}，跳过（需要4列）")
                continue

            header_row = table.rows[0]
            header_text = ' '.join(cell.text for cell in header_row.cells).lower()

            # 验证3：关键词匹配（至少3个）
            keyword_count = sum(1 for keyword in quality_keywords if keyword in header_text)
            if keyword_count < 3:
                logger.debug(f"表格 {table_idx}: 仅匹配 {keyword_count} 个关键词，跳过（需要≥3）")
                continue

            # 验证4：计算完整列名匹配数（加分项）
            full_column_matches = sum(1 for col in core_columns if col in header_text)

            # 综合评分
            score = keyword_count + full_column_matches * 2  # 完整列名匹配权重更高

            logger.debug(f"表格 {table_idx}: 关键词={keyword_count}, 完整列名={full_column_matches}, 评分={score}")

            if score > best_score:
                best_score = score
                best_match_idx = table_idx

        except Exception as e:
            logger.warning(f"检查表格 {table_idx} 时出错: {str(e)}")
            continue

    if best_match_idx is not None:
        logger.info(f"在索引 {best_match_idx} 处找到质量标准表格（评分: {best_score}）")
        return best_match_idx

    logger.warning("未找到符合条件的质量标准表格")
    return None

def merge_cells_in_column(table, col_index: int, start_row: int, end_row: int):
    """
    合并列中从 start_row 到 end_row（包含）的单元格

    注意：此函数通过在合并前清除后续单元格的内容来防止重复内容

    Args:
        table: Word 表格对象
        col_index: 要合并的列索引
        start_row: 起始行索引
        end_row: 结束行索引
    """
    if start_row >= end_row or end_row >= len(table.rows):
        logger.debug(f"跳过列 {col_index}、行 {start_row}-{end_row} 的合并（无效范围）")
        return

    try:
        logger.debug(f"合并列 {col_index}，行 {start_row}-{end_row}")

        # 获取要合并的第一个单元格
        first_cell = table.rows[start_row].cells[col_index]

        # 存储第一个单元格的原始内容
        original_content = first_cell.text.strip()

        # 清除要合并的单元格内容以防止重复
        for row_idx in range(start_row + 1, end_row + 1):
            if row_idx < len(table.rows):
                cell = table.rows[row_idx].cells[col_index]
                cell.text = ""  # 合并前清除内容

        # 现在合并单元格（它们是空的，所以没有重复内容）
        for row_idx in range(start_row + 1, end_row + 1):
            if row_idx < len(table.rows):
                cell_to_merge = table.rows[row_idx].cells[col_index]
                first_cell.merge(cell_to_merge)

        # 确保合并的单元格具有正确的内容
        if first_cell.text.strip() != original_content:
            first_cell.text = original_content

        logger.debug(f"成功合并列 {col_index}，行 {start_row}-{end_row}")

    except Exception as e:
        logger.warning(f"无法合并列 {col_index} 中的单元格 {start_row}-{end_row}: {str(e)}")
        # 如果合并失败，至少清除后续单元格中的重复文本
        try:
            for row_idx in range(start_row + 1, end_row + 1):
                if row_idx < len(table.rows):
                    cell = table.rows[row_idx].cells[col_index]
                    cell.text = ""
        except Exception as e2:
            logger.warning(f"无法清除重复文本: {str(e2)}")

def auto_merge_duplicate_cells(table, target_columns: List[str] = ['类型', '检验项目']):
    """
    自动合并指定列中具有重复内容的单元格

    Args:
        table: Word 表格对象
        target_columns: 要检查合并的列名列表
    """
    if len(table.rows) <= 1:
        logger.debug("行数不足以进行合并")
        return

    # 获取标题行以查找列索引
    header_row = table.rows[0]
    column_mapping = {}

    for col_idx, cell in enumerate(header_row.cells):
        header_text = cell.text.strip()
        for target_col in target_columns:
            if target_col in header_text:
                column_mapping[target_col] = col_idx
                break

    logger.debug(f"用于合并的列映射: {column_mapping}")

    # 合并"类型"列
    if '类型' in column_mapping:
        col_idx = column_mapping['类型']
        logger.debug(f"处理「类型」列（索引 {col_idx}）进行合并")

        current_type = ""
        merge_start = -1

        for row_idx in range(1, len(table.rows)):  # 跳过标题
            cell_text = table.rows[row_idx].cells[col_idx].text.strip()

            if cell_text != current_type:
                # 发现不同的类型，如需要则合并上一组
                if merge_start != -1 and row_idx - merge_start > 1:
                    merge_cells_in_column(table, col_idx, merge_start, row_idx - 1)

                current_type = cell_text
                merge_start = row_idx

        # 处理最后一组
        if merge_start != -1 and len(table.rows) - merge_start > 1:
            merge_cells_in_column(table, col_idx, merge_start, len(table.rows) - 1)

    # 在相同的"类型"内合并"检验项目"列
    if '检验项目' in column_mapping and '类型' in column_mapping:
        type_col_idx = column_mapping['类型']
        item_col_idx = column_mapping['检验项目']

        logger.debug(f"处理「检验项目」列（索引 {item_col_idx}）在相同「类型」内进行合并")

        current_type = ""
        current_item = ""
        item_start = -1

        for row_idx in range(1, len(table.rows)):
            type_text = table.rows[row_idx].cells[type_col_idx].text.strip()
            item_text = table.rows[row_idx].cells[item_col_idx].text.strip()

            if type_text != current_type:
                # 不同的类型，如需要则合并上一项目组
                if item_start != -1 and row_idx - item_start > 1 and current_item:
                    merge_cells_in_column(table, item_col_idx, item_start, row_idx - 1)

                current_type = type_text
                current_item = item_text
                item_start = row_idx
            elif item_text == current_item and item_text != "" and current_item != "":
                # 相同类型中的相同项目 - 继续该组
                continue
            else:
                # 相同类型中的不同项目，如需要则合并上一组
                if item_start != -1 and row_idx - item_start > 1 and current_item:
                    merge_cells_in_column(table, item_col_idx, item_start, row_idx - 1)

                current_item = item_text
                item_start = row_idx

        # 处理最后一组
        if item_start != -1 and len(table.rows) - item_start > 1 and current_item:
            merge_cells_in_column(table, item_col_idx, item_start, len(table.rows) - 1)

def fill_word_document_table(doc_path: str, output_path: str, table_data: List[List[str]],
                           source_headers: List[str],
                           table_index: Optional[int] = None,
                           auto_merge: bool = True) -> str:
    """
    使用质量标准数据填充 Word 文档表格

    Args:
        doc_path: 输入 Word 文档路径
        output_path: 保存输出文档的路径
        table_data: 行列表，每行是单元格值列表
        source_headers: 源表格的列头列表（用于列映射）
        table_index: 要填充的特定表格索引（None 表示自动检测）
        auto_merge: 是否自动合并重复单元格

    Returns:
        成功消息或错误描述
    """
    try:
        logger.info(f"加载文档: {doc_path}")
        doc = Document(doc_path)

        # 查找目标表格
        if table_index is None:
            table_index = find_quality_standards_table(doc)
            if table_index is None:
                return "错误：文档中未找到质量标准表格"

        if table_index >= len(doc.tables):
            return f"错误：表格索引 {table_index} 在文档中未找到（仅有 {len(doc.tables)} 个表格）"

        target_table = doc.tables[table_index]
        logger.info(f"使用表格 {table_index}，有 {len(target_table.rows)} 行和 {len(target_table.columns)} 列")

        # 验证数据
        if not table_data:
            return "错误：未提供表格数据"

        # 获取目标表格的列头
        target_headers = [cell.text.strip() for cell in target_table.rows[0].cells]
        logger.info(f"源表格列头: {source_headers}")
        logger.info(f"目标表格列头: {target_headers}")

        # 创建列映射
        column_mapping = create_column_mapping(source_headers, target_headers)

        if not column_mapping:
            return "错误：无法创建列映射（源列和目标列不匹配）"

        logger.info(f"用 {len(table_data)} 行数据填充表格")

        # 步骤 1：清除现有表格内容（标题除外）
        clear_table_content(target_table, keep_header=True)

        # 步骤 2：插入所需数量的行
        num_data_rows = len(table_data)
        insert_table_rows(target_table, num_data_rows)

        # 步骤 3：用数据填充表格（使用列映射）
        for i, row_data in enumerate(table_data):
            row_index = i + 1  # 跳过标题行
            if row_index < len(target_table.rows):
                row = target_table.rows[row_index]

                # 使用列映射填充单元格
                for source_col_idx, target_col_idx in column_mapping.items():
                    if source_col_idx < len(row_data) and target_col_idx < len(row.cells):
                        cell_text = row_data[source_col_idx]
                        cell = row.cells[target_col_idx]

                        # 恢复单元格内容的格式
                        restore_formatting_to_cell(cell, cell_text)

                        logger.debug(f"填充单元格 [{row_index}, {target_col_idx}]: '{cell_text[:50]}...'")

            else:
                logger.warning(f"跳过行 {row_index}，表格没有足够的行")

        # 步骤 4：如果需要，自动合并重复单元格
        if auto_merge:
            logger.info("执行自动单元格合并...")
            # 使用目标表格的列头来确定需要合并的列
            merge_columns = [col for col in ['类型', '检验项目'] if col in target_headers]
            if merge_columns:
                auto_merge_duplicate_cells(target_table, merge_columns)

        # 步骤 5：保存修改后的文档
        doc.save(output_path)
        logger.info(f"文档保存到: {output_path}")

        return f"成功用 {len(table_data)} 行填充表格并保存到 {output_path}"

    except Exception as e:
        error_msg = f"填充 Word 文档表格时出错: {str(e)}"
        logger.error(error_msg)
        logger.error(f"完整回溯: {traceback.format_exc()}")
        return error_msg

def fill_quality_standards_from_markdown(doc_path: str, output_path: str, markdown_content: str,
                                       table_index: Optional[int] = None, auto_merge: bool = True) -> str:
    """
    使用 Markdown 内容中的质量标准数据填充 Word 文档表格

    Args:
        doc_path: 输入 Word 文档路径
        output_path: 保存输出文档的路径
        markdown_content: Markdown 表格内容字符串
        table_index: 要填充的特定表格索引（None 表示自动检测）
        auto_merge: 是否自动合并重复单元格

    Returns:
        成功消息或错误描述
    """
    try:
        logger.info("解析 Markdown 表格内容...")
        table_data, source_headers = parse_markdown_table_from_string(markdown_content)

        if not table_data:
            return "错误：在 Markdown 内容中未找到有效的表格数据"

        return fill_word_document_table(doc_path, output_path, table_data, source_headers, table_index, auto_merge=auto_merge)

    except Exception as e:
        error_msg = f"处理 Markdown 内容时出错: {str(e)}"
        logger.error(error_msg)
        logger.error(f"完整回溯: {traceback.format_exc()}")
        return error_msg

def fill_quality_standards_from_file(doc_path: str, output_path: str, markdown_file_path: str,
                                   table_index: Optional[int] = None, auto_merge: bool = True) -> str:
    """
    使用 Markdown 文件中的质量标准数据填充 Word 文档表格

    Args:
        doc_path: 输入 Word 文档路径
        output_path: 保存输出文档的路径
        markdown_file_path: 包含表格数据的 Markdown 文件路径
        table_index: 要填充的特定表格索引（None 表示自动检测）
        auto_merge: 是否自动合并重复单元格

    Returns:
        成功消息或错误描述
    """
    try:
        logger.info(f"读取 Markdown 文件: {markdown_file_path}")
        table_data, source_headers = parse_markdown_table_from_file(markdown_file_path)

        if not table_data:
            return "错误：在 Markdown 文件中未找到有效的表格数据"

        return fill_word_document_table(doc_path, output_path, table_data, source_headers, table_index, auto_merge=auto_merge)

    except Exception as e:
        error_msg = f"处理 Markdown 文件时出错: {str(e)}"
        logger.error(error_msg)
        logger.error(f"完整回溯: {traceback.format_exc()}")
        return error_msg

def fill_quality_standards_inplace(doc_path: str, markdown_content: str,
                                 table_index: Optional[int] = None, auto_merge: bool = True) -> str:
    """
    使用 Markdown 内容中的质量标准数据填充 Word 文档表格（就地修改文件）

    Args:
        doc_path: 要就地修改的 Word 文档路径
        markdown_content: Markdown 表格内容字符串
        table_index: 要填充的特定表格索引（None 表示自动检测）
        auto_merge: 是否自动合并重复单元格

    Returns:
        成功消息或错误描述
    """
    return fill_quality_standards_from_markdown(doc_path, doc_path, markdown_content, table_index, auto_merge)