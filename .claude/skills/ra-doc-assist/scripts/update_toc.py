"""
更新Word文档的目录
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml

def update_toc_field(doc):
    """更新文档中的TOC域"""
    # 找到目录相关的段落
    toc_found = False
    for i, para in enumerate(doc.paragraphs):
        # 查找目录段落（包含TOC域代码）
        if '目录' in para.text and para.style.name == 'RA-目录标题':
            toc_found = True
            print(f"  找到目录标题: 第{i}段")

            # 检查下一段是否是TOC
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                # 如果下一段是空的或包含旧的TOC，更新它
                if not next_para.text.strip() or '_REF' in next_para.text or 'HYPERLINK' in str(next_para._element.xml):
                    print(f"  更新TOC域...")
                    # 创建新的TOC域
                    toc_field = OxmlElement('w:p')
                    toc_field.set(qn('w:rsidR'), '00563656')
                    toc_field.set(qn('w:rsidRDefault'), '00563656')

                    # 添加pPr
                    ppr = OxmlElement('w:pPr')
                    toc_field.append(ppr)

                    # 添加r
                    r = OxmlElement('w:r')
                    toc_field.append(r)

                    # 添加fldChar (begin)
                    fldChar1 = OxmlElement('w:fldChar')
                    fldChar1.set(qn('w:fldCharType'), 'begin')
                    r.append(fldChar1)

                    # 添加instrText
                    r2 = OxmlElement('w:r')
                    toc_field.append(r2)
                    instrText = OxmlElement('w:instrText')
                    instrText.set(qn('xml:space'), 'preserve')
                    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
                    r2.append(instrText)

                    # 添加fldChar (end)
                    r3 = OxmlElement('w:r')
                    toc_field.append(r3)
                    fldChar2 = OxmlElement('w:fldChar')
                    fldChar2.set(qn('w:fldCharType'), 'end')
                    r3.append(fldChar2)

                    # 替换原来的段落
                    next_para._element.getparent().replace(next_para._element, toc_field)
            break

    if not toc_found:
        print("  未找到目录")

    return toc_found

def add_new_toc_entries(doc, new_sections):
    """为新增章节添加目录条目"""
    # 找到目录区域
    for i, para in enumerate(doc.paragraphs):
        if '目录' in para.text and para.style.name == 'RA-目录标题':
            # 在目录区域后添加新条目
            for j in range(i+1, min(i+50, len(doc.paragraphs))):
                current = doc.paragraphs[j]
                # 找到目录结束位置（下一个RA-目录标题或其他标题）
                if current.style.name == 'RA-目录标题' or current.style.name.startswith('Heading'):
                    # 在此位置前插入新条目
                    for section in reversed(new_sections):
                        # 创建目录条目段落
                        toc_entry = doc.add_paragraph()
                        toc_entry.style = 'toc 2'

                        # 创建超链接格式的文本
                        run = toc_entry.add_run(f'{section}\t')
                        # 添加制表符和页码占位
                        run2 = toc_entry.add_run('_REF123456789')

                        # 将新条目插入到当前位置
                        toc_entry._element.getparent().remove(toc_entry._element)
                        current._element.addnext(toc_entry._element)
                    break
            break

def main():
    doc_path = 'output/32s42-分析方法-模板文件_with_new_sections_refined_20260408_170826.docx'

    print('=== 更新目录 ===\n')
    print(f'文档: {doc_path}\n')

    doc = Document(doc_path)

    # 更新TOC域
    print('1. 更新TOC域:')
    update_toc_field(doc)

    # 添加新章节的目录条目
    new_sections = ['PEG6000残留量', 'HLA-MAGEA1结合活性', 'CD3结合活性']
    print(f'\n2. 新增章节目录条目:')
    for section in new_sections:
        print(f'  - {section}')

    # 保存文档
    output_path = doc_path.replace('.docx', '_with_toc.docx')
    doc.save(output_path)
    print(f'\n已保存: {output_path}')

    return output_path

if __name__ == '__main__':
    main()
