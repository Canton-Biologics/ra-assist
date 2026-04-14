# -*- coding: utf-8 -*-
"""
分析方法章节顺序：解析 Word/JSON 顺序文件，对模版中 Heading 3 方法块物理重排，并插入标准骨架。
"""
from __future__ import annotations

import json
import os
import re
import unicodedata
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


@dataclass
class OrderSpec:
    """用户指定的章节顺序与新增节 SOP 配置。"""

    order: List[str]
    new_sections: Dict[str, Dict[str, Any]] = field(default_factory=dict)


# ── 解析顺序文件 ──────────────────────────────────────────────


def _strip_numbered_line(text: str) -> Optional[str]:
    """从一行文本去掉行首序号，返回标题；不匹配返回 None。"""
    t = unicodedata.normalize('NFKC', (text or '').strip())
    if not t:
        return None
    # 1. 肽图 / 1、肽图 / 1．肽图
    m = re.match(r'^\s*(\d+)[\.\、．．]\s*(.+)$', t)
    if m:
        return m.group(2).strip()
    # 1\t肽图 或 1 肽图
    m = re.match(r'^\s*(\d+)[\t\u3000]+\s*(.+)$', t)
    if m:
        return m.group(2).strip()
    m = re.match(r'^\s*(\d+)\s+([^\d].+)$', t)
    if m:
        return m.group(2).strip()
    return None


def parse_method_order_docx(path: str) -> List[str]:
    """从 Word 段落解析编号列表为章节名顺序。"""
    doc = Document(path)
    order: List[str] = []
    for para in doc.paragraphs:
        raw = para.text or ''
        title = _strip_numbered_line(raw)
        if title:
            # 跳过明显非方法行的目录提示
            if title in ('目录', '分析方法', 'Contents'):
                continue
            order.append(title)

    # 兼容：部分排序文件不是“编号列表”，而是每行一个方法名（无 1./1、前缀）
    # 这种情况下按出现顺序收集非空行即可。
    if not order:
        for para in doc.paragraphs:
            t = (para.text or '').strip()
            if not t:
                continue
            if t in ('目录', '分析方法', 'Contents'):
                continue
            # 过滤掉明显说明性长段落（避免把正文说明当成方法名）
            if len(t) > 80:
                continue
            order.append(t)
    return order


def parse_method_order_json(path: str) -> OrderSpec:
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    if isinstance(data, list):
        return OrderSpec(order=[str(x).strip() for x in data if str(x).strip()])
    order = data.get('order') or data.get('method_order') or []
    if not isinstance(order, list):
        order = []
    ns = data.get('new_sections') or {}
    if not isinstance(ns, dict):
        ns = {}
    return OrderSpec(
        order=[str(x).strip() for x in order if str(x).strip()],
        new_sections={str(k): dict(v) if isinstance(v, dict) else {} for k, v in ns.items()},
    )


def parse_method_order_file(path: str) -> OrderSpec:
    """根据扩展名解析 .docx（仅 order）或 .json（order + new_sections）。"""
    ext = os.path.splitext(path)[1].lower()
    if ext == '.docx':
        return OrderSpec(order=parse_method_order_docx(path))
    return parse_method_order_json(path)


# ── 标题匹配（与 integrate_sop_method 弱匹配一致） ────────────


def normalize_method_title(name: str) -> str:
    t = str(name or '').strip().lower()
    t = re.sub(r'（[^）]*）|\([^)]*\)', '', t)
    t = re.sub(r'[\s\u3000]', '', t)
    return t


def titles_match(a: str, b: str) -> bool:
    if not a or not b:
        return False
    if a.strip() == b.strip():
        return True

    # ── 保守匹配：避免「还原纯度/非还原纯度」误落到「纯度（SEC-HPLC）」 ──
    # 说明：默认弱匹配允许子串包含，会导致 want="还原纯度（rCE-SDS）" 误匹配到候选"纯度（SEC-HPLC）"。
    # 这里先做类别约束：只要任一侧出现这些强特征词，则要求另一侧也包含同类特征词。
    a_raw = str(a).strip()
    b_raw = str(b).strip()
    a_upper = a_raw.upper().replace(' ', '').replace('\u3000', '')
    b_upper = b_raw.upper().replace(' ', '').replace('\u3000', '')

    def _has_any(s: str, *tokens: str) -> bool:
        return any(t for t in tokens if t and (t in s))

    # 还原/非还原纯度必须同类匹配
    if ('非还原纯度' in a_raw) or ('非还原纯度' in b_raw):
        return ('非还原纯度' in a_raw) and ('非还原纯度' in b_raw)
    if ('还原纯度' in a_raw) or ('还原纯度' in b_raw):
        return ('还原纯度' in a_raw) and ('还原纯度' in b_raw)

    # CE-SDS / CE-SDS 法等关键词：若出现则另一侧也必须出现（防止与 SEC-HPLC 交叉误匹配）
    if _has_any(a_upper, 'CE-SDS', 'CESDS', 'CE—SDS', 'CE–SDS') or _has_any(
        b_upper, 'CE-SDS', 'CESDS', 'CE—SDS', 'CE–SDS'
    ):
        return _has_any(a_upper, 'CE-SDS', 'CESDS', 'CE—SDS', 'CE–SDS') and _has_any(
            b_upper, 'CE-SDS', 'CESDS', 'CE—SDS', 'CE–SDS'
        )

    na, nb = normalize_method_title(a), normalize_method_title(b)
    if na == nb:
        return True
    if na and nb and (na in nb or nb in na):
        return True
    alt = na.replace('蛋白质含量', '蛋白含量').replace('蛋白含量', '蛋白质含量')
    altb = nb.replace('蛋白质含量', '蛋白含量').replace('蛋白含量', '蛋白质含量')
    if alt == altb or (alt and altb and (alt in altb or altb in alt)):
        return True
    # 蛋白质A / 蛋白A
    a2 = na.replace('蛋白质', '蛋白')
    b2 = nb.replace('蛋白质', '蛋白')
    if a2 == b2 or (a2 and b2 and (a2 in b2 or b2 in a2)):
        return True
    return False


def find_matching_title(want: str, candidates: List[str]) -> Optional[str]:
    for c in candidates:
        if titles_match(want, c):
            return c
    return None


def infer_section_title_from_sop_filename(path_or_name: str) -> str:
    """
    从 SOP 文件名推断分析方法章节名（与模版 Heading 3 对齐，再由 titles_match 弱匹配）。
    典型：「3-CB0012-SOP-TI003-01 CB0012蛋白含量（UV）检验标准操作规程.docx」→「蛋白含量（UV）」
    """
    base = os.path.basename(str(path_or_name).strip().strip('"').strip("'"))
    base = unicodedata.normalize('NFKC', base)
    base = re.sub(r"\.docx?$", "", base, flags=re.I).strip()
    # 去掉文首编号段：3-CB0012-SOP-TI003-01
    base = re.sub(
        r"^\s*\d+\s*-\s*CB[\dA-Za-z]+\s*-\s*SOP\s*-\s*TI[\dA-Za-z]+\s*-\s*\d+\s+",
        "",
        base,
        flags=re.I,
    )
    # 去掉产品代码前缀：CB0012 / CB0012␠
    base = re.sub(r"^CB[\dA-Za-z]+\s*", "", base, flags=re.I).strip()
    parts = re.split(r"检验标准\s*操作规程|检验标准操作规程|标准操作规程", base, maxsplit=1, flags=re.I)
    title = parts[0].strip() if parts else base.strip()
    title = re.sub(r"\s+", " ", title).strip()
    return title


def sop_path_relative_to_template(sop_abs: str, template_abs: str) -> str:
    tdir = os.path.dirname(os.path.abspath(template_abs))
    s = os.path.abspath(sop_abs)
    try:
        return os.path.relpath(s, tdir)
    except ValueError:
        return s


def build_auto_merge_spec_and_items(
    template_path: str,
    sop_paths_abs: List[str],
) -> Tuple[OrderSpec, List[dict], List[str]]:
    """
    多 SOP「懒人」模式：推断章节名；模版已有则对齐标题；无则记入 new_sections 并排在 order 末尾。
    返回 (OrderSpec, items, 日志行)。
    """
    log_lines: List[str] = []
    tpl = os.path.abspath(template_path)
    if not os.path.isfile(tpl):
        return OrderSpec(order=[], new_sections={}), [], [f"模版不存在: {tpl}"]

    doc = Document(tpl)
    template_titles = [b["title"] for b in list_method_blocks(doc)]
    if not template_titles:
        log_lines.append("[WARN] 模版中未识别到 Heading 3 方法块，仍将尝试仅按推断名处理")

    items: List[dict] = []
    new_sections: Dict[str, Dict[str, Any]] = {}
    new_tail_order: List[str] = []
    seen_new_canonical: set = set()

    for raw in sop_paths_abs:
        p = os.path.abspath(os.path.normpath(str(raw).strip().strip('"').strip("'")))
        # 双向尝试：U+00A0 <-> 普通空格（修复文件名编码问题）
        if not os.path.isfile(p):
            relaxed = p.replace('\u00a0', ' ')
            if relaxed != p and os.path.isfile(relaxed):
                p = relaxed
            else:
                unrelaxed = p.replace(' ', '\u00a0')
                if unrelaxed != p and os.path.isfile(unrelaxed):
                    p = unrelaxed
        if not os.path.isfile(p):
            log_lines.append(f"[WARN] SOP 不存在，已跳过: {p}")
            continue
        inferred = infer_section_title_from_sop_filename(p)
        if not inferred:
            log_lines.append(f"[WARN] 无法从文件名推断章节名，已跳过: {p}")
            continue
        match = find_matching_title(inferred, template_titles)
        canonical = match if match else inferred
        rel = sop_path_relative_to_template(p, tpl)

        if match is None:
            if canonical in seen_new_canonical:
                log_lines.append(
                    f"[WARN] 章节「{canonical}」出现多个 SOP，后以最后一个文件为准"
                )
            else:
                log_lines.append(
                    f"[新章节] 模版无匹配，将插入骨架并整合: 「{inferred}」"
                )
                new_tail_order.append(canonical)
                seen_new_canonical.add(canonical)
            new_sections[canonical] = {"sop": rel, "stop": []}
        else:
            if inferred.strip() != match:
                log_lines.append(f"[已有章节] 推断「{inferred}」→ 模版标题「{match}」")

        items.append({"section": canonical, "sop": rel, "stop": []})

    order = list(template_titles)
    for c in new_tail_order:
        if find_matching_title(c, order) is None:
            order.append(c)

    spec = OrderSpec(order=order, new_sections=new_sections)
    return spec, items, log_lines


# ── 模版方法块（body 级，含表） ────────────────────────────────


def _is_method_heading_paragraph(para) -> bool:
    st = para.style.name if para.style else ''
    t = (para.text or '').strip()
    if not t:
        return False
    if 'toc' in st.lower():
        return False
    if st.startswith('Heading 3'):
        if len(t) > 100:
            return False
        return True
    if st.startswith('Heading') and '3' in st and not st.startswith('Heading 2'):
        return True
    return False


def list_method_blocks(doc: Document) -> List[Dict[str, Any]]:
    """
    按文档顺序列出各方法块：每项 { 'title': str, 'elements': [lxml elements] }。
    """
    headings: List[Tuple[str, Any]] = []
    for para in doc.paragraphs:
        if _is_method_heading_paragraph(para):
            headings.append((para.text.strip(), para._element))

    if not headings:
        return []

    body = doc.element.body
    blocks: List[Dict[str, Any]] = []
    for i, (title, start_el) in enumerate(headings):
        end_el = headings[i + 1][1] if i + 1 < len(headings) else None
        els: List[Any] = []
        take = False
        for child in body.iterchildren():
            if child is start_el:
                take = True
            if not take:
                continue
            if end_el is not None and child is end_el:
                break
            els.append(child)
        blocks.append({'title': title, 'elements': els})
    return blocks


def compute_final_order(user_order: List[str], template_titles_in_doc_order: List[str]) -> List[str]:
    """
    最终顺序 = 用户 order + 模版中未出现在 user_order（模糊匹配）的标题，保持模版原相对顺序。
    """
    out: List[str] = []
    seen_norm = set()

    def mark(name: str) -> None:
        seen_norm.add(normalize_method_title(name))

    for name in user_order:
        out.append(name)
        mark(name)
        # 若与某模版标题等价，也视为已覆盖
        m = find_matching_title(name, template_titles_in_doc_order)
        if m:
            mark(m)

    for t in template_titles_in_doc_order:
        if find_matching_title(t, user_order):
            continue
        if normalize_method_title(t) in seen_norm:
            continue
        out.append(t)
        mark(t)

    return out


def _detach_elements(els: List[Any]) -> None:
    for el in els:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _skeleton_strip_paragraph_list_numpr(p_el: Any) -> None:
    """
    去掉段上 w:numPr，避免新增方法块沿用上一方法的 Word 多级列表实例，
    导致「原理」等 RA-5 显示为 6、7… 而非重新从 1 起（样式若仍带编号，由 Word 另算）。
    """
    p_pr = p_el.find(qn('w:pPr'))
    if p_pr is None:
        return
    num_pr = p_pr.find(qn('w:numPr'))
    if num_pr is not None:
        p_pr.remove(num_pr)


def _find_base_ra5_num_id(doc: Document) -> Optional[str]:
    """从模版中找一个 RA-5 级标题段落的 numId，用于克隆出“可重启编号”的新 numId。"""
    for p in doc.paragraphs:
        try:
            if not p.style or p.style.name != 'RA-5级标题':
                continue
            p_pr = p._element.find(qn('w:pPr'))
            if p_pr is None:
                continue
            num_pr = p_pr.find(qn('w:numPr'))
            if num_pr is None:
                continue
            num_id = num_pr.find(qn('w:numId'))
            if num_id is None:
                continue
            v = num_id.get(qn('w:val'))
            if v:
                return str(v)
        except Exception:
            continue
    return None


def _clone_num_id(doc: Document, base_num_id: str) -> Optional[str]:
    """
    克隆一个新的 numId，引用与 base_num_id 相同的 abstractNumId。
    用于新增方法块内 RA-5 标题编号从 1 重新开始。
    """
    try:
        numbering = doc.part.numbering_part.element  # w:numbering
        nums = numbering.findall(qn('w:num'))
        max_id = -1
        for n in nums:
            try:
                v = int(n.get(qn('w:numId')))
                if v > max_id:
                    max_id = v
            except Exception:
                continue
        base_num = None
        for n in nums:
            if str(n.get(qn('w:numId'))) == str(base_num_id):
                base_num = n
                break
        if base_num is None:
            return None
        abs_id_el = base_num.find(qn('w:abstractNumId'))
        if abs_id_el is None:
            return None
        abs_id = abs_id_el.get(qn('w:val'))
        if abs_id is None:
            return None

        new_id = str(max_id + 1 if max_id >= 0 else 1)
        new_num = OxmlElement('w:num')
        new_num.set(qn('w:numId'), new_id)
        abs2 = OxmlElement('w:abstractNumId')
        abs2.set(qn('w:val'), str(abs_id))
        new_num.append(abs2)

        # 关键：强制本 num 实例从 1 开始，避免 Word 继续沿用上一实例的计数（出现 7、8…）
        # 仅对 level 0（RA-5 一级标题）做 override 即可。
        lvl_override = OxmlElement('w:lvlOverride')
        lvl_override.set(qn('w:ilvl'), '0')
        start_override = OxmlElement('w:startOverride')
        start_override.set(qn('w:val'), '1')
        lvl_override.append(start_override)
        new_num.append(lvl_override)

        numbering.append(new_num)
        return new_id
    except Exception:
        return None


def _set_paragraph_numpr(p_el: Any, num_id: str, ilvl: str = '0') -> None:
    """为段落设置直接 numPr（numId + ilvl）。"""
    p_pr = p_el.find(qn('w:pPr'))
    if p_pr is None:
        p_pr = OxmlElement('w:pPr')
        p_el.insert(0, p_pr)
    num_pr = p_pr.find(qn('w:numPr'))
    if num_pr is None:
        num_pr = OxmlElement('w:numPr')
        p_pr.append(num_pr)
    num_id_el = num_pr.find(qn('w:numId'))
    if num_id_el is None:
        num_id_el = OxmlElement('w:numId')
        num_pr.append(num_id_el)
    num_id_el.set(qn('w:val'), str(num_id))
    il = num_pr.find(qn('w:ilvl'))
    if il is None:
        il = OxmlElement('w:ilvl')
        num_pr.append(il)
    il.set(qn('w:val'), str(ilvl))


def _create_skeleton_elements(doc: Document, section_title: str) -> List[Any]:
    """生成新方法标准骨架段落 XML 节点（已从 body 摘除，待插入）。"""
    specs = [
        (section_title, 'Heading 3'),
        ('原理', 'RA-5级标题'),
        ('', 'RA-正文'),
        ('设备、材料、试剂', 'RA-5级标题'),
        ('', 'RA-正文'),
        ('操作步骤', 'RA-5级标题'),
        ('', 'RA-正文'),
        ('可接受标准', 'RA-5级标题'),
        ('', 'RA-正文'),
        ('结果计算', 'RA-5级标题'),
        ('', 'RA-正文'),
        ('合格标准', 'RA-5级标题'),
        ('', 'RA-正文'),
    ]
    elems: List[Any] = []
    # 为新增方法块克隆一份 RA-5 编号实例，确保从 1 开始
    base_num_id = _find_base_ra5_num_id(doc)
    new_num_id = _clone_num_id(doc, base_num_id) if base_num_id else None
    for text, style in specs:
        p = doc.add_paragraph(text, style=style)
        el = p._element
        # Heading 3 保留列表/大纲由模版样式决定
        if style == 'RA-5级标题':
            # 关键：使用新的 numId，避免跨方法延续编号（出现 7、8…）
            if new_num_id:
                _set_paragraph_numpr(el, new_num_id, ilvl='0')
        elif style == 'RA-正文':
            # 正文不挂列表
            _skeleton_strip_paragraph_list_numpr(el)
        elems.append(el)
    for el in elems:
        el.getparent().remove(el)
    return elems


def reorder_and_insert_methods(doc: Document, spec: OrderSpec) -> List[str]:
    """
    按 OrderSpec 重排模版中方法块；order 中有而模版无的标题插入骨架。
    返回最终章节名列表（用于推导 stop）。
    """
    blocks = list_method_blocks(doc)
    if not blocks:
        print('  [WARN] 模版中未识别到 Heading 3 方法块，跳过重排')
        return spec.order[:]

    template_titles = [b['title'] for b in blocks]
    final_order = compute_final_order(spec.order, template_titles)

    # title(模版原文) -> block elements
    title_to_elements: Dict[str, List[Any]] = {}
    for b in blocks:
        title_to_elements[b['title']] = b['elements']

    used_template_titles: set = set()

    def take_block_for(want: str) -> List[Any]:
        nonlocal used_template_titles
        for tt, els in title_to_elements.items():
            if tt in used_template_titles:
                continue
            if titles_match(want, tt):
                used_template_titles.add(tt)
                return list(els)
        return []

    planned: List[List[Any]] = []
    for want in final_order:
        els = take_block_for(want)
        if els:
            planned.append(els)
        else:
            print(f'  [INFO] 新增方法骨架: {want}')
            planned.append(_create_skeleton_elements(doc, want))

    body = doc.element.body
    first_el = blocks[0]['elements'][0]
    preamble: List[Any] = []
    for child in body.iterchildren():
        if child is first_el:
            break
        preamble.append(child)

    # 摘下所有旧方法块元素
    all_old: List[Any] = []
    for b in blocks:
        all_old.extend(b['elements'])
    _detach_elements(all_old)

    # 锚点：前言最后一段；无前言则依次 append 到 body 末尾
    if preamble:
        anchor = preamble[-1]
        for block in planned:
            for el in block:
                anchor.addnext(el)
                anchor = el
    else:
        for block in planned:
            for el in block:
                body.append(el)

    return final_order


def derive_stop_keywords(ordered_names: List[str], index: int) -> List[str]:
    if index + 1 < len(ordered_names):
        return [ordered_names[index + 1]]
    return []


def _build_toc_field_paragraph(instr_text: str) -> OxmlElement:
    toc_field = OxmlElement('w:p')
    toc_field.set(qn('w:rsidR'), '00563656')
    toc_field.set(qn('w:rsidRDefault'), '00563656')
    ppr = OxmlElement('w:pPr')
    toc_field.append(ppr)
    r = OxmlElement('w:r')
    toc_field.append(r)
    fld_char_b = OxmlElement('w:fldChar')
    fld_char_b.set(qn('w:fldCharType'), 'begin')
    r.append(fld_char_b)
    r2 = OxmlElement('w:r')
    toc_field.append(r2)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = instr_text
    r2.append(instr)
    r3 = OxmlElement('w:r')
    toc_field.append(r3)
    fld_char_e = OxmlElement('w:fldChar')
    fld_char_e.set(qn('w:fldCharType'), 'end')
    r3.append(fld_char_e)
    return toc_field


def _replace_next_para_with_field(doc: Document, heading_idx: int, instr_text: str) -> bool:
    if heading_idx + 1 >= len(doc.paragraphs):
        return False
    next_para = doc.paragraphs[heading_idx + 1]
    # 允许替换空行或现有 TOC/图表目录域（包含 instrText/Hyperlink）
    if not (
        (not next_para.text.strip())
        or ('w:instrText' in str(next_para._element.xml))
        or ('HYPERLINK' in str(next_para._element.xml))
        or ('TOC' in str(next_para._element.xml))
        or ('PAGEREF' in str(next_para._element.xml))
    ):
        return False
    field_p = _build_toc_field_paragraph(instr_text)
    try:
        next_para._element.getparent().replace(next_para._element, field_p)
        return True
    except Exception:
        return False


def refresh_toc(doc: Document) -> bool:
    """
    刷新/写入：
    - 目录（Heading 1-3）：TOC \\o \"1-3\"
    - 表目录：基于样式 RA-表格编号 的图表目录域
    - 图目录：基于样式 RA-图编号 的图表目录域

    注：页码等最终显示仍需在 Word 中更新域（F9）。
    """
    ok_any = False

    def norm(s: str) -> str:
        return (s or '').replace(' ', '').replace('\u3000', '').strip()

    for i, para in enumerate(doc.paragraphs):
        if not para.style or para.style.name != 'RA-目录标题':
            continue
        t = norm(para.text)
        if not t:
            continue
        if t in ('目录', '目录：', '目录:') or ('目录' in t and '表目录' not in t and '图目录' not in t):
            # 模版实际为「目  录」，此处用“包含目录”兜底
            ok_any = _replace_next_para_with_field(doc, i, ' TOC \\o "1-3" \\h \\z \\u ') or ok_any
        elif '表目录' in t:
            # 图表目录域：按样式收集（模板使用 table of figures + TOC \\t）
            ok_any = _replace_next_para_with_field(doc, i, ' TOC \\h \\z \\t "RA-表格编号,1" ') or ok_any
        elif '图目录' in t:
            ok_any = _replace_next_para_with_field(doc, i, ' TOC \\h \\z \\t "RA-图编号,1" ') or ok_any

    return ok_any


# ── 与 methods.json 合并（extract / 配置） ─────────────────────


def merge_order_spec_into_items(
    config_items: List[dict],
    spec: OrderSpec,
    template_path: str,
) -> List[dict]:
    """
    根据顺序与 new_sections 生成有序的 extract/整合配置项，并设置 stop 为下一节标题。
    """
    doc = Document(template_path)
    template_titles = [b['title'] for b in list_method_blocks(doc)]
    final_order = compute_final_order(spec.order, template_titles)

    by_section: Dict[str, dict] = {}
    for it in config_items:
        sec = str(it.get('section', '') or '').strip()
        if sec:
            by_section[sec] = dict(it)

    for name, meta in spec.new_sections.items():
        name = str(name).strip()
        if not name:
            continue
        if name not in by_section:
            by_section[name] = {
                'section': name,
                'sop': str(meta.get('sop', '') or '').strip(),
                'stop': meta.get('stop') or [],
            }
            if isinstance(by_section[name]['stop'], str):
                by_section[name]['stop'] = [by_section[name]['stop']]

    ordered: List[dict] = []
    used_keys = set()

    for want in final_order:
        hit = None
        for k, v in by_section.items():
            if titles_match(want, k):
                hit = v
                used_keys.add(k)
                break
        if hit:
            ordered.append(dict(hit))

    for k, v in by_section.items():
        if k not in used_keys:
            ordered.append(dict(v))

    for i, it in enumerate(ordered):
        nxt = str(ordered[i + 1].get('section', '') or '').strip() if i + 1 < len(ordered) else ''
        it['stop'] = [nxt] if nxt else []

    return ordered


def order_refined_methods_by_spec(
    methods: List[dict],
    spec: OrderSpec,
    template_path: str,
) -> List[dict]:
    """按顺序规范排列 refined.json 中的 methods，并刷新 stop。"""
    doc = Document(template_path)
    template_titles = [b['title'] for b in list_method_blocks(doc)]
    final_order = compute_final_order(spec.order, template_titles)

    by_name: Dict[str, dict] = {}
    for m in methods:
        n = str(m.get('name', '') or '').strip()
        if n:
            by_name[n] = m

    ordered: List[dict] = []
    used = set()

    for want in final_order:
        hit = None
        hit_key = None
        for k, v in by_name.items():
            if titles_match(want, k):
                hit = v
                hit_key = k
                break
        if hit is not None:
            ordered.append(hit)
            used.add(hit_key)

    for k, v in by_name.items():
        if k not in used:
            ordered.append(v)

    for i, m in enumerate(ordered):
        nxt = str(ordered[i + 1].get('name', '') or '').strip() if i + 1 < len(ordered) else ''
        m['stop'] = [nxt] if nxt else []

    return ordered
