# -*- coding: utf-8 -*-
"""
分析方法整合器

将 SOP 第四章「程序」内容提取后，与分析方法标准模板进行比对整合，
输出含有内部 SOP 内容的分析方法文档。
"""
import os
import re
import shutil
import io
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import parse_xml
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from .sop_extractor import SOPExtractor


# SOP 第四章 → 分析方法 六章节 字段映射
SOP_TO_METHOD = {
    'principle': '1. 原理',
    'materials_and_equipment': '2. 材料和设备',
    'sample_prep': '3. 操作步骤-样品处理',
    'procedure': '3. 操作步骤-测定法',
    'suitability_criteria': '4. 试验成立标准',
    'result_calculation': '5. 结果计算',
    'acceptance_criteria': '6. 合格标准',
}


def _is_section_title(text: str, keywords: List[str]) -> bool:
    """判断是否为章节标题行（应过滤）"""
    t = text.strip()
    if not t or len(t) > 100:
        return False
    for kw in keywords:
        if t == kw or t.rstrip('：:') == kw:
            return True
        if re.match(r'^[\d一二三四五六七八九十]+[\.\s、．]?\s*' + re.escape(kw) + r'\s*$', t):
            return True
    return False


def _filter_content(lines: List[str], skip_titles: List[str]) -> List[str]:
    """过滤内容：去除标题行、SOP 编号引用"""
    result = []
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            continue
        if _is_section_title(cleaned, skip_titles):
            continue
        if cleaned.startswith('3-') and 'SOP' in cleaned and len(cleaned) > 40:
            continue
        result.append(cleaned)
    return result


def _normalize_for_match(text: str) -> str:
    """归一化文本用于来源匹配。"""
    t = str(text or '').strip().lower()
    t = re.sub(r'[\s\u3000]+', '', t)
    t = re.sub(r'[，。；：、,.!！?？（）()\[\]【】\-—_~`\'"“”‘’<>《》/\\]', '', t)
    return t


def _normalize_display_text(text: str) -> str:
    """
    显示兼容归一化：
    - 将部分字体不支持的乘号字符替换为通用字符，避免 Word 显示方框。
    """
    return str(text or '').replace('⨯', '×')


def _line_has_source(line: str, source_pool: List[str]) -> bool:
    """判断单行是否可在来源池中追溯。"""
    n_line = _normalize_for_match(line)
    if not n_line:
        return False
    for src in source_pool:
        n_src = _normalize_for_match(src)
        if not n_src:
            continue
        if n_line in n_src or n_src in n_line:
            return True
    return False


# 整合后常为长行（如「主要设备：A、B…」），无法逐行与 source_pool 子串匹配；勿做行级过滤以免回退成模板原文
# acceptance_criteria：精简结果常为短句（如「应符合规定」），与 SOP 原文子串匹配易失败；勿因过滤整段清空
_SKIP_LINE_SOURCE_FILTER_KEYS = frozenset({'materials_and_equipment', 'principle', 'acceptance_criteria'})


def _filter_refined_by_source(
    refined_content: Dict[str, List[str]],
    source_pool: Optional[List[str]]
) -> Dict[str, List[str]]:
    """
    过滤无来源的 refined 句子，仅保留可追溯内容。
    若 source_pool 为空，返回原始内容（不做硬过滤）。
    """
    if not source_pool:
        return refined_content

    out: Dict[str, List[str]] = {}
    for key, lines in refined_content.items():
        if isinstance(lines, str):
            line_list = [lines] if lines.strip() else []
        elif isinstance(lines, list):
            line_list = [str(x).strip() for x in lines if str(x).strip()]
        else:
            line_list = []
        if key in _SKIP_LINE_SOURCE_FILTER_KEYS:
            out[key] = line_list
            continue
        valid = [line for line in line_list if _line_has_source(line, source_pool)]
        out[key] = valid
    return out


def _split_existing_to_list(existing_text: str) -> List[str]:
    """将模板已有章节文本拆分为行列表。"""
    if not existing_text:
        return []
    return [x.strip() for x in str(existing_text).splitlines() if x.strip()]


def _is_sec_hplc_purity_section(section_name: Optional[str]) -> bool:
    """模版章节「纯度（SEC-HPLC）」等同义写法识别。"""
    if not section_name:
        return False
    t = str(section_name).strip()
    if '纯度' not in t:
        return False
    u = t.upper().replace(' ', '').replace('\u3000', '')
    return 'SEC' in u and 'HPLC' in u


def _is_main_equipment_heading_line(line: str) -> bool:
    s = str(line or '').strip()
    return bool(re.match(r'^主要设备\s*[:：]', s))


def _ensure_material_equipment_lines(
    current_lines: List[str],
    existing_lines: List[str],
    source_pool: Optional[List[str]],
    section_name: Optional[str] = None,
) -> List[str]:
    """
    保证材料和设备中尽量包含 试剂/标准物质/主要设备 三类行。
    先用 current，再用 source_pool（优先 SOP），最后用 existing 补齐。

    特例：章节为纯度（SEC-HPLC）时，「主要设备」行固定为「主要设备：高效液相色谱仪。」
    （与 SKILL/REFERENCE 约定一致；其它试剂/标准物质/材料行仍来自 SOP。）
    """
    out = [str(x).strip() for x in (current_lines or []) if str(x).strip()]
    existing = [str(x).strip() for x in (existing_lines or []) if str(x).strip()]
    source = [str(x).strip() for x in (source_pool or []) if str(x).strip()]

    def _has_reagent_prefixed_line(lines: List[str]) -> bool:
        return any(
            ln.startswith('试剂：') or ln.startswith('主要试剂/耗材')
            for ln in lines
        )

    def pick_line(lines: List[str], *prefixes: str) -> Optional[str]:
        for line in lines:
            for prefix in prefixes:
                if line.startswith(prefix):
                    return line
        return None

    if not any(line.startswith('主要设备：') for line in out):
        candidate = pick_line(source, '主要设备：') or pick_line(existing, '主要设备：')
        if candidate:
            out.append(candidate)
    if not any(line.startswith('标准物质：') for line in out):
        candidate = pick_line(source, '标准物质：') or pick_line(existing, '标准物质：')
        if candidate:
            out.append(candidate)
    if not _has_reagent_prefixed_line(out):
        candidate = (
            pick_line(source, '主要试剂/耗材：', '主要试剂/耗材:')
            or pick_line(source, '试剂：')
            or pick_line(existing, '主要试剂/耗材：', '主要试剂/耗材:')
            or pick_line(existing, '试剂：')
        )
        if candidate:
            out.append(candidate)

    reagent_label = '试剂'
    for line in out:
        if line.startswith('主要试剂/耗材'):
            reagent_label = '主要试剂/耗材'
            break

    # 主要设备行聚合：多行/重复设备名 -> 单行"主要设备：A、B、C。"
    equipment_names: List[str] = []
    standard_lines: List[str] = []
    reagent_items: List[str] = []
    material_lines: List[str] = []  # 新增：材料行
    uncategorized: List[str] = []

    # 第一次遍历：提取各类行
    for line in out:
        if line.startswith('主要设备：'):
            payload = line.split('：', 1)[1] if '：' in line else ''
            parts = re.split(r'[、，,；;。]', payload)
            for p in parts:
                name = p.strip()
                if name:
                    equipment_names.append(name)
        elif line.startswith('标准物质：'):
            standard_lines.append(line)
        elif line.startswith('试剂：') or line.startswith('主要试剂/耗材：') or line.startswith(
            '主要试剂/耗材:'
        ):
            sep = '：' if '：' in line else ':'
            payload = line.split(sep, 1)[1] if sep in line else ''
            parts = re.split(r'[、，,；;。]', payload)
            for p in parts:
                name = p.strip()
                if name:
                    reagent_items.append(name)
        elif line.startswith('材料：'):
            material_lines.append(line)
        else:
            uncategorized.append(line)

    def _dedup_order(items: List[str]) -> List[str]:
        dedup: List[str] = []
        seen = set()
        for n in items:
            k = _normalize_for_match(n)
            if not k or k in seen:
                continue
            seen.add(k)
            dedup.append(n)
        return dedup

    # 模板内呈现顺序：主要设备 -> 标准物质 -> 材料 -> 试剂类（与 refine 整合输出一致）
    processed_lines: List[str] = []
    if equipment_names:
        eq_dedup = _dedup_order(equipment_names)
        if eq_dedup:
            processed_lines.append(f"主要设备：{'、'.join(eq_dedup)}。")

    preferred_std = ''
    for s in source:
        if s.startswith('标准物质：'):
            preferred_std = s if s.endswith('。') else (s + '。')
            break
    if not preferred_std and standard_lines:
        preferred_std = standard_lines[0] if standard_lines[0].endswith('。') else (standard_lines[0] + '。')
    if preferred_std:
        processed_lines.append(preferred_std)

    processed_lines.extend(material_lines)

    if reagent_items:
        reag_dedup = _dedup_order(reagent_items)
        if reag_dedup:
            processed_lines.append(f"{reagent_label}：{'、'.join(reag_dedup)}。")

    processed_lines.extend(uncategorized)

    out = processed_lines

    # 去重
    dedup = []
    seen = set()
    for line in out:
        n = _normalize_for_match(line)
        if n and n not in seen:
            seen.add(n)
            dedup.append(line)

    if section_name and _is_sec_hplc_purity_section(section_name):
        dedup = [line for line in dedup if not _is_main_equipment_heading_line(line)]
        dedup.append('主要设备：高效液相色谱仪。')

    return dedup


def _pick_material_equipment_from_source(source_pool: Optional[List[str]]) -> List[str]:
    """从来源池中提取材料和设备相关行（优先用于 SOP 回填）。"""
    prefixes = ['主要试剂/耗材', '试剂', '标准物质', '主要设备', '材料']
    out = []
    for line in (source_pool or []):
        s = str(line).strip()
        if not s:
            continue
        if any(s.startswith(p) for p in prefixes):
            out.append(s)
    # 去重
    dedup = []
    seen = set()
    for x in out:
        n = _normalize_for_match(x)
        if n and n not in seen:
            seen.add(n)
            dedup.append(x)
    return dedup


def _insert_paragraph_after(paragraph: Paragraph, text: str = '', style=None) -> Paragraph:
    """在当前段落后插入新段落（用于操作步骤区内多段正文，避免挤在同一段内用 \\n 硬换行）。"""
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(str(text))
    if style is not None:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def _paragraph_element_index(doc: Document, para: Paragraph) -> int:
    for idx, p in enumerate(doc.paragraphs):
        if p._element is para._element:
            return idx
    return -1


def _is_protected_style(style_name: str) -> bool:
    """这些样式不应在正文写入时被覆盖。"""
    s = style_name or ''
    return (
        s.startswith('3.2.S') or
        'Heading' in s or
        'RA-5' in s or
        'RA-6' in s or
        'RA-表格编号' in s or
        'RA-图编号' in s
    )


def _field_value_is_effectively_empty(val: Any) -> bool:
    """判断 refined 字段是否应视为空（用于模板回退）。"""
    if val is None:
        return True
    if isinstance(val, str):
        return not val.strip()
    if isinstance(val, list):
        return not val or all(not str(x).strip() for x in val)
    return True


def _collect_writable_indices(doc: Document, start_idx: int, end_idx: int) -> tuple:
    """收集章节范围内可用于写正文的段落索引。

    Returns:
        (可写段落索引列表, 表格标题段落索引列表, 表格索引映射表)
        表格索引映射表: {段落索引: 表格索引}，表示该表格标题段落后面紧跟的表格
    """
    indices: List[int] = []
    table_caption_paras: List[int] = []  # RA-表格编号样式的段落索引
    table_map: Dict[int, int] = {}  # 段落索引 -> 表格索引

    # 首先扫描整个文档，记录所有表格的位置
    para_to_table = {}  # 段落索引 -> 后面紧跟的表格索引
    body = doc.element.body
    elem_idx = 0
    para_idx = -1
    table_idx = -1
    last_para_idx = -1

    for child in body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == 'p':
            para_idx += 1
            last_para_idx = para_idx
        elif tag == 'tbl':
            table_idx += 1
            if last_para_idx >= 0:
                para_to_table[last_para_idx] = table_idx
                last_para_idx = -1
        elem_idx += 1

    # 然后收集可写段落索引
    for i in range(start_idx + 1, end_idx):
        if i >= len(doc.paragraphs):
            break
        para = doc.paragraphs[i]
        style = para.style.name if para.style else ""

        # 检查是否是RA-表格编号样式
        is_table_caption = 'RA-表格编号' in style

        if _is_protected_style(style):
            # 如果是表格标题，记录下来
            if is_table_caption:
                table_caption_paras.append(i)
            continue

        # 含嵌入图的段落保留给后续 _replace_section_images，勿当正文槽位写入或清空
        if _paragraph_has_image(para):
            continue

        indices.append(i)

    return indices, table_caption_paras, para_to_table


def _convert_gradient_rows(rows: List[List[str]]) -> Optional[List[List[str]]]:
    """
    将 SOP 参数表中的“洗脱梯度”行转换为模板梯度表格式：
    [时间(min), 流动相A(%), 流动相B(%), 曲率]
    """
    if not rows or len(rows) < 3:
        return None

    gradient_rows: List[List[str]] = []
    for r in rows:
        if len(r) < 6:
            continue
        if str(r[0]).strip() != '洗脱梯度':
            continue
        time_v = str(r[1]).strip()
        a_v = str(r[3]).strip()
        b_v = str(r[4]).strip()
        curve_v = str(r[5]).strip()
        if time_v in ('时间（min）', '时间(min)', 'Time(min)'):
            continue
        if not time_v:
            continue
        gradient_rows.append([time_v, a_v, b_v, curve_v])

    if not gradient_rows:
        return None

    return [['时间（min）', '流动相A（%）', '流动相B（%）', '曲率']] + gradient_rows


def _replace_table_rows_keep_header(table, rows: List[List[str]]) -> None:
    """替换目标表格的数据行，保留表头行。"""
    if not rows:
        return
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for src_row in rows[1:]:
        new_row = table.add_row()
        max_cols = min(len(new_row.cells), len(src_row))
        for i in range(max_cols):
            new_row.cells[i].text = str(src_row[i])


def _replace_table_with_xml(dst_table, table_xml: str) -> bool:
    """用源表 XML 整表替换目标表（保留样式）。"""
    if not table_xml:
        return False
    try:
        new_tbl = parse_xml(table_xml)
        _ensure_tbl_xml_borders(new_tbl)
        dst_tbl = dst_table._tbl
        dst_tbl.addprevious(new_tbl)
        parent = dst_tbl.getparent()
        parent.remove(dst_tbl)
        return True
    except Exception:
        return False


def _get_table_context_category(doc: Document, table_idx: int) -> str:
    """通过表格前后的段落上下文判断表格类别。

    Returns:
        'suitability_criteria' 或 'acceptance_criteria' 或 ''
    """
    if table_idx >= len(doc.tables):
        return ''

    # 获取表格元素
    table_elem = doc.tables[table_idx]._element

    # 遍历body元素，找到表格的位置，并检查前后段落
    body = doc.element.body
    table_position = None

    # 找到表格在body中的位置
    elem_idx = 0
    for child in body.iterchildren():
        if child == table_elem:
            table_position = elem_idx
            break
        elem_idx += 1

    if table_position is None:
        return ''

    # 检查表格前后各10个元素
    check_range = 10
    start_check = max(0, table_position - check_range)
    end_check = table_position + check_range + 1

    # 收集表格前后的段落文本
    before_paragraphs = []
    after_paragraphs = []

    elem_idx = 0
    for child in body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if start_check <= elem_idx <= end_check:
            if tag == 'p':
                # 找到对应的段落对象
                para = None
                for p in doc.paragraphs:
                    if p._element == child:
                        para = p
                        break
                if para:
                    text = para.text.strip()
                    if elem_idx < table_position:
                        before_paragraphs.append(text)
                    elif elem_idx > table_position:
                        after_paragraphs.append(text)
                        if len(after_paragraphs) >= 5:  # 只收集前5个后续段落
                            break

        elem_idx += 1
        if elem_idx > end_check or (after_paragraphs and len(after_paragraphs) >= 5):
            break

    # 检查前面的段落
    all_before = ' '.join(before_paragraphs[-5:])  # 检查最后5个前面的段落
    all_after = ' '.join(after_paragraphs[:5])  # 检查前5个后面的段落

    # 判断逻辑
    if '试验成立标准' in all_before or '系统适用性可接受标准' in all_before:
        return 'suitability_criteria'
    elif '试验成立标准' in all_after or '系统适用性可接受标准' in all_after:
        # 如果表格后面紧跟"试验成立标准"段落，也可能是
        # 但需要进一步判断
        if '合格标准' in all_after and '结果计算' not in all_after:
            return 'acceptance_criteria'
        return 'suitability_criteria'
    elif '合格标准' in all_before and '试验成立' not in all_before:
        return 'acceptance_criteria'
    elif '合格标准' in all_after and '试验成立' not in all_after and '结果计算' not in all_after:
        return 'acceptance_criteria'
    # 「可接受标准」标题下的两列表为试验成立/检测可接受，不是产品「合格标准」段
    elif '可接受标准' in all_before and '合格标准' not in all_before:
        return 'suitability_criteria'

    return ''


def _is_suitability_target_table(rows: List[List[str]], doc: Document = None, table_idx: int = -1) -> bool:
    """判断是否为系统适用性可接受标准目标表。

    优先根据上下文判断，如果无法判断则根据表格内容。
    """
    # 如果提供了doc和table_idx，优先使用上下文判断
    if doc is not None and table_idx >= 0:
        context_category = _get_table_context_category(doc, table_idx)
        if context_category == 'suitability_criteria':
            return True
        elif context_category == 'acceptance_criteria':
            return False

    # 降级到基于内容的判断
    if not rows:
        return False
    header = " ".join(str(c).strip() for c in rows[0]).lower()
    max_cols = max((len(r) for r in rows), default=0)
    if max_cols != 2:
        return False
    return (
        ('可接受标准' in header and ('样品' in header or '名称' in header)) or
        ('acceptable standard' in header and ('sample' in header or 'name' in header))
    )


def _is_acceptance_target_table(rows: List[List[str]], doc: Document = None, table_idx: int = -1) -> bool:
    """判断是否为产品合格标准目标表。

    优先根据上下文判断，如果无法判断则根据表格内容。
    """
    # 如果提供了doc和table_idx，优先使用上下文判断
    if doc is not None and table_idx >= 0:
        context_category = _get_table_context_category(doc, table_idx)
        if context_category == 'acceptance_criteria':
            return True
        elif context_category == 'suitability_criteria':
            return False

    # 降级：两列「样品/名称+可接受标准」表在 32s42 中属于可接受标准/系统适用性小节，不是产品合格标准段
    if not rows:
        return False
    header = " ".join(str(c).strip() for c in rows[0]).lower()
    max_cols = max((len(r) for r in rows), default=0)
    if max_cols == 2 and (
        ('可接受标准' in header and ('样品' in header or '名称' in header)) or
        ('acceptable standard' in header and ('sample' in header or 'name' in header))
    ):
        return False
    return False


def _insert_table_before_paragraph(doc: Document, anchor_para, rows: List[List[str]]) -> None:
    """在锚点前插入表格。"""
    if not rows:
        return
    max_cols = max((len(r) for r in rows), default=0)
    if max_cols <= 0:
        return
    table = doc.add_table(rows=0, cols=max_cols)
    for src_row in rows:
        row = table.add_row()
        for i in range(min(len(row.cells), len(src_row))):
            row.cells[i].text = str(src_row[i])
    if anchor_para is not None:
        anchor_para._p.addprevious(table._tbl)


def _insert_table_xml_before_paragraph(anchor_para, table_xml: str) -> bool:
    """在锚点前插入源表 XML（保留源表样式）。"""
    if anchor_para is None or not table_xml:
        return False
    try:
        tbl = parse_xml(table_xml)
        _ensure_tbl_xml_borders(tbl)
        anchor_para._p.addprevious(tbl)
        return True
    except Exception:
        return False


def _next_anchor_paragraph_index(
    anchor_pos: Dict[str, int], anchor_name: str, section_end_idx: int
) -> int:
    """当前锚点标题之后、下一 RA-5 锚点标题之前（不含）的段落上界，缺省为章节 end_idx。"""
    order = ['原理', '材料和设备', '操作步骤', '试验成立标准', '结果计算', '合格标准']
    if anchor_name not in order:
        return section_end_idx
    i0 = order.index(anchor_name)
    for j in range(i0 + 1, len(order)):
        nxt = order[j]
        if nxt in anchor_pos:
            return anchor_pos[nxt]
    return section_end_idx


def _has_table_between_anchor_paragraphs(doc: Document, lo: int, hi: int) -> bool:
    """与 _get_table_indices_in_section 一致：上一段索引严格落在 (lo, hi) 的表格。"""
    body = doc.element.body
    para_idx = -1
    for child in body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == 'p':
            para_idx += 1
        elif tag == 'tbl':
            if lo < para_idx < hi:
                return True
    return False


def _insert_unmatched_replacement_tables_in_section(
    doc: Document,
    prepared_tables: List[Dict[str, Any]],
    used_sources: set,
    start_idx: int,
    end_idx: int,
) -> None:
    """
    章节内原无（或未替换成功）占位表时，将仍未消费的 SOP 表补插到对应 RA-5 标题之后。
    解决「仅骨架新方法」无 w:tbl 导致可接受标准/结果计算表无法替换的问题。
    """
    cat_to_anchor = {
        'suitability_criteria': '试验成立标准',
        'result_calculation': '结果计算',
        'acceptance_criteria': '合格标准',
    }
    anchor_pos = _detect_anchor_positions(doc, start_idx, end_idx)
    span_has_table: Dict[str, bool] = {}
    for an in set(cat_to_anchor.values()):
        if an not in anchor_pos:
            continue
        p0 = anchor_pos[an]
        p1 = _next_anchor_paragraph_index(anchor_pos, an, end_idx)
        span_has_table[an] = _has_table_between_anchor_paragraphs(doc, p0, p1)

    last_attach_el: Dict[str, Any] = {}
    for src_i, src in enumerate(prepared_tables):
        if src_i in used_sources:
            continue
        cat = str(src.get('category') or '').strip().lower()
        if cat not in cat_to_anchor:
            continue
        aname = cat_to_anchor[cat]
        if aname not in anchor_pos:
            continue
        if span_has_table.get(aname):
            continue
        base_el = doc.paragraphs[anchor_pos[aname]]._element
        attach_el = last_attach_el.get(aname, base_el)
        rows = src.get('rows', []) or []
        xml = str(src.get('table_xml') or '').strip()
        inserted = False
        if xml:
            try:
                tbl = parse_xml(xml)
                _ensure_tbl_xml_borders(tbl)
                attach_el.addnext(tbl)
                last_attach_el[aname] = tbl
                inserted = True
            except Exception:
                inserted = False
        if not inserted and rows:
            p1 = _next_anchor_paragraph_index(anchor_pos, aname, end_idx)
            if p1 < len(doc.paragraphs):
                tgt = doc.paragraphs[p1]
                _insert_table_before_paragraph(doc, tgt, rows)
                pe = tgt._element
                prev = pe.getprevious()
                while prev is not None:
                    if prev.tag.rsplit('}', 1)[-1] == 'tbl':
                        last_attach_el[aname] = prev
                        break
                    prev = prev.getprevious()
                inserted = True
        if inserted:
            used_sources.add(src_i)
            print(f"  [INFO] 小节「{aname}」内原无表，已补插 SOP 表: category={cat}")


def _ensure_tbl_xml_borders(tbl_elem) -> None:
    """
    为表格 XML 补齐全边框，防止跨文档样式丢失导致无边框。
    """
    try:
        tbl_pr = tbl_elem.find(qn('w:tblPr'))
        if tbl_pr is None:
            tbl_pr = OxmlElement('w:tblPr')
            tbl_elem.insert(0, tbl_pr)

        old = tbl_pr.find(qn('w:tblBorders'))
        if old is not None:
            tbl_pr.remove(old)

        borders = OxmlElement('w:tblBorders')
        for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'), 'single')
            e.set(qn('w:sz'), '4')
            e.set(qn('w:space'), '0')
            e.set(qn('w:color'), 'auto')
            borders.append(e)
        tbl_pr.append(borders)
    except Exception:
        # 边框补齐失败不阻断主流程
        return


def _has_usable_acceptance_replacement(tables: Optional[List[Dict[str, Any]]]) -> bool:
    """SOP 侧是否存在可用于替换模版「合格标准」表的数据（有分类且至少一行数据）。"""
    for tb in tables or []:
        if not isinstance(tb, dict):
            continue
        if str(tb.get('category') or '').strip() != 'acceptance_criteria':
            continue
        rows = tb.get('rows') or []
        if len(rows) < 2:
            continue
        if any(any(str(c or '').strip() for c in row) for row in rows[1:]):
            return True
    return False


def _pick_acceptance_lines_from_source(source_pool: Optional[List[str]]) -> List[str]:
    """从来源池中优先挑选合格标准相关句子。"""
    out = []
    for line in (source_pool or []):
        s = str(line).strip()
        if not s:
            continue
        if any(k in s for k in ['进样针数', '进样体积', '备注', '色谱条件', '操作步骤']):
            continue
        if any(k in s for k in ['不低于', '不高于', '≤', '≥', '无干扰', '一致', '差值', '应']):
            out.append(s)
    # 去重
    dedup = []
    seen = set()
    for x in out:
        n = _normalize_for_match(x)
        if n and n not in seen:
            seen.add(n)
            dedup.append(x)
    return dedup[:6]


def _enforce_required_method_lines(
    refined: Dict[str, List[str]],
    required_source: Optional[Dict[str, Any]]
) -> Dict[str, List[str]]:
    """
    强约束补齐：确保关键章节至少包含 SOP 原文关键句。
    仅补齐缺失，不覆盖已有 refined。
    """
    if not isinstance(required_source, dict):
        return refined

    out: Dict[str, List[str]] = {}
    for key, lines in refined.items():
        if isinstance(lines, str):
            out[key] = [lines] if lines.strip() else []
        elif isinstance(lines, list):
            out[key] = [_normalize_display_text(str(x).strip()) for x in lines if str(x).strip()]
        else:
            out[key] = []

    def source_lines(key: str) -> List[str]:
        val = required_source.get(key, [])
        if isinstance(val, str):
            return [val] if val.strip() else []
        if isinstance(val, list):
            return [str(x).strip() for x in val if str(x).strip()]
        return []

    def append_if_missing(key: str, candidates: List[str], max_add: int = 2) -> None:
        existing_norm = {_normalize_for_match(x) for x in out.get(key, [])}
        added = 0
        for c in candidates:
            n = _normalize_for_match(c)
            if not n or n in existing_norm:
                continue
            out.setdefault(key, []).append(c)
            existing_norm.add(n)
            added += 1
            if added >= max_add:
                break

    # 关键样品处理链
    if len(out.get('sample_prep', [])) < 2:
        cand = [x for x in source_lines('sample_prep') if any(k in x for k in ['供试品', '系统适用性', '空白溶液', 'FB溶液', '离心', '孵育'])]
        append_if_missing('sample_prep', cand, max_add=3)

    # 操作步骤关键锚点（含 UV/蛋白含量等非色谱方法）
    if len(out.get('procedure', [])) < 2:
        proc_kw = [
            '色谱条件', '平衡系统', '进样序列', '进样测试',
            '波长', '280', '吸光度', '分光光度', '紫外', '校正', '测定', '比色皿',
            '读数', '平行测定', '空白校正',
        ]
        cand = [x for x in source_lines('procedure') if any(k in x for k in proc_kw)]
        append_if_missing('procedure', cand, max_add=8)

    # 试验成立标准需有判定语句
    if not any(any(k in x for k in ['RSD', '≤', '≥', '相关系数', '分离度', '无干扰']) for x in out.get('suitability_criteria', [])):
        cand = [x for x in source_lines('suitability_criteria') if any(k in x for k in ['RSD', '≤', '≥', '相关系数', '分离度', '无干扰', '符合要求'])]
        append_if_missing('suitability_criteria', cand, max_add=3)

    # 结果计算防截断：仅回填“计算表达”类语句，避免把操作步骤灌入结果计算
    src_calc = source_lines('result_calculation')
    cur_calc = out.get('result_calculation', [])
    if src_calc:
        calc_keep_keywords = [
            '=', '%', '计算公式', '面积归一法', '主峰%', '聚体%', '低分子量', 'RRT', 'RT',
            'Slope', 'Conc', 'A=', '朗伯-比尔', 'Beer-Lambert'
        ]
        calc_exclude_keywords = [
            '变性还原', '烷基化', '稀释', '酶解', '终止', '孵育', '离心',
            '平衡系统', '进样序列', '运行序列', '进样测试',
            '操作步骤', '使用Empower软件', '申请表', '申请批准'
        ]

        def is_calc_line(line: str) -> bool:
            s = str(line).strip()
            if not s:
                return False
            if any(k in s for k in calc_exclude_keywords):
                return False
            return any(k in s for k in calc_keep_keywords)

        src_calc_filtered = [x for x in src_calc if is_calc_line(x)]

        # 判定为“过度精简/截断”的条件：
        # 1) 当前行数过少；或 2) 缺少核心计算锚点
        lacks_calc_anchor = not any(
            any(k in x for k in ['=', '%', '计算公式', '面积归一法', '主峰%', 'RRT', 'RT'])
            for x in cur_calc
        )
        too_short = (len(cur_calc) <= 1)
        if too_short or lacks_calc_anchor:
            merged: List[str] = []
            seen = set()
            # 先保留当前已有，再补来源，避免完全覆盖用户已有句子
            for x in (cur_calc + src_calc_filtered):
                s = str(x).strip()
                if not s:
                    continue
                n = _normalize_for_match(s)
                if not n or n in seen:
                    continue
                seen.add(n)
                merged.append(s)
            # 限制上限，防止段落膨胀
            out['result_calculation'] = merged[:16]

    # 合格标准至少 1 条
    if not out.get('acceptance_criteria'):
        append_if_missing('acceptance_criteria', source_lines('acceptance_criteria'), max_add=1)

    return out


def _normalize_table_text(text: str) -> str:
    """归一化表格文本用于匹配。"""
    t = str(text or '').strip().lower()
    t = re.sub(r'[\s\u3000]+', '', t)
    t = re.sub(r'[，。；：、,.!！?？（）()\[\]【】\-—_~`\'"“”‘’<>《》/\\]', '', t)
    return t


def _table_feature_set(rows: List[List[str]]) -> set:
    """提取表格特征集合（表头 + 前几行首列）。"""
    if not rows:
        return set()
    features = set()
    header = rows[0] if rows else []
    for cell in header:
        n = _normalize_table_text(cell)
        if n:
            features.add(n)
    for r in rows[1:4]:
        if not r:
            continue
        n = _normalize_table_text(r[0])
        if n:
            features.add(n)
    return features


def _table_match_score(src_rows: List[List[str]], dst_rows: List[List[str]]) -> int:
    """按表头/结构综合计算表格匹配分数。"""
    if not src_rows or not dst_rows:
        return 0
    src_header = " ".join([str(x).strip() for x in src_rows[0]])
    dst_header = " ".join([str(x).strip() for x in dst_rows[0]])
    src_header_norm = _normalize_table_text(src_header)
    dst_header_norm = _normalize_table_text(dst_header)

    score = 0
    if src_header_norm and src_header_norm == dst_header_norm:
        score += 8

    src_features = _table_feature_set(src_rows)
    dst_features = _table_feature_set(dst_rows)
    overlap = len(src_features & dst_features)
    score += overlap * 2

    src_cols = max((len(r) for r in src_rows), default=0)
    dst_cols = max((len(r) for r in dst_rows), default=0)
    if src_cols > 0 and src_cols == dst_cols:
        score += 1

    # 强语义关键词加权：梯度表 / 可接受标准表
    grad_keys = {'梯度', '洗脱', '流动相a', '流动相b', '曲率', '时间'}
    acc_keys = {'可接受标准', '系统适用性', '检验项目', '检验方法', '质量标准', 'rrt', 'rt'}
    src_all = _normalize_table_text(" ".join(" ".join(str(c) for c in row) for row in src_rows))
    dst_all = _normalize_table_text(" ".join(" ".join(str(c) for c in row) for row in dst_rows))
    if any(_normalize_table_text(k) in src_all for k in grad_keys) and any(_normalize_table_text(k) in dst_all for k in grad_keys):
        score += 3
    if any(_normalize_table_text(k) in src_all for k in acc_keys) and any(_normalize_table_text(k) in dst_all for k in acc_keys):
        score += 3
    return score


def _replace_matching_tables(doc: Document, replacement_tables: List[Dict[str, Any]]) -> None:
    """按表头关键词匹配并替换模板表格内容。"""
    if not replacement_tables:
        return
    # 预处理：提取可映射到模板梯度洗脱表的标准化表格
    prepared_tables: List[Dict[str, Any]] = []
    for src in replacement_tables:
        rows = src.get('rows', []) if isinstance(src, dict) else []
        converted = _convert_gradient_rows(rows)
        if converted:
            prepared_tables.append({'header': converted[0], 'rows': converted})
        prepared_tables.append(src)

    used_targets = set()
    for src in prepared_tables:
        rows = src.get('rows', [])
        if not rows:
            continue
        best_idx = -1
        best_score = 0
        for i, table in enumerate(doc.tables):
            if i in used_targets or not table.rows:
                continue
            dst_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            score = _table_match_score(rows, dst_rows)
            if score > best_score:
                best_score = score
                best_idx = i
        if best_idx >= 0 and best_score >= 3:
            _replace_table_rows_keep_header(doc.tables[best_idx], rows)
            used_targets.add(best_idx)


def _get_table_indices_in_section(doc: Document, start_idx: int, end_idx: int) -> List[int]:
    """
    获取指定章节范围内的表格索引（仅正文 body 下）。
    通过遍历文档 body 元素顺序，记录每个表格所在的段落锚点。
    """
    body = doc.element.body
    para_idx = -1
    table_idx = -1
    section_table_indices: List[int] = []

    for child in body.iterchildren():
        tag = child.tag.rsplit('}', 1)[-1]
        if tag == 'p':
            para_idx += 1
        elif tag == 'tbl':
            table_idx += 1
            # 表格出现在 start/end 章节标题之间，视为当前章节表格
            if start_idx < para_idx < end_idx:
                section_table_indices.append(table_idx)

    return section_table_indices


def _replace_matching_tables_in_section(
    doc: Document,
    replacement_tables: List[Dict[str, Any]],
    section_table_indices: List[int],
    start_idx: int,
    end_idx: int
) -> None:
    """仅在指定章节范围内做表格替换，避免跨章节覆盖。"""
    if not replacement_tables:
        return

    def is_gradient_target_table(rows: List[List[str]]) -> bool:
        if not rows:
            return False
        header = " ".join(str(c).strip() for c in rows[0]).lower()
        flat = " ".join(" ".join(str(c).strip() for c in r) for r in rows).lower()
        keys = ['时间', '流动相a', '流动相b', '曲率', 'gradient', 'elution']
        hit = sum(1 for k in keys if k in header or k in flat)
        return hit >= 3

    def is_result_calc_target_table(rows: List[List[str]]) -> bool:
        if not rows:
            return False
        header = " ".join(str(c).strip() for c in rows[0]).lower()
        flat = " ".join(" ".join(str(c).strip() for c in r) for r in rows).lower()
        keys = ['计算', '结果', '公式', '参数', '含量', '浓度', 'result', 'calculation', 'formula', 'slope', 'conc']
        return sum(1 for k in keys if k in header or k in flat) >= 2

    # 复用 _replace_matching_tables 的核心逻辑，限制目标索引
    prepared_tables: List[Dict[str, Any]] = []
    for src in replacement_tables:
        if not isinstance(src, dict):
            continue
        rows = src.get('rows', []) or []
        category = str(src.get('category') or '').strip().lower()
        # 当前处理系统适用性标准 + 可接受标准 + 结果计算 + 梯度洗脱（procedure）表迁移
        if category and category not in ('suitability_criteria', 'acceptance_criteria', 'result_calculation', 'procedure'):
            continue
        if rows and category == 'procedure':
            converted = _convert_gradient_rows(rows)
            if converted:
                rows = converted
            if not is_gradient_target_table(rows):
                continue
            prepared_tables.append({
                'category': 'procedure',
                'rows': rows,
                'table_xml': src.get('table_xml', ''),
                'table_title': src.get('table_title', ''),
            })
            continue
        if rows and category == 'acceptance_criteria' and not _is_acceptance_target_table(rows):
            continue
        if rows and category == 'result_calculation':
            prepared_tables.append({
                'category': 'result_calculation',
                'rows': rows,
                'table_xml': src.get('table_xml', ''),
                'table_title': src.get('table_title', ''),
            })
            continue
        prepared_tables.append(src)
    if not prepared_tables:
        return

    used_sources = set()
    target_set = set(section_table_indices)
    # 目标驱动：逐个模板表格找最佳 SOP 表格，确保"同章节同类表格"优先以 SOP 覆盖
    for idx in section_table_indices:
        if idx not in target_set:
            continue
        dst_table = doc.tables[idx]
        if not dst_table.rows:
            continue
        dst_rows = [[cell.text.strip() for cell in row.cells] for row in dst_table.rows]

        # 优先级：先检查内容特征（梯度表、结果计算表），再检查上下文（suitability、acceptance）
        dst_category = ''
        # 1. 梯度表检查（内容特征优先）
        if is_gradient_target_table(dst_rows):
            dst_category = 'procedure'
        # 2. 结果计算表检查（内容特征优先）
        elif is_result_calc_target_table(dst_rows):
            dst_category = 'result_calculation'
        # 3. 上下文判断（suitability、acceptance）
        elif _is_suitability_target_table(dst_rows, doc, idx):
            dst_category = 'suitability_criteria'
        elif _is_acceptance_target_table(dst_rows, doc, idx):
            dst_category = 'acceptance_criteria'

        if not dst_category:
            continue
        print(f"  [DEBUG] 表格{idx}: 分类={dst_category}, 行数={len(dst_rows)}, 列数={len(dst_rows[0]) if dst_rows else 0}")

        best_src = None
        best_score = 0
        for src_i, src in enumerate(prepared_tables):
            if src_i in used_sources:
                continue
            src_category = str(src.get('category') or '').strip().lower()
            if src_category and src_category != dst_category:
                continue
            src_rows = src.get('rows', [])
            if not src_rows:
                continue
            score = _table_match_score(src_rows, dst_rows)
            if score > best_score:
                best_score = score
                best_src = src_i
        print(f"  [DEBUG] 最佳匹配: src_i={best_src}, score={best_score}")

        if best_src is not None and best_score >= 3:
            src_best = prepared_tables[best_src]
            print(f"  [INFO] 替换表格{idx}: category={dst_category}")
            # 可接受标准/结果计算表优先整表 XML 替换（保留样式）；梯度表仅替换数据
            if dst_category in ('acceptance_criteria', 'result_calculation'):
                if _replace_table_with_xml(dst_table, str(src_best.get('table_xml') or '')):
                    used_sources.add(best_src)
                    continue
            _replace_table_rows_keep_header(dst_table, src_best.get('rows', []))
            used_sources.add(best_src)
            continue

        # 兜底：按列数一致做弱匹配（仅可接受标准表）
        if dst_category != 'acceptance_criteria':
            continue
        dst_cols = max((len(r.cells) for r in dst_table.rows), default=0)
        if dst_cols <= 0:
            continue
        for src_i, src in enumerate(prepared_tables):
            if src_i in used_sources:
                continue
            src_rows = src.get('rows', [])
            if not src_rows:
                continue
            src_cols = max((len(r) for r in src_rows), default=0)
            if src_cols == dst_cols:
                _replace_table_rows_keep_header(dst_table, src_rows)
                used_sources.add(src_i)
                break

    _insert_unmatched_replacement_tables_in_section(
        doc, prepared_tables, used_sources, start_idx, end_idx
    )


def _paragraph_has_image(para) -> bool:
    """检测段落是否含嵌入图（兼容 w:drawing / blip / pic）。"""
    try:
        xml = para._element.xml
        if any(
            x in xml
            for x in ('pic:pic', 'w:drawing', 'a:blip', 'graphic', 'pic:cNvPr')
        ):
            return True
        for run in para.runs:
            rx = run._element.xml
            if 'pic:pic' in rx or 'graphic' in rx or 'w:drawing' in rx:
                return True
    except Exception:
        pass
    return False


def _flatten_section_images_dedup(
    section_images: Optional[Dict[str, List[Dict[str, Any]]]]
) -> List[Dict[str, Any]]:
    """
    合并 SOP 各小节提取的图片列表，按稳定顺序去重（sha1 优先），供与模板图位匹配。
    """
    if not section_images:
        return []
    order = (
        'suitability_criteria',
        'typical_figure',
        'result_calculation',
        'procedure',
        'acceptance_criteria',
    )
    seen_sha: set = set()
    seen_rid: set = set()
    out: List[Dict[str, Any]] = []
    keys_ordered = [k for k in order if k in section_images]
    keys_ordered.extend(k for k in sorted(section_images.keys()) if k not in keys_ordered)
    for k in keys_ordered:
        for item in section_images.get(k) or []:
            if not isinstance(item, dict):
                continue
            sha = str(item.get('sha1', '') or '')
            rid = str(item.get('rel_id', '') or '')
            if sha:
                if sha in seen_sha:
                    continue
                seen_sha.add(sha)
            elif rid:
                if rid in seen_rid:
                    continue
                seen_rid.add(rid)
            else:
                continue
            out.append(item)
    return out


def _extract_image_caption(
    doc: Document,
    image_para_idx: int,
    section_start: int,
    section_end: int
) -> str:
    """从图片周围的段落中提取标题/上下文。

    优先级：
    1. 图片段落的下一个段落（通常图题在图片后面）
    2. 图片段落的前一个段落
    3. 图片段落本身（如果有文本）
    """
    # 检查下一个段落（优先）
    if image_para_idx + 1 < min(section_end, len(doc.paragraphs)):
        next_para = doc.paragraphs[image_para_idx + 1]
        next_text = next_para.text.strip()
        # 如果下一个段落有文本且不包含图片，则认为是图题
        if next_text and not any('pic:pic' in run._element.xml or 'w:drawing' in run._element.xml for run in next_para.runs):
            return next_text

    # 检查前一个段落
    if image_para_idx - 1 >= section_start:
        prev_para = doc.paragraphs[image_para_idx - 1]
        prev_text = prev_para.text.strip()
        if prev_text:
            return prev_text

    # 使用图片段落本身的文本
    current_para = doc.paragraphs[image_para_idx]
    current_text = current_para.text.strip()
    if current_text:
        return current_text

    return ""


def _calculate_caption_similarity(caption1: str, caption2: str) -> float:
    """计算两个标题之间的相似度（0-1之间）。

    使用关键词匹配和编辑距离的组合方法。
    """
    if not caption1 or not caption2:
        return 0.0

    # 转换为小写并去除空格
    c1 = caption1.lower().strip()
    c2 = caption2.lower().strip()

    # 完全匹配
    if c1 == c2:
        return 1.0

    # 包含关系
    if c1 in c2 or c2 in c1:
        return 0.8

    # 提取关键词（去除常见的停用词）
    stopwords = {'的', '是', '在', '和', '与', '或', '等', '图', '图表', '示意图', 'fig', 'figure',
                 '样品', '供试', '检测', '测定', '分析'}

    # 专业术语词典（权重更高）
    technical_terms = {
        # 2字术语
        '系统', '适用', '试验', '色谱', '图谱', '标准', '计算', '结果', '公式',
        '浓度', '含量', '纯度', '活性', '效能', '曲线', '线性', '关系',
        '峰', '保留', '时间', '面积', '高度', '梯度', '洗脱', '流动',
        '波长', '吸收', '吸光', '光程', '斜率', '截距', '相关', '系数',
        '肽图', '典型', '适应',
        # 3字术语
        '适用性', '流动相', '洗脱', '保留时', '色谱图', '测定图', '验图',
        '适应性',
        # 4字术语
        '系统适用', '适用性试', '浓度计算', '结果计算', '线性关系',
        '纯度检测', '含量测定', '标准曲线', '相关系数', '验图谱', '典型图谱',
    }

    def extract_keywords(text):
        """提取关键词，支持中文和英文，使用最大匹配算法"""
        import re

        keywords = set()

        # 1. 优先匹配专业术语（最大正向匹配）
        remaining_text = text
        i = 0
        while i < len(remaining_text):
            matched = False
            # 尝试匹配4字、3字、2字术语
            for length in [4, 3, 2]:
                if i + length <= len(remaining_text):
                    candidate = remaining_text[i:i+length]
                    # 检查是否是中文且在专业术语表中
                    if re.match(r'^[\u4e00-\u9fa5]+$', candidate) and candidate in technical_terms:
                        keywords.add(candidate)
                        i += length
                        matched = True
                        break
            if not matched:
                i += 1

        # 2. 提取其他中文词组（2-4个连续的汉字）
        chinese_pattern = re.compile(r'[\u4e00-\u9fa5]{2,4}')
        chinese_words = chinese_pattern.findall(text)
        keywords.update(chinese_words)

        # 3. 提取英文单词（2个字符以上）
        english_pattern = re.compile(r'[a-zA-Z]{2,}')
        english_words = english_pattern.findall(text)
        keywords.update([w.lower() for w in english_words])

        # 4. 提取数字和单位的组合（如"2.0%"、"pH7.0"）
        number_pattern = re.compile(r'\d+\.?\d*[a-zA-Z%]|pH[\d.]+')
        number_words = number_pattern.findall(text)
        keywords.update(number_words)

        # 过滤停用词，但保留专业术语
        filtered = []
        for w in keywords:
            if len(w) <= 1:
                continue
            # 保留专业术语
            if w in technical_terms:
                filtered.append(w)
            # 保留不在停用词表中的词
            elif w not in stopwords:
                filtered.append(w)

        return set(filtered)

    keywords1 = extract_keywords(c1)
    keywords2 = extract_keywords(c2)

    if not keywords1 or not keywords2:
        return 0.0

    # 计算关键词重叠度
    intersection = keywords1 & keywords2
    union = keywords1 | keywords2

    if not union:
        return 0.0

    # Jaccard相似度
    jaccard = len(intersection) / len(union)

    # 计算专业术语匹配数（权重更高）
    tech_intersection = intersection & technical_terms
    tech_matches = len(tech_intersection)

    # 组合相似度：Jaccard相似度 + 专业术语加权
    # 专业术语匹配可以显著提高相似度
    similarity = jaccard * 0.6 + (tech_matches / max(len(technical_terms), 1)) * 0.4

    # 如果有专业术语匹配，额外提升相似度
    if tech_matches > 0:
        similarity = min(similarity + tech_matches * 0.15, 1.0)

    return min(similarity, 1.0)


def _match_images_by_caption(
    doc: Document,
    section_start: int,
    section_end: int,
    sop_images: List[Dict[str, Any]],
    similarity_threshold: float = 0.3
) -> List[Tuple[int, Dict[str, Any], float]]:
    """根据标题语义匹配模板图片和SOP图片。

    Returns:
        匹配结果列表：[(模板段落索引, SOP图片字典, 相似度), ...]
    """
    # 找到模板中包含图片的段落
    template_image_infos = []
    for idx in range(section_start, min(section_end, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        if not _paragraph_has_image(para):
            continue
        caption = _extract_image_caption(doc, idx, section_start, section_end)
        template_image_infos.append((idx, caption))

    if not template_image_infos or not sop_images:
        return []

    # 计算SOP图片的标题（从caption字段或周围的文本推断）
    sop_image_infos = []
    for sop_img in sop_images:
        caption = str(sop_img.get('caption', '') or '').strip()
        # 如果没有caption，尝试从filename推断
        if not caption:
            filename = str(sop_img.get('filename', '') or '').strip()
            caption = filename
        sop_image_infos.append((sop_img, caption))

    # 匹配模板图片和SOP图片
    matches = []
    used_template_indices = set()
    used_sop_indices = set()

    # 为每个模板图片找到最佳匹配的SOP图片
    for tmpl_idx, tmpl_caption in template_image_infos:
        if tmpl_idx in used_template_indices:
            continue

        best_match = None
        best_similarity = 0.0
        best_sop_idx = -1

        for sop_idx, (sop_img, sop_caption) in enumerate(sop_image_infos):
            if sop_idx in used_sop_indices:
                continue

            similarity = _calculate_caption_similarity(tmpl_caption, sop_caption)

            if similarity > best_similarity and similarity >= similarity_threshold:
                best_similarity = similarity
                best_match = sop_img
                best_sop_idx = sop_idx

        if best_match is not None:
            matches.append((tmpl_idx, best_match, best_similarity))
            used_template_indices.add(tmpl_idx)
            used_sop_indices.add(best_sop_idx)

    # 图题差异大时：若模板与 SOP 图片张数一致，按文档顺序一一替换（仍须 rel_id 在 blob 映射中存在）
    if not matches and len(template_image_infos) == len(sop_image_infos) and template_image_infos:
        matches = [
            (tmpl_idx, sop_image_infos[j][0], 0.5)
            for j, (tmpl_idx, _) in enumerate(template_image_infos)
        ]

    # 按相似度降序排序
    matches.sort(key=lambda x: x[2], reverse=True)

    return matches


def _load_sop_image_blob_map(sop_path: Optional[str]) -> Dict[str, bytes]:
    """读取 SOP 文档图片二进制：rel_id -> blob。"""
    if not sop_path or not os.path.exists(sop_path):
        return {}
    try:
        sop_doc = Document(sop_path)
    except Exception:
        return {}
    blobs: Dict[str, bytes] = {}
    rels = getattr(sop_doc.part, 'rels', {})
    for rel_id, rel in rels.items():
        reltype = str(getattr(rel, 'reltype', '') or '')
        if '/image' not in reltype:
            continue
        part = getattr(rel, 'target_part', None)
        if not part:
            continue
        blob = getattr(part, 'blob', b'') or b''
        if blob:
            blobs[rel_id] = blob
    return blobs


def _find_subsection_end_anchor(
    doc: Document,
    start_idx: int,
    end_idx: int,
    target_keywords: List[str],
    next_keywords: List[str],
    strict_heading: bool = False
):
    """查找目标小节结束处锚点段落（用于在其前插图）。"""
    def _is_heading_like(i: int) -> bool:
        if i < 0 or i >= len(doc.paragraphs):
            return False
        para = doc.paragraphs[i]
        text = para.text.strip()
        style = para.style.name if para.style else ''
        if style.startswith('3.2.S') or 'Heading' in style or 'RA-5' in style or 'RA-6' in style:
            return True
        return bool(re.match(r'^\d+(\.\d+)*\s*', text))

    target_start = -1
    for i in range(start_idx + 1, end_idx):
        t = doc.paragraphs[i].text.strip()
        if strict_heading and not _is_heading_like(i):
            continue
        if any(k in t for k in target_keywords):
            target_start = i
            break
    if target_start < 0:
        return doc.paragraphs[end_idx] if end_idx < len(doc.paragraphs) else None

    for j in range(target_start + 1, end_idx):
        t = doc.paragraphs[j].text.strip()
        if strict_heading and not _is_heading_like(j):
            continue
        if any(k in t for k in next_keywords):
            return doc.paragraphs[j]
    return doc.paragraphs[end_idx] if end_idx < len(doc.paragraphs) else None


def _insert_section_images(
    doc: Document,
    anchor_para,
    section_images: List[Dict[str, Any]],
    image_blob_map: Dict[str, bytes],
    body_style=None,
    max_images: int = 6
) -> int:
    """在锚点前插入图片及图题。"""
    if not anchor_para or not section_images or not image_blob_map:
        return 0
    inserted = 0
    seen_sha = set()
    for item in section_images:
        if inserted >= max_images:
            break
        rel_id = str(item.get('rel_id', '') or '')
        sha1 = str(item.get('sha1', '') or '')
        if sha1 and sha1 in seen_sha:
            continue
        blob = image_blob_map.get(rel_id)
        if not blob:
            continue

        p_img = anchor_para.insert_paragraph_before('')
        r = p_img.add_run()
        try:
            r.add_picture(io.BytesIO(blob), width=Inches(5.8))
        except Exception:
            continue
        if body_style is not None:
            try:
                p_img.style = body_style
            except Exception:
                pass

        caption = str(item.get('caption', '') or '').strip()
        if caption:
            p_cap = anchor_para.insert_paragraph_before(caption)
            if body_style is not None:
                try:
                    p_cap.style = body_style
                except Exception:
                    pass
        if sha1:
            seen_sha.add(sha1)
        inserted += 1
    return inserted


def _replace_section_images(
    doc: Document,
    start_idx: int,
    end_idx: int,
    section_images: List[Dict[str, Any]],
    image_blob_map: Dict[str, bytes],
    body_style=None,
    similarity_threshold: float = 0.3,
    delete_unmatched: bool = False
) -> int:
    """根据标题语义相近的图片替换模板中的图片，不新增图片。

    Args:
        doc: 文档对象
        start_idx: 章节起始段落索引
        end_idx: 章节结束段落索引
        section_images: SOP中的图片列表
        image_blob_map: 图片数据映射
        body_style: 正文样式
        similarity_threshold: 语义相似度阈值（0-1），默认0.3
        delete_unmatched: 是否删除未匹配的图片，默认False（保留不匹配的图片）

    Returns:
        替换的图片数量
    """
    if not section_images or not image_blob_map:
        return 0

    # 使用语义匹配找到对应的图片
    matches = _match_images_by_caption(
        doc, start_idx, end_idx, section_images, similarity_threshold
    )

    if not matches:
        return 0

    # 替换匹配的图片
    replaced = 0
    seen_sha = set()

    for tmpl_idx, sop_img, similarity in matches:
        rel_id = str(sop_img.get('rel_id', '') or '')
        sha1 = str(sop_img.get('sha1', '') or '')

        if sha1 and sha1 in seen_sha:
            continue

        blob = image_blob_map.get(rel_id)
        if not blob:
            continue

        # 找到对应的图片段落
        para = doc.paragraphs[tmpl_idx]

        # 清除段落内容并添加新图片
        para.clear()
        r = para.add_run()
        try:
            r.add_picture(io.BytesIO(blob), width=Inches(5.8))
        except Exception:
            continue

        if body_style is not None:
            try:
                para.style = body_style
            except Exception:
                pass

        # 添加或替换图题
        sop_caption = str(sop_img.get('caption', '') or '').strip()
        if sop_caption:
            # 检查下一个段落是否是图题
            if tmpl_idx + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[tmpl_idx + 1]
                next_text = next_para.text.strip()
                # 如果下一个段落是图题（不包含图片），则替换它
                if next_text and not _paragraph_has_image(next_para):
                    if tmpl_idx + 1 <= end_idx:
                        next_para.clear()
                        next_para.add_run(sop_caption)
                        if body_style is not None:
                            try:
                                next_para.style = body_style
                            except Exception:
                                pass
                    else:
                        # 添加新图题段落
                        new_cap_para = doc.paragraphs[tmpl_idx + 1].insert_paragraph_before(sop_caption)
                        if body_style is not None:
                            try:
                                new_cap_para.style = body_style
                            except Exception:
                                pass
                else:
                    # 添加新图题段落
                    new_cap_para = doc.paragraphs[tmpl_idx + 1].insert_paragraph_before(sop_caption)
                    if body_style is not None:
                        try:
                            new_cap_para.style = body_style
                        except Exception:
                            pass

        if sha1:
            seen_sha.add(sha1)
        replaced += 1

    # 处理未匹配的图片
    if delete_unmatched:
        # 找到所有图片段落
        all_image_indices = []
        for idx in range(start_idx, min(end_idx, len(doc.paragraphs))):
            if _paragraph_has_image(doc.paragraphs[idx]):
                all_image_indices.append(idx)

        # 找到未匹配的图片并删除
        matched_indices = set(m[0] for m in matches)
        for idx in all_image_indices:
            if idx not in matched_indices:
                doc.paragraphs[idx].clear()

    return replaced


def build_integrated_content(
    sop_data: Dict[str, Any],
    style_mode: str = 'flat'
) -> List[Dict[str, str]]:
    """
    将 SOP 提取内容构建为分析方法格式。

    Args:
        sop_data: extract_procedure_chapter4() 的返回结果
        style_mode: 'flat' 使用 RA-正文 扁平格式；'hierarchical' 使用 RA-5/6级标题

    Returns:
        [{'style': str, 'text': str}, ...]
    """
    out = []
    RA_BODY = 'RA-正文'
    RA_H5 = 'RA-5级标题'
    RA_H6 = 'RA-6级标题'

    # 实验设备/试验设备 作为小节标题可过滤，设备具体名称从表格/段落提取
    skip_titles = [
        '实验原理', '实验材料', '实验设备', '试验设备', '样品处理', '操作步骤',
        '数据处理', '计算', '可接受标准', '系统适用性', 'Experimental', 'Principle',
        'Material', 'Sample Preparation', 'Operation', 'Data Processing', 'Calculate',
        'Acceptable Standard', 'System Suitability',
    ]

    # 1. 原理
    if style_mode == 'hierarchical':
        out.append({'style': RA_H5, 'text': '原理'})
    else:
        out.append({'style': RA_BODY, 'text': '原理：'})
    principle = _filter_content(sop_data.get('principle', []), skip_titles)
    if principle:
        if style_mode == 'flat':
            out[-1]['text'] = '原理：' + principle[0]
            for p in principle[1:5]:
                out.append({'style': RA_BODY, 'text': p})
        else:
            for p in principle[:6]:
                out.append({'style': RA_BODY, 'text': p})
    else:
        out[-1]['text'] = out[-1]['text'] + '按标准检验方法进行测定。'

    # 2. 材料和设备
    if style_mode == 'hierarchical':
        out.append({'style': RA_H5, 'text': '设备、材料、试剂'})
    else:
        out.append({'style': RA_BODY, 'text': '材料和设备'})
    mat = _filter_content(sop_data.get('materials_and_equipment', []), skip_titles)
    chrom = sop_data.get('chromatography_conditions', {})
    if chrom.get('色谱柱') and not any('色谱柱' in m for m in mat):
        mat.append(f"色谱柱：{chrom['色谱柱']}")
    for m in (mat if mat else ['设备、材料、试剂根据检验方法配置。']):
        out.append({'style': RA_BODY, 'text': m})

    # 3. 操作步骤
    if style_mode == 'hierarchical':
        out.append({'style': RA_H5, 'text': '操作步骤'})
        out.append({'style': RA_H6, 'text': '样品处理'})
    else:
        out.append({'style': RA_BODY, 'text': '操作步骤'})
    sample = _filter_content(sop_data.get('sample_prep', []), skip_titles)
    for s in sample[:15]:
        out.append({'style': RA_BODY, 'text': s})

    # 色谱条件 + 操作步骤
    if chrom:
        cond_parts = []
        for k, v in chrom.items():
            if v:
                cond_parts.append(f"{k}：{v}")
        if cond_parts:
            out.append({'style': RA_BODY, 'text': '色谱条件：' + '；'.join(cond_parts[:5]) + '。'})
    proc = _filter_content(sop_data.get('procedure', []), skip_titles)
    for p in proc[:8]:
        if any(kw in p for kw in ['平衡', '序列', '进样', '数据处理', '积分', 'Empower']):
            out.append({'style': RA_BODY, 'text': p})

    if style_mode == 'hierarchical':
        out.append({'style': RA_H6, 'text': '测定法'})
        if not proc and not chrom:
            out.append({'style': RA_BODY, 'text': '按检验方法操作步骤进行测定。'})

    # 4. 试验成立标准
    if style_mode == 'hierarchical':
        out.append({'style': RA_H5, 'text': '试验成立标准'})
    else:
        out.append({'style': RA_BODY, 'text': '试验成立标准：'})
    suit = _filter_content(sop_data.get('suitability_criteria', []), skip_titles)
    if suit:
        if style_mode == 'flat':
            out[-1]['text'] = out[-1]['text'] + suit[0]
            for line in suit[1:8]:
                out.append({'style': RA_BODY, 'text': line})
        else:
            for line in suit[:8]:
                out.append({'style': RA_BODY, 'text': line})
    else:
        out[-1]['text'] = out[-1]['text'] + '见表。'

    # 5. 结果计算
    if style_mode == 'hierarchical':
        out.append({'style': RA_H5, 'text': '结果计算'})
    else:
        out.append({'style': RA_BODY, 'text': '结果计算：'})
    calc = _filter_content(sop_data.get('result_calculation', []), skip_titles)
    if calc:
        if style_mode == 'flat':
            out[-1]['text'] = out[-1]['text'] + calc[0]
            for c in calc[1:10]:
                out.append({'style': RA_BODY, 'text': c})
        else:
            for c in calc[:10]:
                out.append({'style': RA_BODY, 'text': c})
    else:
        out[-1]['text'] = out[-1]['text'] + '按检验方法规定计算。'

    # 6. 合格标准
    if style_mode == 'hierarchical':
        out.append({'style': RA_H5, 'text': '合格标准'})
    else:
        out.append({'style': RA_BODY, 'text': '合格标准：'})
    accept = _filter_content(sop_data.get('acceptance_criteria', []), skip_titles)
    if accept:
        acc_text = accept[0]
        for a in accept[1:]:
            if any(kw in a for kw in ['≥', '≤', '%', 'mg/mL', '不低于', '不高于', '一致', '参照']):
                acc_text = a
                break
        out[-1]['text'] = out[-1]['text'] + acc_text
    else:
        out[-1]['text'] = out[-1]['text'] + '参照相应质量标准执行。'

    return out


# ─────────────────────────────────────────────────────────────
# Phase 1 / Phase 3: 提取 JSON / 从 JSON 写入（供 Claude CLI 精简流程使用）
# ─────────────────────────────────────────────────────────────

def _extract_section_text_from_doc(
    doc: Document,
    section_name: str,
    stop_keywords: List[str]
) -> Dict[str, str]:
    """
    从 docx 中提取指定方法章节内的段落文本，按六章节关键词分组。

    返回 {"principle": "...", "materials_and_equipment": "...", ...}
    """
    start_idx, end_idx = find_method_section_range(doc, section_name, stop_keywords)
    if start_idx < 0:
        return {}

    # 六章节关键词 → key 映射
    section_keywords = {
        'principle': ['原理'],
        'materials_and_equipment': ['主要材料和设备', '设备、材料、试剂', '材料和设备', '材料与设备'],
        'procedure': ['操作步骤', '测定法'],
        'suitability_criteria': ['可接受标准', '试验成立标准', '系统适用性'],
        'result_calculation': ['计算公式', '结果计算'],
        'acceptance_criteria': ['合格标准'],
    }

    result = {}
    current_key = 'principle'  # 默认从原理开始
    current_lines = []

    for i in range(start_idx + 1, end_idx):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue
        # 检测是否进入新小节
        matched_key = None
        for key, kws in section_keywords.items():
            for kw in kws:
                if text.startswith(kw) or text == kw:
                    matched_key = key
                    break
            if matched_key:
                break
        if matched_key and matched_key != current_key:
            if current_lines:
                result[current_key] = '\n'.join(current_lines)
            current_key = matched_key
            current_lines = []
            # 如果该行不只是标题，也保留内容部分
            for kw in section_keywords[matched_key]:
                if text.startswith(kw + '：'):
                    text = text[len(kw) + 1:].strip()
                    break
                elif text == kw:
                    text = ''
                    break
            if text:
                current_lines.append(text)
        else:
            current_lines.append(text)

    if current_lines:
        result[current_key] = '\n'.join(current_lines)

    return result


def extract_for_refinement(
    template_path: str,
    sop_path: str,
    section_name: str,
    stop_keywords: List[str],
    ref_template_path: Optional[str] = None
) -> dict:
    """
    Phase 1: 提取 SOP 原始内容 + 模板已有内容 + 参考模板样本。

    Args:
        template_path: 目标分析方法模板路径
        sop_path: SOP 文档路径
        section_name: 模板中方法章节名称
        stop_keywords: 下一章节关键词
        ref_template_path: 参考模板路径（可选，用于提取风格样本）

    Returns:
        结构化 dict，可直接 json.dumps 输出
    """
    if not os.path.exists(sop_path):
        raise FileNotFoundError(f"SOP 不存在: {sop_path}")

    # 1. 提取 SOP 第四章原始内容
    extractor = SOPExtractor(sop_path)
    sop_data = extractor.extract_procedure_chapter4(section_name)
    # 转为纯字符串列表（去除非 JSON 序列化的内容）
    sop_raw = {}
    for key in ['principle', 'materials_and_equipment', 'sample_prep',
                'procedure', 'suitability_criteria', 'result_calculation',
                'acceptance_criteria']:
        val = sop_data.get(key, [])
        if isinstance(val, list):
            sop_raw[key] = [str(v) for v in val if str(v).strip()]
        elif isinstance(val, dict):
            sop_raw[key] = [f"{k}：{v}" for k, v in val.items() if v]
        else:
            sop_raw[key] = [str(val)] if val else []
    # 色谱条件合并到 procedure
    chrom = sop_data.get('chromatography_conditions', {})
    if chrom:
        chrom_text = '色谱条件：' + '；'.join(
            f"{k}：{v}" for k, v in chrom.items() if v
        )
        sop_raw.setdefault('procedure', []).insert(0, chrom_text)

    method_tables = extractor.extract_method_related_tables()

    # 蛋白质含量稳态特判：将“样品名称/可接受标准”表并入试验成立标准来源
    if ('蛋白' in section_name and '含量' in section_name):
        patched_tables: List[Dict[str, Any]] = []
        extra_suitability: List[str] = []
        for tb in method_tables:
            if not isinstance(tb, dict):
                continue
            rows = tb.get('rows', []) or []
            header = " ".join(str(c).strip() for c in (rows[0] if rows else []))
            if ('可接受标准' in header) and any(k in header for k in ['样品', '名称']):
                extra_suitability.extend(_explode_acceptance_table_rows(rows))
                tb2 = dict(tb)
                tb2['category'] = 'suitability_criteria'
                patched_tables.append(tb2)
            else:
                patched_tables.append(tb)
        if extra_suitability:
            cur = sop_raw.get('suitability_criteria', []) or []
            seen = {_normalize_sentence_for_dedup(x) for x in cur}
            for line in extra_suitability:
                n = _normalize_sentence_for_dedup(line)
                if n and n not in seen:
                    seen.add(n)
                    cur.append(line)
            sop_raw['suitability_criteria'] = cur
        method_tables = patched_tables

    result = {
        'name': section_name,
        'stop': stop_keywords,
        'sop_path': os.path.abspath(sop_path),
        'sop_raw': sop_raw,
        'sop_tables': method_tables,
        'sop_section_images': sop_data.get('section_images', {}),
    }

    # 2. 提取模板已有内容
    if os.path.exists(template_path):
        doc = Document(template_path)
        result['template_existing'] = _extract_section_text_from_doc(
            doc, section_name, stop_keywords
        )

    # 3. 提取参考模板风格样本
    if ref_template_path and os.path.exists(ref_template_path):
        ref_doc = Document(ref_template_path)
        result['reference_style'] = _extract_section_text_from_doc(
            ref_doc, section_name, stop_keywords
        )

    return result


def build_refined_content(
    refined_data: Dict[str, List[str]],
    style_mode: str = 'flat',
    independent_title_sections: set = None
) -> List[Dict[str, str]]:
    """
    Phase 3: 将 LLM 精简后的内容构建为可写入模板的格式。

    与 build_integrated_content 逻辑相同，但输入是精简后的文本而非原始 SOP 数据。

    Args:
        refined_data: {"principle": [...], "materials_and_equipment", [...], ...}
        style_mode: 'flat' | 'hierarchical'
        independent_title_sections: 有独立标题的小节名称集合，如 {'材料和设备', '操作步骤'}

    Returns:
        [{'style': str, 'text': str, 'section_key': str (optional)}, ...]（不再输出 level；写回时 RA-正文会清除列表编号）
    """
    if independent_title_sections is None:
        independent_title_sections = set()

    out = []
    RA_BODY = 'RA-正文'
    RA_H5 = 'RA-5级标题'
    RA_H6 = 'RA-6级标题'

    # 章节配置：(key, flat标题, hierarchical标题, 默认文本)
    # sample_prep：hier_title=None — 模版已有 RA-5「操作步骤」及 RA-6「样品处理」，
    # 此处只输出 3.1 下正文（空白溶液/系统适用性溶液/供试品制备），勿再插入重复「操作步骤」标题行。
    sections = [
        ('principle', '原理：', '原理', None),
        ('materials_and_equipment', '材料和设备', '材料和设备', None),
        ('sample_prep', '操作步骤', None, None),
        ('procedure', None, None, None),  # 3.2 操作步骤正文，紧随 3.1 之后写入同一大节
        ('suitability_criteria', '试验成立标准：', '试验成立标准', None),
        ('result_calculation', '结果计算：', '结果计算', None),
        ('acceptance_criteria', '合格标准：', '合格标准', None),
    ]

    for key, flat_title, hier_title, default_text in sections:
        lines = refined_data.get(key, [])
        if isinstance(lines, str):
            lines = [lines] if lines.strip() else []

        # procedure 合并到 sample_prep 后面，不单独加标题
        if key == 'procedure':
            for line in lines:
                if line.strip():
                    out.append({'style': RA_BODY, 'text': line.strip(), 'section_key': key})
            continue

        # 处理章节内容（不写 level：避免 _set_paragraph_level 挂多级列表导致 1.1、2.1）
        if lines:
            # hierarchical模式
            if style_mode == 'hierarchical' and hier_title:
                # 检查是否是独立标题的小节
                if hier_title in independent_title_sections:
                    # 独立标题：先添加标题行，再添加所有内容行
                    out.append({'style': RA_BODY, 'text': hier_title, 'section_key': key})
                    for line in lines:
                        if line.strip():
                            out.append({'style': RA_BODY, 'text': line.strip(), 'section_key': key})
                else:
                    # 合并标题：将标题和第一行内容合并
                    first_line = hier_title + '：' + lines[0]
                    out.append({'style': RA_BODY, 'text': first_line, 'section_key': key})
                    for line in lines[1:]:
                        if line.strip():
                            out.append({'style': RA_BODY, 'text': line.strip(), 'section_key': key})
            # flat模式：标题与首行合并
            elif style_mode == 'flat' and flat_title and flat_title.endswith('：'):
                out.append({'style': RA_BODY, 'text': flat_title + lines[0], 'section_key': key})
                for line in lines[1:]:
                    if line.strip():
                        out.append({'style': RA_BODY, 'text': line.strip(), 'section_key': key})
            else:
                # 直接添加所有行
                for line in lines:
                    if line.strip():
                        out.append({'style': RA_BODY, 'text': line.strip(), 'section_key': key})
        elif default_text:
            # 使用默认文本
            if style_mode == 'hierarchical' and hier_title:
                out.append({'style': RA_BODY, 'text': hier_title + '：' + default_text, 'section_key': key})
            elif flat_title and flat_title.endswith('：'):
                out.append({'style': RA_BODY, 'text': flat_title + default_text, 'section_key': key})
            else:
                out.append({'style': RA_BODY, 'text': default_text, 'section_key': key})

    return out


def _set_paragraph_level(para, level: int) -> None:
    """设置段落的编号层级（Lv0/Lv1）。

    Args:
        para: 段落对象
        level: 层级（0=顶层标题，1=子项内容）
    """
    try:
        # 获取或创建段落属性的pPr元素
        pPr = para._element.pPr
        if pPr is None:
            from docx.oxml import OxmlElement
            pPr = OxmlElement('w:pPr')
            para._element.insert(0, pPr)

        # 设置编号层级
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            # 兼容修复：目标段落无编号时，尝试继承相邻段落的 numId
            from docx.oxml import OxmlElement
            numPr = OxmlElement('w:numPr')
            pPr.append(numPr)

            def _extract_num_id(p_elem):
                pPr2 = p_elem.find(qn('w:pPr'))
                if pPr2 is None:
                    return None
                numPr2 = pPr2.find(qn('w:numPr'))
                if numPr2 is None:
                    return None
                numId2 = numPr2.find(qn('w:numId'))
                if numId2 is None:
                    return None
                return numId2.get(qn('w:val'))

            inherited_num_id = None

            # 向前查找
            prev = para._element.getprevious()
            while prev is not None:
                if prev.tag.rsplit('}', 1)[-1] == 'p':
                    inherited_num_id = _extract_num_id(prev)
                    if inherited_num_id:
                        break
                prev = prev.getprevious()

            # 向后兜底查找
            if not inherited_num_id:
                nxt = para._element.getnext()
                while nxt is not None:
                    if nxt.tag.rsplit('}', 1)[-1] == 'p':
                        inherited_num_id = _extract_num_id(nxt)
                        if inherited_num_id:
                            break
                    nxt = nxt.getnext()

            if inherited_num_id:
                numId = OxmlElement('w:numId')
                numId.set(qn('w:val'), str(inherited_num_id))
                numPr.append(numId)

        if numPr is not None:
            # 设置ilvl
            ilvl = numPr.find(qn('w:ilvl'))
            if ilvl is None:
                from docx.oxml import OxmlElement
                ilvl = OxmlElement('w:ilvl')
                numPr.append(ilvl)
            ilvl.set(qn('w:val'), str(level))
    except Exception as e:
        # 如果设置失败，不影响主流程
        pass


def _get_paragraph_level(para):
    """读取段落编号层级，未编号返回 None。"""
    try:
        pPr = para._element.pPr
        if pPr is None:
            return None
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            return None
        ilvl = numPr.find(qn('w:ilvl'))
        if ilvl is None:
            return None
        return ilvl.get(qn('w:val'))
    except Exception:
        return None


def _clear_list_numbering_paragraph(para) -> None:
    """
    去掉段落的 Word 列表编号（w:numPr）。
    RA-正文写入后必须调用，避免出现与节标题「1 原理 / 2 设备」叠加的 1.1、2.1 等子编号。
    """
    try:
        pPr = para._element.pPr
        if pPr is None:
            return
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
    except Exception:
        pass


def _clear_direct_paragraph_indent(para) -> None:
    """
    仅清理“当前被写入且带编号段落”的直接缩进属性，
    避免 firstLine/hanging 残留导致局部缩进异常。
    """
    # 严格模板继承模式：不主动修改缩进
    return


def _prioritize_items_for_template_slots(
    items: List[Dict[str, Any]],
    max_slots: int
) -> List[Dict[str, Any]]:
    """
    模板槽位不足时的保底裁剪：
    - 优先保留核心标题（尤其结果计算/合格标准）
    - 再按优先级保留关键内容，最后按原顺序输出
    """
    if max_slots <= 0 or len(items) <= max_slots:
        return items

    title_keys = [
        '原理', '材料和设备', '设备、材料、试剂', '材料与设备',
        '操作步骤', '试验成立标准', '结果计算', '合格标准',
    ]

    def score(idx_item):
        idx, item = idx_item
        txt = str(item.get('text', '') or '')
        lvl = int(item.get('level', 0) or 0)
        s = 0
        # 顶层标题最高
        if lvl == 0 and any(k in txt for k in title_keys):
            s += 1000
        # 合格标准与结果计算优先
        if '合格标准' in txt:
            s += 900
        if '结果计算' in txt:
            s += 800
        # 公式/限度值优先
        if any(k in txt for k in ['=', '≥', '≤', '%', 'mg/mL', '~']):
            s += 300
        # 前序内容略优先，避免只保留尾部
        s += max(0, 200 - idx)
        return s

    indexed = list(enumerate(items))
    must_keep = set()
    for i, it in indexed:
        txt = str(it.get('text', '') or '')
        lvl = int(it.get('level', 0) or 0)
        if lvl == 0 and any(k in txt for k in title_keys):
            must_keep.add(i)

    # 先放必保项
    selected = set(sorted(must_keep)[:max_slots])
    if len(selected) < max_slots:
        rest = [x for x in indexed if x[0] not in selected]
        rest.sort(key=score, reverse=True)
        for i, _ in rest:
            if len(selected) >= max_slots:
                break
            selected.add(i)

    kept = [items[i] for i in sorted(selected)]
    return kept


def _normalize_sentence_for_dedup(text: str) -> str:
    """用于句子去重的归一化（保留语义，忽略标点空白差异）。"""
    t = str(text or '').strip().lower()
    t = re.sub(r'[\s\u3000]+', '', t)
    t = re.sub(r'[，。；：、,.!！?？（）()\[\]【】\-—_~`\'"“”‘’<>《》/\\]', '', t)
    return t


def _split_semantic_units(text: str) -> List[str]:
    """
    语义分句：
    - 优先按常见小标题键（如“供试品溶液:”）切分
    - 再按句号/分号做二次切分
    """
    s = str(text or '').strip()
    if not s:
        return []

    label_pat = (
        r'(供试品制备[:：]|供试品溶液[:：]|空白溶液[:：]|系统适用性溶液[:：]|'
        r'FB溶液[:：]|操作步骤[:：]|色谱条件[:：]|结果计算[:：]|试验成立标准[:：]|合格标准[:：])'
    )
    parts = re.split(label_pat, s)

    units: List[str] = []
    if len(parts) > 1:
        # re.split 返回: [prefix, label1, content1, label2, content2, ...]
        prefix = parts[0].strip()
        if prefix:
            units.extend([x.strip() for x in re.split(r'[；;。]\s*', prefix) if x.strip()])
        i = 1
        while i + 1 < len(parts):
            label = parts[i].strip()
            content = parts[i + 1].strip()
            merged = (label + content).strip()
            if merged:
                units.append(merged)
            i += 2
    else:
        units.extend([x.strip() for x in re.split(r'[；;。]\s*', s) if x.strip()])

    # 去重并保序
    dedup: List[str] = []
    seen = set()
    for u in units:
        n = _normalize_sentence_for_dedup(u)
        if not n or n in seen:
            continue
        seen.add(n)
        dedup.append(u)
    return dedup


def _merge_text_semantic(base: str, extras: List[str]) -> str:
    """将基文本与超量文本按语义去重并换行拼接。"""
    merged_units: List[str] = []
    for piece in [base] + list(extras or []):
        merged_units.extend(_split_semantic_units(piece))
    # 二次去重保序
    out: List[str] = []
    seen = set()
    for u in merged_units:
        n = _normalize_sentence_for_dedup(u)
        if not n or n in seen:
            continue
        seen.add(n)
        out.append(u)
    return '\n'.join(out)


def _explode_acceptance_table_rows(rows: List[List[str]]) -> List[str]:
    """将两列表（样品名称/可接受标准）拆为试验成立标准行列表。"""
    out: List[str] = []
    if not rows or len(rows) < 2:
        return out
    for r in rows[1:]:
        if len(r) < 2:
            continue
        name = str(r[0] or '').strip()
        criteria = str(r[1] or '').strip()
        if not name or not criteria:
            continue
        parts = re.split(r'[\n\r]+|；|;', criteria)
        for p in parts:
            s = str(p).strip().rstrip('。')
            if s:
                out.append(f"{name}：{s}")
    dedup: List[str] = []
    seen = set()
    for x in out:
        n = _normalize_sentence_for_dedup(x)
        if not n or n in seen:
            continue
        seen.add(n)
        dedup.append(x)
    return dedup


def _anchor_for_section_key(section_key: str) -> str:
    """字段 key 到章节锚点标题的映射。"""
    mapping = {
        'principle': '原理',
        'materials_and_equipment': '材料和设备',
        'sample_prep': '操作步骤',
        'procedure': '操作步骤',
        'suitability_criteria': '试验成立标准',
        'result_calculation': '结果计算',
        'acceptance_criteria': '合格标准',
    }
    return mapping.get(section_key, '')


# 规范锚点名 → 段落中可能出现的标题表述（长语在前，避免子串误匹配）
# 与 input/32s42-分析方法-模板文件.docx 中 RA-5 标题对齐（含「主要材料和设备」「计算公式」「可接受标准」）
_CANONICAL_ANCHOR_PHRASES: Dict[str, List[str]] = {
    '原理': ['原理'],
    '材料和设备': [
        '主要材料和设备', '设备、材料、试剂', '材料和设备', '材料与设备', '设备、材料',
    ],
    '操作步骤': ['操作步骤'],
    '试验成立标准': ['试验成立标准', '系统适用性', '可接受标准'],
    '结果计算': ['计算公式', '结果计算'],
    '合格标准': ['合格标准'],
}


def _text_matches_anchor_phrases(txt: str, phrases: List[str], allow_contains: bool = True) -> bool:
    t = txt.strip()
    if not t:
        return False
    for ph in phrases:
        if t == ph or t.startswith(ph + '：') or t.startswith(ph + ':'):
            return True
    if allow_contains:
        for ph in phrases:
            if ph in t:
                return True
    return False


def _anchor_heading_texts_to_skip(anchor: str) -> set:
    """写入时与模板标题重复的纯标题行，避免再写一遍。"""
    s = {anchor}
    for ph in _CANONICAL_ANCHOR_PHRASES.get(anchor, []):
        s.add(ph)
        s.add(ph + '：')
    return s


def _ensure_principle_reflects_sop(
    filtered_refined: Dict[str, Any],
    required_source: Optional[Dict[str, Any]],
) -> None:
    """
    若精简结果未体现 SOP 第四章「实验原理」核心表述，则用 required_source 中的原理覆盖。
    （避免 LLM/规则仍沿用模版占位句而忽略 SOP。）
    """
    if not required_source:
        return
    sop_p = required_source.get('principle')
    if sop_p is None:
        return
    if isinstance(sop_p, str):
        lines = [sop_p.strip()] if sop_p.strip() else []
    else:
        lines = [str(x).strip() for x in sop_p if str(x).strip()]
    if not lines:
        return
    cur = filtered_refined.get('principle') or []
    if isinstance(cur, str):
        cur = [cur] if str(cur).strip() else []
    cur_join = _normalize_for_match(''.join(cur))
    sop_join = _normalize_for_match(''.join(lines))
    if len(sop_join) < 12:
        return
    if sop_join not in cur_join:
        filtered_refined['principle'] = lines[:5]


def _ensure_anchor_sections_not_empty(
    doc: Document,
    start_idx: int,
    end_idx: int,
    filtered_refined: Dict[str, Any],
    styles: Dict[str, Any],
) -> int:
    """
    确保关键小节（试验成立标准/结果计算/合格标准）在“新增骨架章节”场景不为空：
    - 若小节标题后到下一小节标题前没有任何非空段落，且 filtered_refined 中有对应内容，则补插 RA-正文 段。
    - 返回插入段落数（用于更新 end_idx）。
    """
    anchor_order = ['原理', '材料和设备', '操作步骤', '试验成立标准', '结果计算', '合格标准']
    key_for_anchor = {
        '试验成立标准': 'suitability_criteria',
        '结果计算': 'result_calculation',
        '合格标准': 'acceptance_criteria',
    }
    anchor_pos = _detect_anchor_positions(doc, start_idx, end_idx)
    existing_sorted = sorted(
        [(a, anchor_pos[a]) for a in anchor_order if a in anchor_pos],
        key=lambda x: x[1],
    )
    ranges: Dict[str, Tuple[int, int]] = {}
    for i, (a, p_idx) in enumerate(existing_sorted):
        n_idx = existing_sorted[i + 1][1] if i + 1 < len(existing_sorted) else end_idx
        ranges[a] = (p_idx, n_idx)

    inserted = 0
    # 重要：按“从后往前”插入，避免前面小节插入段落导致后面小节的段落索引整体右移
    targets: List[Tuple[int, str, str]] = []
    for anchor, k in key_for_anchor.items():
        if anchor in ranges:
            a0, _a1 = ranges[anchor]
            targets.append((a0, anchor, k))
    targets.sort(key=lambda x: x[0], reverse=True)

    for _a0, anchor, k in targets:
        a0, a1 = ranges[anchor]
        # 判断分区内是否已有正文
        has_body = False
        for j in range(a0 + 1, min(a1, len(doc.paragraphs))):
            txt = doc.paragraphs[j].text.strip()
            st = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ''
            if txt and ('RA-5' not in st and 'RA-6' not in st and 'Heading' not in st):
                has_body = True
                break
        if has_body:
            continue
        lines = filtered_refined.get(k) or []
        if isinstance(lines, str):
            lines = [lines] if lines.strip() else []
        lines = [str(x).strip() for x in lines if str(x).strip()]
        if not lines:
            continue
        # 在下一小节标题前插入（或在章节末尾）
        insert_before_idx = a1 if a1 < len(doc.paragraphs) else len(doc.paragraphs) - 1
        cur = doc.paragraphs[insert_before_idx] if insert_before_idx < len(doc.paragraphs) else None
        body_st = styles.get('RA-正文')
        for text in reversed(lines):
            if cur is not None:
                new_p = cur.insert_paragraph_before(text)
            else:
                new_p = doc.add_paragraph(text)
            if body_st:
                try:
                    new_p.style = body_st
                except Exception:
                    pass
            _clear_list_numbering_paragraph(new_p)
            inserted += 1
        print(f"  [INFO] 小节「{anchor}」原无正文，已补插 {len(lines)} 段")

    return inserted


def _sanitize_ra_heading_paragraphs(
    doc: Document,
    start_idx: int,
    end_idx: int,
    styles: Dict[str, Any],
) -> int:
    """
    修复异常写回导致的“RA-5/RA-6 标题段落携带正文”的问题。

    期望：
    - RA-5级标题 段落仅包含标题文本（如「合格标准」），不应含换行或冒号后的正文
    - 正文应落在紧随其后的 RA-正文 段落（必要时插入新的 RA-正文）

    返回：插入的正文段落数量（用于更新 end_idx）。
    """
    RA_H5 = 'RA-5级标题'
    RA_H6 = 'RA-6级标题'
    RA_BODY = 'RA-正文'
    body_style = styles.get(RA_BODY)
    inserted = 0

    # 标题关键字集合（用于识别“标题：正文”中的标题部分）
    title_keywords: List[str] = []
    for a in ('原理', '材料和设备', '操作步骤', '试验成立标准', '结果计算', '合格标准'):
        title_keywords.append(a)
        title_keywords.extend(_CANONICAL_ANCHOR_PHRASES.get(a, []))
    title_keywords = [x for x in dict.fromkeys(title_keywords) if x]

    def _split_heading_and_body(txt: str) -> Tuple[str, str]:
        t = (txt or '').replace('\r', '').strip()
        if not t:
            return '', ''
        if '\n' in t:
            lines = [x.strip() for x in t.split('\n') if x.strip()]
            head = (lines[0] if lines else '').strip().rstrip('：:')
            body = '\n'.join(lines[1:]).strip()
            return head, body
        # “标题：正文”也视为异常（标题段落不应承载正文）
        for kw in title_keywords:
            if t.startswith(kw + '：') or t.startswith(kw + ':'):
                head = kw.strip().rstrip('：:')
                body = t[len(kw) + 1 :].strip()
                return head, body
        # 若只是标题末尾多了冒号，也规整掉
        return t.strip().rstrip('：:'), ''

    i = start_idx
    while i < min(end_idx, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        st = p.style.name if p.style else ''
        if st not in (RA_H5, RA_H6):
            i += 1
            continue
        head, body = _split_heading_and_body(p.text or '')
        if not head:
            i += 1
            continue
        if not body and (p.text or '').strip() == head:
            i += 1
            continue

        # 重置标题段落文本为纯标题
        p.text = head

        if not body:
            i += 1
            continue

        # 将正文写入下一段落（优先复用紧随其后的空 RA-正文）
        next_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None
        body_lines = [x.strip() for x in body.split('\n') if x.strip()]
        if not body_lines:
            i += 1
            continue

        if next_para is not None and (next_para.style and next_para.style.name == RA_BODY) and (not next_para.text.strip()):
            next_para.text = body_lines[0]
            _clear_list_numbering_paragraph(next_para)
            if body_style:
                try:
                    next_para.style = body_style
                except Exception:
                    pass
            inserted += 1
            for extra in body_lines[1:]:
                np = next_para.insert_paragraph_before(extra)
                if body_style:
                    try:
                        np.style = body_style
                    except Exception:
                        pass
                _clear_list_numbering_paragraph(np)
                inserted += 1
        else:
            anchor = next_para if next_para is not None else None
            for extra in reversed(body_lines):
                if anchor is not None:
                    np = anchor.insert_paragraph_before(extra)
                else:
                    np = doc.add_paragraph(extra)
                if body_style:
                    try:
                        np.style = body_style
                    except Exception:
                        pass
                _clear_list_numbering_paragraph(np)
                inserted += 1

        i += 1

    return inserted


def _detect_anchor_positions(doc: Document, start_idx: int, end_idx: int) -> Dict[str, int]:
    """
    在方法章节内定位 6 个锚点标题（优先 ilvl=0 标题行）。
    支持新模板「设备、材料、试剂」与旧表述「材料和设备」等价。
    """
    canonical_order = ['原理', '材料和设备', '操作步骤', '试验成立标准', '结果计算', '合格标准']
    pos: Dict[str, int] = {}

    def try_assign(i: int, txt: str, require_ilvl0: bool) -> None:
        para = doc.paragraphs[i]
        ilvl = _get_paragraph_level(para)
        if require_ilvl0 and ilvl != '0':
            return
        st = para.style.name if para.style else ''
        is_subsection_style = 'RA-5' in st or 'RA-6' in st or ('RA-' in st and '标题' in st)
        # 非标题样式的长段落里允许包含“系统适用性”等词，不能据此误判为锚点标题
        allow_contains = bool(ilvl == '0' or is_subsection_style or len(txt.strip()) <= 20)
        for canonical in canonical_order:
            if canonical in pos:
                continue
            phrases = _CANONICAL_ANCHOR_PHRASES.get(canonical, [canonical])
            if _text_matches_anchor_phrases(txt, phrases, allow_contains=allow_contains):
                pos[canonical] = i
                break

    for i in range(start_idx + 1, min(end_idx, len(doc.paragraphs))):
        txt = doc.paragraphs[i].text.strip()
        if not txt:
            continue
        try_assign(i, txt, require_ilvl0=True)

    for i in range(start_idx + 1, min(end_idx, len(doc.paragraphs))):
        txt = doc.paragraphs[i].text.strip()
        if not txt:
            continue
        try_assign(i, txt, require_ilvl0=False)

    return pos


def write_refined_to_template(
    template_path: str,
    section_name: str,
    stop_keywords: List[str],
    refined_content: Dict[str, List[str]],
    output_path: Optional[str] = None,
    style_mode: str = 'flat',
    source_pool: Optional[List[str]] = None,
    replacement_tables: Optional[List[Dict[str, Any]]] = None,
    required_source: Optional[Dict[str, Any]] = None,
    replacement_section_images: Optional[Dict[str, List[Dict[str, Any]]]] = None,
    sop_path: Optional[str] = None
) -> str:
    """
    Phase 3: 将精简后的内容写入模板对应章节。

    Args:
        template_path: 分析方法模板路径
        section_name: 模板中方法章节名称
        stop_keywords: 下一章节关键词
        refined_content: {"principle": [...], "materials_and_equipment": [...], ...}
        output_path: 输出路径
        style_mode: 'flat' | 'hierarchical'

    Returns:
        输出文件路径
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板不存在: {template_path}")

    if not output_path:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        base = os.path.splitext(os.path.basename(template_path))[0]
        output_path = os.path.join(os.path.dirname(template_path), f"{base}_{ts}.docx")

    if os.path.normpath(os.path.abspath(template_path)) != os.path.normpath(os.path.abspath(output_path)):
        shutil.copy(template_path, output_path)
    doc = Document(output_path)

    # 查找模板章节
    start_idx, end_idx = find_method_section_range(doc, section_name, stop_keywords)
    if start_idx < 0:
        raise ValueError(f"未找到章节: {section_name}")

    # 来源过滤：仅写入可在 SOP/模板中追溯的句子
    filtered_refined = _filter_refined_by_source(refined_content, source_pool)
    # 强约束补齐：关键章节不允许缺失
    filtered_refined = _enforce_required_method_lines(filtered_refined, required_source)

    # 兜底：若过滤后某字段为空，则回退到模板已有字段文本（避免引入无来源默认句）
    existing = _extract_section_text_from_doc(doc, section_name, stop_keywords)
    for key in ['principle', 'materials_and_equipment', 'sample_prep',
                'procedure', 'suitability_criteria', 'result_calculation',
                'acceptance_criteria']:
        val = filtered_refined.get(key, [])
        if isinstance(val, str):
            val = [val] if val.strip() else []
        if _field_value_is_effectively_empty(val):
            if key == 'materials_and_equipment':
                source_lines = _pick_material_equipment_from_source(source_pool)
                if source_lines:
                    filtered_refined[key] = source_lines
                    continue
            existing_lines = _split_existing_to_list(existing.get(key, ''))
            if existing_lines:
                filtered_refined[key] = [_normalize_display_text(x) for x in existing_lines]

    # 材料和设备：确保不丢"主要设备"等关键行
    filtered_refined['materials_and_equipment'] = _ensure_material_equipment_lines(
        filtered_refined.get('materials_and_equipment', []),
        _split_existing_to_list(existing.get('materials_and_equipment', '')),
        source_pool,
        section_name=section_name,
    )

    _ensure_principle_reflects_sop(filtered_refined, required_source)

    # 检测模板中存在的小节标题
    template_sections = _detect_template_sections(doc, start_idx, end_idx)
    print(f"  [DEBUG] 模板中检测到的小节: {template_sections}")

    # 检测模板中哪些小节有独立标题（不与内容合并）
    independent_title_sections = _detect_independent_title_sections(doc, start_idx, end_idx)
    if independent_title_sections:
        print(f"  [DEBUG] 检测到独立标题的小节: {independent_title_sections}")

    # 根据模板存在的小节，过滤refined内容
    filtered_refined = _filter_refined_by_template_sections(filtered_refined, template_sections, style_mode)

    # 检查章节中是否有表格，如果有表格则清空文字内容（避免表格+文字重复）
    section_table_indices = _get_table_indices_in_section(doc, start_idx, end_idx)
    has_suitability_table = False
    has_acceptance_table = False
    for idx in section_table_indices:
        if idx < len(doc.tables):
            table = doc.tables[idx]
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if _is_suitability_target_table(rows, doc, idx):
                has_suitability_table = True
            if _is_acceptance_target_table(rows, doc, idx):
                has_acceptance_table = True

    pre_clear_suitability: List[str] = []
    if has_suitability_table:
        pre_clear_suitability = [
            str(x).strip() for x in (filtered_refined.get('suitability_criteria', []) or []) if str(x).strip()
        ]

    # 如果章节中已有表格，清空对应的文字内容（用空列表，便于后续「模板回退」识别为缺失）
    if has_suitability_table and 'suitability_criteria' in filtered_refined:
        print(f"  [INFO] 章节中已有suitability_criteria表格，清空文字内容")
        filtered_refined['suitability_criteria'] = []
    if has_acceptance_table and 'acceptance_criteria' in filtered_refined:
        if _has_usable_acceptance_replacement(replacement_tables):
            print(
                "  [INFO] 章节中已有合格标准表格且 SOP 侧有可替换表数据，清空正文以免与表重复"
            )
            filtered_refined['acceptance_criteria'] = []
        else:
            print(
                "  [INFO] 模版有合格标准表但 SOP 侧无匹配表数据，保留合格标准正文，避免第 6 节整节为空"
            )

    # 蛋白质含量：表内两列「样品/名称+可接受标准」回填试验成立要点（在表格清空之后执行）
    if ('蛋白' in str(section_name) and '含量' in str(section_name)) and has_suitability_table:
        table_lines: List[str] = []
        for tb in (replacement_tables or []):
            if not isinstance(tb, dict):
                continue
            cat = str(tb.get('category') or '').strip()
            rows = tb.get('rows', []) or []
            header = " ".join(str(c).strip() for c in (rows[0] if rows else []))
            if cat in ('suitability_criteria', 'acceptance_criteria') and ('可接受标准' in header) and any(k in header for k in ['样品', '名称']):
                table_lines.extend(_explode_acceptance_table_rows(rows))
        if table_lines:
            cur = filtered_refined.get('suitability_criteria', []) or []
            seen = {_normalize_sentence_for_dedup(x) for x in cur}
            for line in table_lines:
                n = _normalize_sentence_for_dedup(line)
                if n and n not in seen:
                    seen.add(n)
                    cur.append(line)
            filtered_refined['suitability_criteria'] = cur

    if has_suitability_table and _field_value_is_effectively_empty(filtered_refined.get('suitability_criteria')) and pre_clear_suitability:
        filtered_refined['suitability_criteria'] = list(pre_clear_suitability)

    # 合格标准仍空时：模板占位 → SOP 原文 acceptance 段拆句 → 来源池挑选
    if _field_value_is_effectively_empty(filtered_refined.get('acceptance_criteria')):
        acc_lines = _split_existing_to_list(existing.get('acceptance_criteria', ''))
        if acc_lines:
            filtered_refined['acceptance_criteria'] = [_normalize_display_text(x) for x in acc_lines]
        elif required_source:
            raw = required_source.get('acceptance_criteria', [])
            flat: List[str] = []
            if isinstance(raw, list):
                for x in raw:
                    xs = str(x).strip()
                    if not xs:
                        continue
                    flat.extend(
                        [s.strip() for s in re.split(r'[。；;\n]', xs) if s.strip() and len(s.strip()) >= 4]
                    )
            elif isinstance(raw, str) and raw.strip():
                flat = [
                    s.strip()
                    for s in re.split(r'[。；;\n]', raw.strip())
                    if s.strip() and len(s.strip()) >= 4
                ]
            flat = [_normalize_display_text(x) for x in flat][:8]
            if flat:
                filtered_refined['acceptance_criteria'] = flat
        if _field_value_is_effectively_empty(filtered_refined.get('acceptance_criteria')):
            picked = _pick_acceptance_lines_from_source(source_pool)
            if picked:
                filtered_refined['acceptance_criteria'] = picked[:6]

    # 构建格式化内容
    content_list = build_refined_content(filtered_refined, style_mode, independent_title_sections)

    # 缓存样式
    styles = {}
    for name in ['RA-5级标题', 'RA-6级标题', 'RA-正文']:
        try:
            styles[name] = doc.styles[name]
        except Exception:
            styles[name] = None

    # 写入内容（复用 integrate_sop_into_template 的写入逻辑）
    valid_items = [x for x in content_list if x.get('text', '').strip()]
    writable_indices, table_caption_paras, para_to_table = _collect_writable_indices(doc, start_idx, end_idx)

    # 构建表格标题段落的集合，用于在清空段落时保护
    table_caption_set = set(table_caption_paras)

    next_section_para = doc.paragraphs[end_idx] if end_idx < len(doc.paragraphs) else None

    # 先定位锚点，再清理“见表”段落，避免锚点丢失导致分区跳过
    anchor_order = ['原理', '材料和设备', '操作步骤', '试验成立标准', '结果计算', '合格标准']
    anchor_pos = _detect_anchor_positions(doc, start_idx, end_idx)

    # 预处理：识别并清空包含"见表"的段落，避免它们影响写入顺序
    see_table_paragraphs = []  # 包含"见表"的段落索引
    for i in range(start_idx + 1, end_idx):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text.strip()
            if '见表' in text:
                see_table_paragraphs.append(i)
                # 清空这些段落，保留"见表"引用
                doc.paragraphs[i].clear()

    # 写入时跳过表格标题段落，避免在表格标题和表格之间插入内容
    actually_written_indices = set()
    # 稳规则：按章节锚点分区写入，禁止跨区混写

    # 构建每个锚点的范围（按文档出现顺序，避免「计算公式」在「可接受标准」之前时段区间颠倒）
    ranges: Dict[str, Tuple[int, int]] = {}
    existing = sorted(
        [(a, anchor_pos[a]) for a in anchor_order if a in anchor_pos],
        key=lambda x: x[1],
    )
    for i, (a, p_idx) in enumerate(existing):
        n_idx = existing[i + 1][1] if i + 1 < len(existing) else end_idx
        ranges[a] = (p_idx, n_idx)

    grouped: Dict[str, List[Dict[str, Any]]] = {a: [] for a in anchor_order}
    for item in valid_items:
        sec_key = str(item.get('section_key', '') or '')
        anchor = _anchor_for_section_key(sec_key)
        if anchor in grouped:
            grouped[anchor].append(item)

    skip_clear_elements = set()
    suitability_insert_count = 0
    procedure_insert_count = 0

    for anchor in anchor_order:
        items = grouped.get(anchor, [])
        if not items:
            continue
        # 模板已存在锚点标题行，避免重复写入同名纯标题（含「设备、材料、试剂」等新模板标题）
        skip_heads = _anchor_heading_texts_to_skip(anchor)
        items = [it for it in items if str(it.get('text', '')).strip() not in skip_heads]
        if not items:
            continue
        # 分区内去重，避免同义/重复句反复写入
        dedup_items: List[Dict[str, Any]] = []
        seen_item_text = set()
        for it in items:
            t = str(it.get('text', '') or '').strip()
            n = _normalize_sentence_for_dedup(t)
            if not n or n in seen_item_text:
                continue
            seen_item_text.add(n)
            dedup_items.append(it)
        items = dedup_items
        if not items:
            continue
        if anchor not in ranges:
            print(f"  [WARN] 章节 {section_name} 未定位锚点: {anchor}，该分区内容跳过")
            # 修复：如果该分区有内容需要写入，尝试创建锚点或使用备用策略
            if items and anchor in anchor_pos:
                print(f"  [INFO] 尝试为缺失锚点 '{anchor}' 创建写入槽位")
                # 使用原始锚点位置（可能需要调整偏移）
                original_pos = anchor_pos.get(anchor)
                if original_pos is not None and original_pos < len(doc.paragraphs):
                    # 创建临时范围：从锚点到文档末尾或下一个主要章节
                    next_anchor_idx = end_idx
                    for other_anchor in anchor_order:
                        if other_anchor != anchor and other_anchor in anchor_pos:
                            other_pos = anchor_pos[other_anchor]
                            if other_pos > original_pos:
                                next_anchor_idx = min(next_anchor_idx, other_pos)
                    ranges[anchor] = (original_pos, next_anchor_idx)
                    r_start, r_end = ranges[anchor]
                    local_slots = [idx for idx in writable_indices if (r_start < idx < r_end and idx not in table_caption_set)]
                    # 如果仍然没有槽位，使用锚点行本身
                    if not local_slots and original_pos < len(doc.paragraphs):
                        local_slots = [original_pos]
            else:
                continue
        r_start, r_end = ranges[anchor]

        # 当前分区内可写槽位（不含表格标题）
        local_slots = [idx for idx in writable_indices if (r_start < idx < r_end and idx not in table_caption_set)]
        # 「可接受标准」与表题之间常无正文段：在表题（或合格标准标题）前插入，避免回写到 RA-5 标题行
        if anchor == '试验成立标准' and items and not local_slots:
            cap_in_range = [c for c in table_caption_paras if r_start < c < r_end]
            ins_idx: Optional[int] = cap_in_range[0] if cap_in_range else None
            if ins_idx is None and r_end < len(doc.paragraphs):
                ins_idx = r_end
            if ins_idx is not None:
                cur = doc.paragraphs[ins_idx]
                for it in reversed(items):
                    text = str(it.get('text', '')).strip()
                    style_name = it.get('style', 'RA-正文')
                    new_p = cur.insert_paragraph_before(text)
                    skip_clear_elements.add(new_p._element)
                    cur = new_p
                    if styles.get(style_name):
                        try:
                            new_p.style = styles[style_name]
                        except Exception:
                            pass
                    if style_name == 'RA-正文':
                        _clear_list_numbering_paragraph(new_p)
                suitability_insert_count += len(items)
                print(f"  [INFO] 分区 {anchor} 在表题/下节标题前插入 {len(items)} 段（模板无试验成立正文槽位）")
                # 插入后段号整体后移，须重算后续分区（尤其「结果计算」「合格标准」）的锚点与可写槽位
                end_adj = end_idx + suitability_insert_count
                new_anchor_pos = _detect_anchor_positions(doc, start_idx, end_adj)

                # 修复：保留原始锚点位置，避免因锚点检测失败导致章节标题丢失
                # 如果新检测中缺失了原有锚点（特别是"合格标准"），使用原始位置
                for canonical in anchor_order:
                    if canonical in anchor_pos and canonical not in new_anchor_pos:
                        # 计算偏移量：插入的段落数
                        offset = suitability_insert_count
                        old_pos = anchor_pos[canonical]
                        # 如果原始锚点位置在插入点之后，需要加上偏移量
                        adjusted_pos = old_pos + offset if old_pos > r_start else old_pos
                        if adjusted_pos < len(doc.paragraphs):
                            # 验证该位置仍然是有效锚点（放宽条件：只要有段落结构即可）
                            para = doc.paragraphs[adjusted_pos]
                            # 即使段落被清空，只要位置合理就保留锚点
                            if adjusted_pos > start_idx and adjusted_pos < end_adj:
                                new_anchor_pos[canonical] = adjusted_pos

                anchor_pos = new_anchor_pos
                existing_sorted = sorted(
                    [(a, anchor_pos[a]) for a in anchor_order if a in anchor_pos],
                    key=lambda x: x[1],
                )
                ranges.clear()
                for ii, (aa, p_idx) in enumerate(existing_sorted):
                    n_idx = existing_sorted[ii + 1][1] if ii + 1 < len(existing_sorted) else end_adj
                    ranges[aa] = (p_idx, n_idx)
                writable_indices, table_caption_paras, para_to_table = _collect_writable_indices(
                    doc, start_idx, end_adj
                )
                table_caption_set = set(table_caption_paras)
                continue
        # 若分区内无槽位，回退使用锚点行本身作为唯一槽位（常见于“标题+内容同一行”的模板）
        if not local_slots and r_start < len(doc.paragraphs):
            local_slots = [r_start]

        overflow_texts: List[str] = []
        for i, item in enumerate(items):
            text = item.get('text', '').strip()
            style_name = item.get('style', 'RA-正文')
            if i >= len(local_slots):
                overflow_texts.append(text)
                continue
            target_para_idx = local_slots[i]
            para = doc.paragraphs[target_para_idx]

            # 关键修复：如果目标段落是RA-5/RA-6标题样式，跳过该槽位以保护标题
            para_style = para.style.name if para.style else ''
            if _is_protected_style(para_style):
                print(f"  [WARN] 跳过受保护的标题段落 idx={target_para_idx} style={para_style}")
                overflow_texts.append(text)
                continue

            actually_written_indices.add(target_para_idx)

            original_style = para.style
            para.clear()
            run = para.add_run(text)
            run.font.bold = False
            run.font.italic = False
            if styles.get(style_name):
                try:
                    para.style = styles[style_name]
                except Exception:
                    if original_style:
                        para.style = original_style
            elif original_style:
                para.style = original_style

            # RA-正文不挂 Word 多级列表（避免出现 1.1、2.1 与节标题编号叠加）
            if style_name == 'RA-正文':
                _clear_list_numbering_paragraph(para)

        # 槽位不足：操作步骤区多段正文插入新段，避免一段内 \\n 导致版式混乱；其它锚点仍语义合并
        if overflow_texts and local_slots:
            merged = [x for x in overflow_texts if str(x).strip()]
            if not merged:
                pass
            elif anchor == '操作步骤':
                n_fill = min(len(items), len(local_slots))
                anchor_last = local_slots[max(n_fill - 1, 0)]
                ref = doc.paragraphs[anchor_last]
                body_st = styles.get('RA-正文')
                for xt in merged:
                    ref = _insert_paragraph_after(ref, str(xt).strip(), style=body_st)
                    if body_st:
                        try:
                            _clear_list_numbering_paragraph(ref)
                        except Exception:
                            pass
                    pi = _paragraph_element_index(doc, ref)
                    if pi >= 0:
                        actually_written_indices.add(pi)
                    skip_clear_elements.add(ref._element)
                    procedure_insert_count += 1
                print(
                    f"  [INFO] 分区 {anchor} 槽位不足，已在区内插入 {len(merged)} 个独立段落（样品处理/操作步骤）"
                )
            else:
                last_idx = local_slots[-1]
                para = doc.paragraphs[last_idx]
                base = para.text.strip()
                base_text = base
                base_plain = base.rstrip('：')
                ov_skip = _anchor_heading_texts_to_skip(anchor)
                if base_plain in anchor_order or base_plain in ov_skip:
                    base_text = base_plain + '：'
                final_text = _merge_text_semantic(base_text, merged)
                para.clear()
                run = para.add_run(final_text)
                run.font.bold = False
                run.font.italic = False
                _clear_list_numbering_paragraph(para)
                actually_written_indices.add(last_idx)
                print(f"  [INFO] 分区 {anchor} 槽位不足，剩余内容已并入最后段落")

    end_idx_adj = end_idx + suitability_insert_count + procedure_insert_count
    _, table_caption_paras_adj, _ = _collect_writable_indices(doc, start_idx, end_idx_adj)

    # 清空多余段落
    # 注意：需要清空所有未被写入的段落，但不包括受保护的段落（如章节标题、表格标题）
    # 获取受保护的段落索引（章节标题、表格标题等不应被清空）
    protected_indices = set()
    for j in range(start_idx, min(end_idx_adj, len(doc.paragraphs))):
        if j < len(doc.paragraphs):
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ""
            if _is_protected_style(style):
                protected_indices.add(j)

    # 表格标题段落也应该被保护（已经在_is_protected_style中处理）
    # 额外保护：紧跟在表格标题后面的段落（通常是表格内容）也不应被清空
    for table_caption_idx in table_caption_paras_adj:
        if table_caption_idx < start_idx or table_caption_idx >= end_idx_adj:
            continue
        protected_indices.add(table_caption_idx)
        # 检查后面是否有紧跟的段落，如果有也要保护（通常是空段落）
        if table_caption_idx + 1 < min(end_idx_adj, len(doc.paragraphs)):
            next_para = doc.paragraphs[table_caption_idx + 1]
            if not next_para.text.strip():
                protected_indices.add(table_caption_idx + 1)

    # 额外保护：模板中原有的顶层编号标题段落（1~6）不清空（含「设备、材料、试剂」）
    title_keywords = list(anchor_order)
    for a in anchor_order:
        title_keywords.extend(_CANONICAL_ANCHOR_PHRASES.get(a, []))
    title_keywords = list(dict.fromkeys(title_keywords))
    for j in range(start_idx, min(end_idx_adj, len(doc.paragraphs))):
        para = doc.paragraphs[j]
        txt = para.text.strip()
        if not txt:
            continue
        ilvl = _get_paragraph_level(para)
        if ilvl == '0' and any(k in txt for k in title_keywords):
            protected_indices.add(j)

    # 清空start_idx到end_idx之间所有未被写入且不受保护的段落
    for j in range(start_idx, min(end_idx_adj, len(doc.paragraphs))):
        para = doc.paragraphs[j]
        if para._element in skip_clear_elements:
            continue
        # 勿清空模板中的图位及紧随其后的图题段（否则 _replace_section_images 无图可换）
        if _paragraph_has_image(para):
            continue
        if j > start_idx and _paragraph_has_image(doc.paragraphs[j - 1]):
            continue
        if j not in actually_written_indices and j not in protected_indices:
            para.clear()
            # 移除编号属性，避免出现空内容的编号行
            pPr = para._element.pPr
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)

    # 新增骨架章节常只有空 RA-正文，占位不足会导致小节整块为空；在清空之后补插关键正文段
    inserted_after_clear = _ensure_anchor_sections_not_empty(
        doc,
        start_idx,
        end_idx_adj,
        filtered_refined,
        styles,
    )
    if inserted_after_clear:
        end_idx_adj += inserted_after_clear

    # 仅替换当前章节范围内的表格（插入段落后重算表格索引与章节上界）
    section_table_indices_adj = _get_table_indices_in_section(doc, start_idx, end_idx_adj)
    print(f"  [DEBUG] 章节范围: start_idx={start_idx}, end_idx={end_idx} (写入后上界 end_idx_adj={end_idx_adj})")
    print(f"  [DEBUG] 章节内表格索引: {section_table_indices_adj}")
    _replace_matching_tables_in_section(
        doc, replacement_tables or [], section_table_indices_adj, start_idx, end_idx_adj
    )

    # 按需求：仅替换现有图片，不新增图片（合并 SOP 多区间提取：系统适用性/典型图谱/结果计算等）
    image_map = _load_sop_image_blob_map(sop_path)
    body_style = styles.get('RA-正文')
    section_images = replacement_section_images or {}
    merged_imgs = _flatten_section_images_dedup(section_images)
    n_replaced = _replace_section_images(
        doc,
        start_idx,
        end_idx_adj,
        merged_imgs,
        image_map,
        body_style=body_style
    )
    if n_replaced:
        print(f"  [INFO] 已从 SOP 替换模板内图片 {n_replaced} 张（合并 {len(merged_imgs)} 张候选图）")

    # 最后兜底：修复 RA-5/RA-6 标题段落夹带正文（应仅保留标题）
    inserted_by_sanitize = _sanitize_ra_heading_paragraphs(doc, start_idx, end_idx_adj, styles)
    if inserted_by_sanitize:
        end_idx_adj += inserted_by_sanitize

    doc.save(output_path)
    return output_path


def find_method_section_range(
    doc: Document,
    section_name: str,
    stop_keywords: List[str]
) -> Tuple[int, int]:
    """
    在模板中查找指定方法章节的起止段落索引。

    匹配逻辑：
    1. 段落文本必须非空且长度 >= 2
    2. section_name 出现在 text 中（正向包含）
    3. 段落样式是标题类：3.2.S* / Heading 2+ / RA-5级标题 等

    Returns:
        (start_idx, end_idx)，未找到返回 (-1, -1)
    """
    start_idx = -1

    def _norm_name(x: str) -> str:
        t = str(x or '').strip().lower()
        t = re.sub(r'[\s\u3000]', '', t)
        t = re.sub(r'[()（）\\[\\]【】<>《》:：,，.。\\-_/]', '', t)
        return t

    def _section_candidates(name: str) -> List[str]:
        cands = [name.strip()]
        # 去掉括号内容：肽图（RP-UPLC） -> 肽图
        no_bracket = re.sub(r'（[^）]*）|\([^)]*\)', '', name).strip()
        if no_bracket and no_bracket not in cands:
            cands.append(no_bracket)
        # 常见同义写法
        if '蛋白含量' in name and '蛋白质含量' not in name:
            cands.append(name.replace('蛋白含量', '蛋白质含量'))
        if '蛋白质含量' in name and '蛋白含量' not in name:
            cands.append(name.replace('蛋白质含量', '蛋白含量'))
        return [c for c in cands if c]

    section_cands = _section_candidates(section_name)
    norm_cands = [_norm_name(c) for c in section_cands]

    def _name_match(text: str) -> bool:
        if not text:
            return False
        t_norm = _norm_name(text)
        for c, c_norm in zip(section_cands, norm_cands):
            if c in text or text in c:
                return True
            if c_norm and t_norm and (c_norm in t_norm or t_norm in c_norm):
                return True
        return False

    # 第一轮：匹配 section_name / 同义候选 与标题文本
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if not text or len(text) < 2:
            continue
        if 'toc' in style.lower() or 'table of' in style.lower():
            continue
        if _name_match(text):
            # 标题长度不应过长（避免匹配到正文段落）
            if len(text) > max(len(section_name), 4) + 40:
                continue
            if (style.startswith('3.2.S') or
                ('Heading' in style and style != 'Heading 1') or
                'RA-5' in style or 'RA-6' in style):
                start_idx = i
                break

    # 第二轮：放宽匹配 — text == section_name 或 text 是 section_name 的子串
    if start_idx < 0:
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            style = para.style.name if para.style else ""
            if not text or len(text) < 2:
                continue
            if 'toc' in style.lower() or 'table of' in style.lower():
                continue
            if _name_match(text):
                start_idx = i
                break

    if start_idx < 0:
        return -1, -1

    def _is_heading_style(style_name: str) -> bool:
        return (
            style_name.startswith('3.2.S') or
            ('Heading' in style_name and style_name != 'Heading 1') or
            'RA-5' in style_name or
            'RA-6' in style_name
        )

    section_style = doc.paragraphs[start_idx].style.name
    end_idx = len(doc.paragraphs)
    for i in range(start_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if not text or 'toc' in style.lower() or 'table of' in style.lower():
            continue

        # stop 关键词优先：显式边界最可靠
        for stop_kw in stop_keywords:
            if stop_kw in text and len(text) < len(stop_kw) + 30:
                if _is_heading_style(style) or style == section_style:
                    return start_idx, i

        # 兜底边界：若 stop 未匹配，遇到同级标题即视为下一方法章节开始
        if style == section_style:
            return start_idx, i

    # 未找到下一同级标题时，默认当前章节到文档结尾（通常为最后一章）
    return start_idx, end_idx


def integrate_sop_into_template(
    template_path: str,
    sop_path: str,
    section_name: str,
    stop_keywords: List[str],
    output_path: Optional[str] = None,
    style_mode: str = 'flat'
) -> str:
    """
    将 SOP 内容整合到分析方法模板的指定章节。

    Args:
        template_path: 分析方法标准模板路径
        sop_path: SOP 文档路径
        section_name: 模板中方法章节名称（如 "纯度（SEC-HPLC）"）
        stop_keywords: 下一章节关键词，用于定位结束
        output_path: 输出路径，默认带时间戳
        style_mode: 'flat' | 'hierarchical'

    Returns:
        输出文件路径
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"模板不存在: {template_path}")
    if not os.path.exists(sop_path):
        raise FileNotFoundError(f"SOP 不存在: {sop_path}")

    if not output_path:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        base = os.path.splitext(os.path.basename(template_path))[0]
        output_path = os.path.join(os.path.dirname(template_path), f"{base}_{ts}.docx")

    if os.path.normpath(os.path.abspath(template_path)) != os.path.normpath(os.path.abspath(output_path)):
        shutil.copy(template_path, output_path)
    doc = Document(output_path)

    # 提取 SOP 第四章内容
    extractor = SOPExtractor(sop_path)
    sop_data = extractor.extract_procedure_chapter4(section_name)

    # 构建整合内容
    content_list = build_integrated_content(sop_data, style_mode)

    # 查找模板章节
    start_idx, end_idx = find_method_section_range(doc, section_name, stop_keywords)
    if start_idx < 0:
        raise ValueError(f"未找到章节: {section_name}")

    # 缓存样式
    styles = {}
    for name in ['RA-5级标题', 'RA-6级标题', 'RA-正文']:
        try:
            styles[name] = doc.styles[name]
        except Exception:
            styles[name] = None

    # 写入内容
    valid_items = [x for x in content_list if x.get('text', '').strip()]
    writable_indices, table_caption_paras, para_to_table = _collect_writable_indices(doc, start_idx, end_idx)

    # 构建表格标题段落的集合，用于在写入时跳过
    table_caption_set = set(table_caption_paras)

    next_section_para = doc.paragraphs[end_idx] if end_idx < len(doc.paragraphs) else None

    # 写入时跳过表格标题段落，避免在表格标题和表格之间插入内容
    written_count = 0
    for i, item in enumerate(valid_items):
        text = item.get('text', '').strip()
        style_name = item.get('style', 'RA-正文')

        # 跳过表格标题段落，不在它们的位置写入内容
        while written_count < len(writable_indices) and writable_indices[written_count] in table_caption_set:
            written_count += 1

        if written_count < len(writable_indices):
            para = doc.paragraphs[writable_indices[written_count]]
            written_count += 1
            # 保存原有段落样式
            original_style = para.style
            # 清除段落内容
            para.clear()
            # 添加新文本，明确不加粗
            run = para.add_run(text)
            run.font.bold = False
            run.font.italic = False
            # 恢复或设置段落样式
            if styles.get(style_name):
                try:
                    para.style = styles[style_name]
                except Exception:
                    if original_style:
                        para.style = original_style
            elif original_style:
                para.style = original_style
            if style_name == 'RA-正文':
                _clear_list_numbering_paragraph(para)
        else:
            # 在下一章节前插入新段落
            if next_section_para is not None:
                new_para = next_section_para.insert_paragraph_before(text)
            else:
                new_para = doc.add_paragraph(text)
            # 设置新段落样式并确保不加粗
            if styles.get(style_name):
                try:
                    new_para.style = styles[style_name]
                except Exception:
                    pass
            # 确保新段落的run不加粗
            for run in new_para.runs:
                run.font.bold = False
                run.font.italic = False
            if style_name == 'RA-正文':
                _clear_list_numbering_paragraph(new_para)

    # 清空多余可写段落（保护表图编号、标题）
    # 重新计算实际写入的段落索引
    actually_written_indices = set()
    wc = 0
    for i, item in enumerate(valid_items):
        # 跳过表格标题段落
        while wc < len(writable_indices) and writable_indices[wc] in table_caption_set:
            wc += 1
        if wc < len(writable_indices):
            actually_written_indices.add(writable_indices[wc])
            wc += 1

    # 清空未写入且不受保护的段落
    for j in range(start_idx, min(end_idx, len(doc.paragraphs))):
        if j < len(doc.paragraphs):
            para = doc.paragraphs[j]
            style = para.style.name if para.style else ""
            is_protected = _is_protected_style(style)
            if _paragraph_has_image(para):
                continue
            if j > start_idx and _paragraph_has_image(doc.paragraphs[j - 1]):
                continue
            if j not in actually_written_indices and not is_protected:
                para.clear()

    # 直写模式同样迁移“系统适用性试验/结果与计算”图片与图题
    image_map = _load_sop_image_blob_map(sop_path)
    body_style = styles.get('RA-正文')
    sec_imgs = sop_data.get('section_images', {}) if isinstance(sop_data, dict) else {}

    suit_anchor = _find_subsection_end_anchor(
        doc, start_idx, end_idx,
        target_keywords=['4 试验成立标准', '试验成立标准', '系统适用性', '可接受标准'],
        next_keywords=['5 结果计算', '结果计算']
    )
    _insert_section_images(
        doc, suit_anchor, sec_imgs.get('suitability_criteria', []), image_map, body_style=body_style, max_images=4
    )

    doc.save(output_path)
    return output_path


def _detect_template_sections(doc: Document, start_idx: int, end_idx: int) -> set:
    """检测模板章节中存在的小节标题。

    Returns:
        存在的小节名称集合，如 {'结果计算', '合格标准'}
    """
    sections = set()
    # 小节关键词映射
    section_keywords = {
        '原理': ['原理', 'principle'],
        '材料和设备': [
            '主要材料和设备', '设备、材料、试剂', '材料和设备', '设备、材料', '材料与设备', '试剂', 'reagent',
        ],
        '操作步骤': ['操作步骤', 'operation', 'procedure'],
        '样品处理': ['样品处理', 'sample preparation'],
        '试验成立标准': ['可接受标准', '试验成立标准', '系统适用性', 'suitability'],
        '结果计算': ['计算公式', '结果计算', 'result calculation'],
        '合格标准': ['合格标准', 'acceptance criteria'],
    }

    for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text.strip().lower()
        if not text:
            continue

        # 检查是否包含某个小节的关键词
        for section_name, keywords in section_keywords.items():
            for keyword in keywords:
                if keyword.lower() in text:
                    sections.add(section_name)
                    break

    return sections


# 与 _CANONICAL_ANCHOR_PHRASES 一致：模板实际标题 → 逻辑小节名（供 build_refined_content 独立标题判断）
_SECTION_INDEP_TITLE_ALIASES: Dict[str, List[str]] = {
    '原理': ['原理'],
    '材料和设备': ['主要材料和设备', '设备、材料、试剂', '材料和设备', '材料与设备'],
    '操作步骤': ['操作步骤'],
    '试验成立标准': ['试验成立标准', '系统适用性', '可接受标准'],
    '结果计算': ['计算公式', '结果计算'],
    '合格标准': ['合格标准'],
}


def _detect_independent_title_sections(
    doc: Document,
    start_idx: int,
    end_idx: int
) -> set:
    """检测模板中哪些小节有独立的标题行（不与内容合并）。

    RA-5/RA-6 级标题常无列表 numPr，需同时认 ilvl=0 与「标题样式」。
    """
    independent_sections: set = set()

    for i in range(start_idx, min(end_idx, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if not text:
            continue

        ilvl = _get_paragraph_level(para)
        is_lv0 = ilvl == '0'
        st = para.style.name if para.style else ''
        is_subsection_style = 'RA-5' in st or 'RA-6' in st or ('RA-' in st and '标题' in st)

        plain = text.rstrip('：')
        if not (is_lv0 or is_subsection_style):
            continue

        for canonical, aliases in _SECTION_INDEP_TITLE_ALIASES.items():
            if plain in aliases:
                independent_sections.add(canonical)
                print(f'  [DEBUG] 检测到独立标题: {canonical} ({plain!r}, lv0={is_lv0}, style={st!r})')
                break

    return independent_sections


def _filter_refined_by_template_sections(
    refined_data: Dict[str, Any],
    template_sections: set,
    style_mode: str
) -> Dict[str, Any]:
    """根据模板存在的小节过滤refined内容。

    只保留模板中存在的小节对应的内容。

    Args:
        refined_data: 精简后的数据
        template_sections: 模板中存在的小节集合
        style_mode: 'flat' | 'hierarchical'

    Returns:
        过滤后的精简数据
    """
    filtered = {}

    # 小节映射
    section_mapping = {
        'principle': '原理',
        'materials_and_equipment': '材料和设备',
        'sample_prep': '操作步骤',  # 样品处理包含在操作步骤中
        'procedure': '操作步骤',
        'suitability_criteria': '试验成立标准',
        'result_calculation': '结果计算',
        'acceptance_criteria': '合格标准',
    }

    for key, value in refined_data.items():
        mapped_section = section_mapping.get(key, '')

        # 如果映射的小节在模板中存在，或者模板为空（保留所有内容），则保留
        if mapped_section in template_sections or not template_sections:
            filtered[key] = value
        else:
            # 模板中没有该小节，跳过
            print(f'  [SKIP] {mapped_section} - 模板中无此小节')

    return filtered
