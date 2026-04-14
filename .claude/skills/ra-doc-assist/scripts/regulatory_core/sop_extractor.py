# -*- coding: utf-8 -*-
"""
SOP 内容提取器

从标准操作规程（SOP）文档中提取结构化内容。
"""
from typing import Dict, List, Optional, Any
import re
import hashlib
import os
from docx import Document


class SOPExtractor:
    """SOP 内容提取器 - 同时提取段落和表格"""

    def __init__(self, sop_path: str):
        """
        初始化提取器

        Args:
            sop_path: SOP 文档路径
        """
        self.doc = Document(sop_path)
        self.paragraphs = []
        self.tables = []
        self.table_xmls = []
        self._image_rel_cache = self._build_image_rel_cache()
        self._extract_all()

    def _build_image_rel_cache(self) -> Dict[str, Dict[str, Any]]:
        """构建图片关系缓存: rel_id -> 图片元信息。"""
        cache: Dict[str, Dict[str, Any]] = {}
        rels = getattr(self.doc.part, 'rels', {})
        for rel_id, rel in rels.items():
            reltype = str(getattr(rel, 'reltype', '') or '')
            if '/image' not in reltype:
                continue
            target_part = getattr(rel, 'target_part', None)
            if not target_part:
                continue
            blob = getattr(target_part, 'blob', b'') or b''
            partname = str(getattr(target_part, 'partname', '') or '')
            content_type = str(getattr(target_part, 'content_type', '') or '')
            cache[rel_id] = {
                'rel_id': rel_id,
                'filename': os.path.basename(partname) if partname else '',
                'content_type': content_type,
                'size_bytes': len(blob),
                'sha1': hashlib.sha1(blob).hexdigest() if blob else '',
            }
        return cache

    def _extract_images_from_paragraph(self, para) -> List[Dict[str, Any]]:
        """提取段落中的图片引用。"""
        refs: List[Dict[str, Any]] = []
        seen = set()
        for node in para._element.iter():
            tag = str(getattr(node, 'tag', '') or '')
            if not tag.endswith('}blip'):
                continue
            rel_id = ''
            for attr_key, attr_val in getattr(node, 'attrib', {}).items():
                key = str(attr_key)
                if key.endswith('}embed') or key.endswith('}link'):
                    rel_id = str(attr_val)
                    break
            if not rel_id or rel_id in seen:
                continue
            seen.add(rel_id)
            refs.append(dict(self._image_rel_cache.get(rel_id, {'rel_id': rel_id})))
        return refs

    def _extract_all(self):
        """提取所有段落和表格内容"""
        # 提取段落
        for idx, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            style = para.style.name if para.style else ""
            images = self._extract_images_from_paragraph(para)
            self.paragraphs.append({
                'index': idx,
                'text': text,
                'style': style,
                'images': images,
            })

        # 提取表格
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            self.tables.append(table_data)
            self.table_xmls.append(table._tbl.xml)

    def _looks_like_figure_caption(self, text: str) -> bool:
        t = str(text or '').strip()
        if not t:
            return False
        return bool(re.match(r'^(图|Figure|FIGURE)\s*[0-9一二三四五六七八九十IVXivx\-\.：: ]+', t))

    def extract_section_images_with_captions(
        self,
        section_keywords: List[str],
        end_keywords: List[str] = None
    ) -> List[Dict[str, Any]]:
        """
        提取指定章节内图片及其标题（图题）。
        返回:
        [
          {"rel_id":"rId8","sha1":"...","filename":"image1.png","caption":"图 1 ..."},
        ]
        """
        in_section = False
        out: List[Dict[str, Any]] = []
        seen = set()

        for i, para in enumerate(self.paragraphs):
            text = para.get('text', '')
            style = str(para.get('style', '')).lower()
            if 'toc' in style:
                continue

            images = para.get('images', []) or []

            if not in_section:
                if any(kw in text for kw in section_keywords):
                    in_section = True
                elif images:
                    # 图常在图题上方：本段仅有图、后续若干段内出现小节关键词时，仍视为该小节内插图
                    look = ' '.join(
                        str(self.paragraphs[j].get('text', '') or '')
                        for j in range(i, min(i + 10, len(self.paragraphs)))
                    )
                    if any(kw in look for kw in section_keywords):
                        in_section = True
                else:
                    continue

            if in_section and end_keywords:
                if any((kw in text and len(text) < 50) for kw in end_keywords):
                    break

            if not images:
                continue

            # 优先同段图题，否则向后看若干段（含「图3 …典型图谱」类）
            caption = text if self._looks_like_figure_caption(text) else ''
            if not caption:
                for j in range(i + 1, min(i + 5, len(self.paragraphs))):
                    t2 = self.paragraphs[j].get('text', '')
                    if self._looks_like_figure_caption(t2):
                        caption = t2
                        break
            if not caption:
                for j in range(i + 1, min(i + 5, len(self.paragraphs))):
                    t2 = str(self.paragraphs[j].get('text', '') or '').strip()
                    if len(t2) > 4 and any(
                        k in t2 for k in ('典型图谱', '系统适用', '谱图', '色谱图', 'Figure', 'FIGURE')
                    ):
                        caption = t2
                        break

            for img in images:
                rel_id = str(img.get('rel_id', '') or '')
                sha1 = str(img.get('sha1', '') or '')
                key = (rel_id, sha1)
                if key in seen:
                    continue
                seen.add(key)
                out.append({
                    'rel_id': rel_id,
                    'sha1': sha1,
                    'filename': img.get('filename', ''),
                    'caption': caption,
                })

        return out

    def get_section_content(
        self,
        section_keywords: List[str],
        end_keywords: List[str] = None
    ) -> List[str]:
        """获取指定章节的内容

        Args:
            section_keywords: 章节开始关键词列表
            end_keywords: 章节结束关键词列表

        Returns:
            章节内容行列表
        """
        def heading_depth(text: str) -> int:
            """
            返回阿拉伯数字分级标题深度：
            - 例如 "4.4" -> 2, "4.4.4" -> 3, "4.4.4.1" -> 4
            - 非分级标题返回 0
            """
            t = str(text or '').strip()
            m = re.match(r'^(\d+(?:\.\d+){1,})(?:[\.、．\s]|$)', t)
            if not m:
                return 0
            return len([x for x in m.group(1).split('.') if x != ''])

        content = []
        in_section = False
        skip_deep_level_block = False

        for para in self.paragraphs:
            text = para['text']
            style = para['style'].lower()

            # 跳过目录样式
            if 'toc' in style:
                continue

            # 检查是否进入目标章节
            if not in_section:
                is_section_start = False
                for kw in section_keywords:
                    # 改进匹配逻辑：关键词在段落中即可，不要求开头或完全相等
                    if kw in text:
                        is_section_start = True
                        break
                if is_section_start:
                    in_section = True
                    continue

            # 检查是否离开目标章节
            if in_section and end_keywords:
                is_section_end = False
                for kw in end_keywords:
                    # 同样改进结束检测
                    if kw in text and len(text) < 50:  # 结束标题通常较短
                        is_section_end = True
                        break
                if is_section_end:
                    break

            # 跳过第4层及以下分级标题块：仅保留至第3层级
            depth = heading_depth(text)
            if in_section:
                if depth >= 4:
                    skip_deep_level_block = True
                    continue
                # 遇到 1~3 级分级标题，说明已回到允许层级，结束跳过状态
                if depth > 0 and depth <= 3:
                    skip_deep_level_block = False
                if skip_deep_level_block:
                    continue

            if in_section and text:
                content.append(text)

        return content

    def get_table_by_header(self, header_keywords: List[str]) -> Optional[List[List[str]]]:
        """通过表头关键词获取表格

        Args:
            header_keywords: 表头关键词列表

        Returns:
            表格数据，或 None
        """
        for table in self.tables:
            if not table:
                continue
            first_row = " ".join(table[0]) if table[0] else ""
            if any(kw in first_row for kw in header_keywords):
                return table
        return None

    def get_chromatography_conditions(self) -> Dict[str, str]:
        """提取色谱条件表格"""
        for table in self.tables:
            if not table:
                continue
            flat_text = " ".join([" ".join(row) for row in table])
            if '色谱柱' in flat_text and ('检测' in flat_text or '波长' in flat_text):
                return self._parse_conditions_table(table)
        return {}

    def _parse_conditions_table(self, table: List[List[str]]) -> Dict[str, str]:
        """解析条件表格为字典"""
        conditions = {}
        for row in table:
            if len(row) >= 2:
                key = row[0].strip()
                value = row[1].strip() if len(row) > 1 else ""
                if key and value:
                    conditions[key] = value
        return conditions

    def get_acceptance_criteria(self) -> Optional[List[List[str]]]:
        """提取可接受标准表格"""
        for table in self.tables:
            if not table:
                continue
            flat_text = " ".join([" ".join(row) for row in table])
            if '可接受标准' in flat_text or '名称' in flat_text:
                if any(kw in flat_text for kw in ['空白', '系统适用性', '供试品']):
                    return table
        return None

    def get_equipment_materials_tables(self) -> List[str]:
        """
        提取实验设备/仪器、材料相关表格内容，转为文本行列表。
        用于补充 materials_and_equipment，防止设备内容丢失。
        """
        lines: List[str] = []

        def _cell_norm(s: str) -> str:
            return re.sub(r"[\s\u3000]+", "", str(s or "")).strip().lower()

        def _looks_like_vendor(s: str) -> bool:
            t = str(s or "").strip()
            if not t:
                return False
            if re.search(
                r"(有限公司|有限责任公司|股份有限公司|集团|公司|GmbH|AG|Inc\.?|Ltd\.?|Co\.?,?\s*Ltd\.?|Corporation|Corp\.?)\b",
                t,
                flags=re.IGNORECASE,
            ):
                return True
            if re.search(
                r"(Eppendorf|Thermo|Beckman|Sartorius|Merck|Waters|Agilent|Tecan|Mettler|Toledo|ProteinSimple)",
                t,
                re.I,
            ):
                return True
            return False

        def _is_noise_name(name: str) -> bool:
            n = str(name or "").strip()
            if not n:
                return True
            if len(n) <= 1:
                return True
            if n in ("—", "-", "NA", "N/A", "无", "无。", "None"):
                return True
            # 纯编号/纯数字
            if re.fullmatch(r"[0-9\.\-_/]+", n):
                return True
            # 纯公司名/供应商名
            if _looks_like_vendor(n) and len(n) <= 40:
                return True
            return False

        def _looks_like_equipment_name(name: str) -> bool:
            t = str(name or "").strip()
            if not t:
                return False
            # 中文设备关键词
            if re.search(r"(仪|机|柜|台|天平|离心机|培养箱|酶标|移液器|混匀|恒温|色谱|分光|电泳|光度计)", t):
                return True
            # 英文设备关键词（偏少量）
            if re.search(r"(incubator|reader|pipette|centrifuge|balance|biosafety|cabinet|workbench)", t, re.I):
                return True
            return False

        def _looks_like_reagent_or_consumable_name(name: str) -> bool:
            t = str(name or "").strip()
            if not t:
                return False
            # 中文试剂/耗材关键词
            if re.search(r"(溶液|缓冲|培养基|试剂|对照|参比|标准品|耗材|色谱柱|滤膜|孔板|试剂盒)", t):
                return True
            # 英文常见试剂/耗材
            if re.search(r"(PBS|FBS|EDTA|Trypsin|Puromycin|kit|buffer|medium|plate|column)", t, re.I):
                return True
            return False

        def _classify_table_by_header(header_cells: List[str], flat_text: str) -> str:
            """
            返回：'equipment' | 'reagent' | ''（不处理）
            """
            header_norm = " ".join(_cell_norm(x) for x in header_cells if str(x).strip())
            flat_norm = _cell_norm(flat_text)

            # 排除审批/签字/记录等表
            if any(k in flat_norm for k in ("signature", "reviewed", "drafted", "批准", "审核", "签字", "签名", "修订历史")):
                return ""
            # 排除条件表/可接受标准表（这些由专用逻辑处理）
            if any(k in flat_norm for k in ("色谱柱", "检测波长", "流动相a", "流动相b", "可接受标准")):
                return ""

            equip_kws = ("设备", "仪器", "instrument", "equipment")
            reagent_kws = ("试剂", "溶液", "solution", "reagent", "耗材", "consumable", "材料", "material")

            # 仅允许“表头驱动”分类，避免把术语表/缩略语表误识别为设备/试剂表
            if any(k in header_norm for k in equip_kws):
                return "equipment"
            if any(k in header_norm for k in reagent_kws):
                return "reagent"
            return ""

        def _choose_name_col(header_cells: List[str]) -> int:
            """
            优先选择表头含“名称/name”的列，否则用第 1 列。
            """
            for idx, c in enumerate(header_cells):
                cn = _cell_norm(c)
                if ("名称" in cn) or (cn == "name") or ("name" in cn and len(cn) <= 12):
                    return idx
            return 0

        def _is_vendor_col(header_cell: str) -> bool:
            cn = _cell_norm(header_cell)
            return any(k in cn for k in ("厂家", "生产厂家", "供应商", "厂商", "manufact", "vendor", "supplier", "公司", "company"))

        def _is_model_col(header_cell: str) -> bool:
            cn = _cell_norm(header_cell)
            return any(k in cn for k in ("型号", "model", "规格", "spec", "货号", "catalog", "cat", "lot", "批号"))

        for table in self.tables:
            if not table or len(table) < 2:
                continue
            flat_text = " ".join([" ".join(str(c) for c in row) for row in table])

            header_cells = [str(c).strip() for c in (table[0] or [])]
            table_kind = _classify_table_by_header(header_cells, flat_text)
            if not table_kind:
                continue

            name_col = _choose_name_col(header_cells)
            vendor_cols = {i for i, h in enumerate(header_cells) if _is_vendor_col(h)}
            model_cols = {i for i, h in enumerate(header_cells) if _is_model_col(h)}

            for row in table[1:]:
                if not row:
                    continue
                # 保持原列位，避免因空单元格错位
                row_cells = [str(c).strip() for c in row]
                if name_col >= len(row_cells):
                    continue
                name = row_cells[name_col].strip()
                if _is_noise_name(name):
                    continue

                # 若 name 列为空但首列有值，兜底用首列
                if not name and row_cells and row_cells[0].strip():
                    name = row_cells[0].strip()
                if _is_noise_name(name):
                    continue

                # 如果 name 看起来像“设备名：厂商”，只保留设备名
                if "：" in name or ":" in name:
                    left, right = re.split(r"[：:]", name, maxsplit=1)
                    if _looks_like_vendor(right.strip()):
                        name = left.strip()

                if _is_noise_name(name):
                    continue

                # 对于 equipment：优先写成“主要设备：xxx”，便于后续归类；reagent 直接输出名称
                if table_kind == "equipment":
                    if not _looks_like_equipment_name(name):
                        # 过滤术语表/缩略语表等误入项
                        continue
                    lines.append(f"主要设备：{name}")
                else:
                    if not _looks_like_reagent_or_consumable_name(name):
                        continue
                    # reagent/consumable：不拼接厂商/型号列，只保留名称
                    lines.append(name)
        return lines

    def _normalize_materials_equipment_lines(self, lines: List[str]) -> List[str]:
        """
        归一化“材料和设备”字段：
        - 仅当输入本身存在参比品相关信息时，才补一条“标准物质：...”
        - 标准物质优先取系统适用性样品/溶液对应的参比品信息
        - 去重并保持顺序
        """
        if not lines:
            return []

        normalized = [str(x).strip() for x in lines if str(x).strip()]

        # 1) 优先从“系统适用性样品/溶液”上下文提取标准物质
        standard_candidate = None
        for i, line in enumerate(normalized):
            if any(k in line for k in ['系统适用性样品', '系统适用性溶液']):
                # 优先看下一行
                if i + 1 < len(normalized):
                    nxt = normalized[i + 1]
                    if any(k in nxt for k in ['工作参比品', '参比品', '对照品', 'reference']):
                        standard_candidate = nxt.rstrip('。')
                        break
            # 兜底：行内直接包含参比品信息
            if any(k in line for k in ['工作参比品', '参比品', '对照品', 'reference']):
                standard_candidate = line.rstrip('。')
                break

        if standard_candidate:
            std_line = f"标准物质：{standard_candidate}。"
            if not any(x.startswith('标准物质：') for x in normalized):
                normalized.append(std_line)

        # 2) 去重（保持顺序）
        dedup = []
        seen = set()
        for line in normalized:
            key = line.replace(' ', '')
            if key in seen:
                continue
            seen.add(key)
            dedup.append(line)
        return dedup

    def extract_method_related_tables(self) -> List[Dict[str, Any]]:
        """
        提取与可接受标准/梯度/结果计算相关的 SOP 表格，用于替换或补插。
        返回结构: [{"header": [...], "rows": [[...], ...]}]
        """
        def collect_acceptance_anchor_table_indices() -> tuple:
            """
            基于文档顺序提取，区分系统适用性标准和产品合格标准：
            - 段落命中"可接受标准"
            - 后续紧邻表格的表头也含"可接受标准"
            - 向前查找最近的一个包含"系统适用性"的段落，如果找到则归类为suitability_criteria
            - 如果没有找到"系统适用性"段落，则归类为acceptance_criteria

            Returns:
                (suitability_indices, acceptance_indices)
            """
            suitability_indices = set()
            acceptance_indices = set()
            body = self.doc.element.body
            table_idx = -1
            hit_acceptance_para = False

            # 记录所有段落索引和文本
            para_texts = {}
            para_idx = -1
            for child in body.iterchildren():
                tag = child.tag.rsplit('}', 1)[-1]
                if tag == 'p':
                    para_idx += 1
                    text = "".join(t for t in child.itertext() if t).strip()
                    para_texts[para_idx] = text

            table_idx = -1
            para_idx = -1
            hit_acceptance_para = False

            for child in body.iterchildren():
                tag = child.tag.rsplit('}', 1)[-1]
                if tag == 'p':
                    para_idx += 1
                    text = para_texts.get(para_idx, '')
                    hit_acceptance_para = ('可接受标准' in text)
                elif tag == 'tbl':
                    table_idx += 1
                    if not hit_acceptance_para:
                        continue
                    if table_idx < len(self.tables) and self.tables[table_idx]:
                        header_text = " ".join(str(c).strip() for c in self.tables[table_idx][0])
                        if '可接受标准' in header_text:
                            # 向前查找最近的一个包含"系统适用性"的段落
                            found_suitability = False
                            for check_idx in range(para_idx - 1, max(-1, para_idx - 10), -1):
                                if check_idx in para_texts and '系统适用性' in para_texts[check_idx]:
                                    found_suitability = True
                                    break

                            if found_suitability:
                                suitability_indices.add(table_idx)
                            else:
                                acceptance_indices.add(table_idx)
                    # 仅消费紧邻的第一张表
                    hit_acceptance_para = False
            return suitability_indices, acceptance_indices

        def collect_calculation_anchor_table_info() -> Dict[int, str]:
            """
            基于文档顺序提取：
            - 仅在 SOP 第4章程序中的“数据处理/计算/结果计算”段落命中
            - 后续紧邻表格，记录该表索引与标题段落文本
            """
            def find_section_spans(section_keywords: List[str], end_keywords: List[str]) -> List[tuple]:
                spans: List[tuple] = []
                in_section = False
                start_idx = -1
                for para in self.paragraphs:
                    idx = int(para.get('index', -1))
                    text = str(para.get('text', '') or '')
                    style = str(para.get('style', '') or '').lower()
                    if 'toc' in style:
                        continue
                    if not in_section:
                        if any(kw in text for kw in section_keywords):
                            in_section = True
                            start_idx = idx
                            continue
                    else:
                        if end_keywords and any((kw in text and len(text) < 50) for kw in end_keywords):
                            spans.append((start_idx, idx))
                            in_section = False
                            start_idx = -1
                if in_section and start_idx >= 0:
                    max_idx = max((int(p.get('index', -1)) for p in self.paragraphs), default=start_idx)
                    spans.append((start_idx, max_idx + 1))
                return spans

            # 仅允许第4章程序里的“数据处理 + 计算”范围
            calc_spans = find_section_spans(
                ['数据处理', 'Data Processing', '计算', 'Calculate', '结果计算', '浓度计算', '数据计算'],
                ['系统适用性', 'System Suitability', '可接受标准', 'Acceptable Standard', '结果判定']
            )
            allowed_para_indices = set()
            for s, e in calc_spans:
                allowed_para_indices.update(range(s, e))

            mapping: Dict[int, str] = {}
            para_text_map = {int(p.get('index', -1)): str(p.get('text', '') or '') for p in self.paragraphs}
            body = self.doc.element.body
            table_idx = -1
            calc_title = ''
            hit_calc_para = False
            para_idx = -1
            for child in body.iterchildren():
                tag = child.tag.rsplit('}', 1)[-1]
                if tag == 'p':
                    para_idx += 1
                    text = para_text_map.get(para_idx, "").strip()
                    # 离开计算章节范围则关闭追踪
                    if para_idx not in allowed_para_indices:
                        hit_calc_para = False
                        calc_title = ''
                        continue

                    # 在计算章节内，命中计算关键词后进入追踪状态
                    if any(k in text for k in ['结果与计算', '结果计算', '数据处理', '计算']):
                        hit_calc_para = True
                        calc_title = text
                        continue

                    # 计算章节内若已进入追踪，持续更新“最近标题/说明行”作为表题
                    if hit_calc_para and text:
                        calc_title = text
                elif tag == 'tbl':
                    table_idx += 1
                    if hit_calc_para:
                        mapping[table_idx] = calc_title or '结果计算'
                        # 保持追踪，支持一个计算段落下连续多张表
            return mapping

        def is_acceptance_related(rows: List[List[str]], flat_text: str) -> bool:
            if not rows:
                return False
            header_text = " ".join(str(c).strip() for c in rows[0]).lower()
            flat_lower = flat_text.lower()
            # sample_kws = ['空白', '供试品', '系统适用性', '样品']  # 兼容纯度场景，暂时停用
            # criteria_kws = ['RSD', '≤', '≥', '应', '无干扰', '一致', '差值', '相关系数', '不低于', '不高于']  # 暂时停用
            # 目标：名称/样品名称 + 可接受标准
            if ('可接受标准' in header_text) and any(k in header_text for k in ['名称', '样品']):
                return True
            # 英文兼容
            if ('acceptable standard' in flat_lower) and ('name' in flat_lower or 'sample' in flat_lower):
                return True
            # 兼容纯度等 SOP：表头未写“可接受标准”，但属于系统适用性判定表（暂时停用）
            # if any(k in header_text for k in ['名称', '样品']) and any(k in flat_text for k in sample_kws):
            #     if any(k in flat_text for k in criteria_kws):
            #         return True
            return False

        def is_gradient_table(rows: List[List[str]], flat_text: str) -> bool:
            if not rows:
                return False
            header_text = " ".join(str(c).strip() for c in rows[0]).lower()
            # 梯度表应同时包含“流动相A/B”（及可选曲率），避免误把“离心时间/孵育时间”等普通参数表判成梯度表
            if ('流动相a' in header_text and '流动相b' in header_text) and ('时间' in header_text or '曲率' in header_text):
                return True
            if ('流动相a' in flat_text or '流动相A' in flat_text) and ('流动相b' in flat_text or '流动相B' in flat_text) and ('时间' in flat_text):
                return True
            return False

        def is_result_calc_table_strict(rows: List[List[str]], title: str, flat_text: str) -> bool:
            """
            计算章节表格严格过滤：
            - 必须命中计算相关关键词
            - 排除色谱条件、进样序列等操作类表格
            """
            if not rows:
                return False
            header_text = " ".join(str(c).strip() for c in rows[0]).lower()
            full = f"{title} {header_text} {flat_text}".lower()

            include_kws = [
                '计算', '结果', '公式', '鉴定结果',
                '峰', 'rrt', 'rt', '浓度', '含量',
                'slope', 'conc', 'calculation', 'result', 'formula'
            ]
            exclude_kws = [
                '色谱柱', '流动相', '梯度', '平衡系统',
                '进样序列', '进样针数', '进样体积', '备注', 'condition',
                '仪器检查', '方法编辑', '参数设定', 'coupler', 'quick check', 'wavelength'
            ]

            if any(k in full for k in exclude_kws):
                return False
            matched = sum(1 for k in include_kws if k in full)
            # 需要至少2个强计算信号，且至少命中一个“硬特征”关键词
            hard_features = ['公式', '峰', 'rrt', 'rt', '浓度', '含量', 'slope', 'conc', 'formula']
            return matched >= 2 and any(k in full for k in hard_features)

        def normalize_to_acceptance_2cols(rows: List[List[str]]) -> List[List[str]]:
            """将 SOP 可接受标准相关表统一为两列：样品名称/可接受标准。"""
            if not rows:
                return rows
            max_cols = max((len(r) for r in rows), default=0)
            header_text = " ".join(str(c).strip() for c in rows[0]).lower()
            # 已是两列且表头语义正确，直接返回
            if max_cols == 2 and ('可接受标准' in header_text):
                return rows

            # 多列表降维：首列作为名称，其余列拼接为可接受标准描述
            out = [['样品名称', '可接受标准']]
            for r in rows[1:]:
                if not r:
                    continue
                name = str(r[0]).strip() if len(r) > 0 else ''
                if not name:
                    continue
                if name in ('……', '...', '…'):
                    continue
                rest = [str(x).strip() for x in r[1:] if str(x).strip()]
                if not rest:
                    continue
                criteria = "；".join(rest)
                if criteria in ('……', '...', '…'):
                    continue
                out.append([name, criteria])
            if len(out) <= 1:
                return []

            # 去重：按名称保留首条
            dedup = [out[0]]
            seen_names = set()
            for r in out[1:]:
                n = r[0].replace(' ', '')
                if n in seen_names:
                    continue
                seen_names.add(n)
                dedup.append(r)

            # 纯度场景优先保留三条关键可接受标准（暂时停用）
            # preferred = ['空白溶液', 'FB溶液', '系统适用性溶液（控制针）']
            # name_to_row = {r[0]: r for r in dedup[1:]}
            # if all(k in name_to_row for k in preferred):
            #     return [dedup[0]] + [name_to_row[k] for k in preferred]
            return dedup

        results: List[Dict[str, Any]] = []
        excluded_keywords = [
            '签字', '签名', '批准', '审核', '起草', '修订历史', '版本',
            '页码', '职责', 'references', 'record'
        ]
        anchored_suitability_indices, anchored_acceptance_indices = collect_acceptance_anchor_table_indices()
        calc_anchor_info = collect_calculation_anchor_table_info()
        seen = set()

        for t_idx, table in enumerate(self.tables):
            if not table or len(table) < 2:
                continue

            rows = [[str(c).strip() for c in row] for row in table]
            flat_text = " ".join([" ".join(row) for row in rows]).strip()
            if not flat_text:
                continue

            flat_lower = flat_text.lower()
            if any(kw in flat_lower for kw in excluded_keywords):
                continue

            category = ''
            if t_idx in anchored_suitability_indices:
                # 系统适用性可接受标准表格
                rows = normalize_to_acceptance_2cols(rows)
                if not rows:
                    continue
                category = 'suitability_criteria'
            elif t_idx in anchored_acceptance_indices:
                # 产品合格标准表格
                rows = normalize_to_acceptance_2cols(rows)
                if not rows:
                    continue
                category = 'acceptance_criteria'
            elif is_acceptance_related(rows, flat_text):
                rows = normalize_to_acceptance_2cols(rows)
                if not rows:
                    continue
                category = 'acceptance_criteria'
            elif is_gradient_table(rows, flat_text):
                category = 'procedure'
            elif t_idx in calc_anchor_info:
                if not is_result_calc_table_strict(rows, calc_anchor_info.get(t_idx, ''), flat_text):
                    continue
                category = 'result_calculation'
            else:
                continue

            header = rows[0]
            signature = ("|".join(header), len(rows), len(rows[0]) if rows and rows[0] else 0)
            if signature in seen:
                continue
            seen.add(signature)
            results.append({
                'category': category,
                'header': header,
                'rows': rows,
                'table_xml': self.table_xmls[t_idx] if t_idx < len(self.table_xmls) else '',
                'table_title': calc_anchor_info.get(t_idx, ''),
            })

        return results

    def extract_sechplc(self) -> Dict[str, Any]:
        """完整提取 SEC-HPLC SOP 内容

        Returns:
            包含各章节内容的字典
        """
        content = {
            'principle': [],
            'equipment_materials': [],
            'sample_prep': [],
            'procedure': [],
            'calculation': [],
            'acceptance': [],
            'chromatography_conditions': {},
            'acceptance_table': None
        }

        # 提取各章节
        content['principle'] = self.get_section_content(
            ['实验原理', 'Experimental Principles'],
            ['实验材料', 'Experiments Material']
        )

        content['equipment_materials'] = self.get_section_content(
            ['实验材料', 'Experiments Material'],
            ['样品处理', 'Sample Preparation']
        )

        content['sample_prep'] = self.get_section_content(
            ['样品处理', 'Sample Preparation'],
            ['操作步骤', 'Operation Steps']
        )

        content['procedure'] = self.get_section_content(
            ['操作步骤', 'Operation Steps'],
            ['数据处理', 'Data Processing']
        )

        content['calculation'] = self.get_section_content(
            ['计算', 'Calculate'],
            ['系统适用性', 'System Suitability']
        )

        content['acceptance'] = self.get_section_content(
            ['可接受标准', 'Acceptable Standard'],
            ['职责', 'Responsibilities']
        )

        content['chromatography_conditions'] = self.get_chromatography_conditions()
        content['acceptance_table'] = self.get_acceptance_criteria()

        return content

    def extract_solovpe(self) -> Dict[str, Any]:
        """完整提取 SoloVPE SOP 内容

        Returns:
            包含各章节内容的字典
        """
        content = {
            'principle': [],
            'equipment_materials': [],
            'sample_prep': [],
            'procedure': [],
            'calculation': [],
            'acceptance': [],
            'instrument_params': {}
        }

        # 提取各章节
        content['principle'] = self.get_section_content(
            ['实验原理', 'Experimental Principle'],
            ['实验材料', 'Experiments Material']
        )

        content['equipment_materials'] = self.get_section_content(
            ['实验材料', 'Experiments Material'],
            ['样品处理', 'Sample Preparation']
        )

        content['sample_prep'] = self.get_section_content(
            ['样品处理', 'Sample Preparation'],
            ['操作步骤', 'Operations']
        )

        content['procedure'] = self.get_section_content(
            ['操作步骤', 'Operations'],
            ['结果与计算', 'Results and Calculations']
        )

        content['calculation'] = self.get_section_content(
            ['结果与计算', 'Results and Calculations'],
            ['可接受标准', 'Acceptable Standard']
        )

        content['acceptance'] = self.get_section_content(
            ['可接受标准', 'Acceptable Standard'],
            ['注意事项', 'Note', '职责', 'Responsibilities']
        )

        # 提取仪器参数表格
        for table in self.tables:
            if not table:
                continue
            flat = " ".join([" ".join(row) for row in table])
            if '参数' in flat and ('Slope' in flat or 'Wavelength' in flat):
                for row in table:
                    if len(row) >= 2:
                        content['instrument_params'][row[0]] = row[1]

        return content

    def extract_peptide_map(self) -> Dict[str, Any]:
        """完整提取肽图 RP-UPLC SOP 内容

        Returns:
            包含各章节内容的字典
        """
        content = {
            'principle': [],
            'equipment_materials': [],
            'sample_prep': [],
            'procedure': [],
            'calculation': [],
            'acceptance': [],
            'chromatography_conditions': {}
        }

        # 提取各章节 - 使用通用关键词
        content['principle'] = self.get_section_content(
            ['实验原理', 'Experimental Principles', '原理'],
            ['实验材料', 'Experiments Material', '试剂与材料']
        )

        content['equipment_materials'] = self.get_section_content(
            ['实验材料', 'Experiments Material', '试剂与材料', '仪器与设备'],
            ['样品处理', 'Sample Preparation', '溶液制备']
        )

        content['sample_prep'] = self.get_section_content(
            ['样品处理', 'Sample Preparation', '溶液制备'],
            ['操作步骤', 'Operation Steps', '色谱条件', '测定']
        )

        content['procedure'] = self.get_section_content(
            ['操作步骤', 'Operation Steps', '测定', '色谱条件'],
            ['数据处理', 'Data Processing', '结果计算', '计算']
        )

        content['calculation'] = self.get_section_content(
            ['计算', 'Calculate', '结果计算', '数据处理'],
            ['系统适用性', 'System Suitability', '可接受标准']
        )

        content['acceptance'] = self.get_section_content(
            ['可接受标准', 'Acceptable Standard', '合格标准'],
            ['职责', 'Responsibilities', '注意事项']
        )

        # 提取色谱条件表格
        content['chromatography_conditions'] = self.get_chromatography_conditions()

        return content

    def extract_generic(self) -> Dict[str, Any]:
        """通用 SOP 内容提取

        Returns:
            包含各章节内容的字典
        """
        content = {
            'principle': [],
            'materials': [],
            'equipment': [],
            'sample_prep': [],
            'procedure': [],
            'calculation': [],
            'acceptance': [],
            'suitability': []
        }

        # 通用关键词匹配
        content['principle'] = self.get_section_content(
            ['实验原理', '原理', 'Principle', 'Experimental Principle'],
            ['实验材料', '材料', 'Material']
        )

        content['materials'] = self.get_section_content(
            ['实验材料', '材料', '试剂', 'Material', 'Reagent'],
            ['设备', '仪器', 'Equipment', 'Instrument']
        )

        content['equipment'] = self.get_section_content(
            ['设备', '仪器', 'Equipment', 'Instrument'],
            ['样品', '操作', 'Sample', 'Procedure']
        )

        content['sample_prep'] = self.get_section_content(
            ['样品处理', '样品制备', 'Sample Preparation'],
            ['操作步骤', '操作', 'Procedure', 'Operation']
        )

        content['procedure'] = self.get_section_content(
            ['操作步骤', '操作', 'Procedure', 'Operation'],
            ['计算', '数据处理', 'Calculation', 'Data']
        )

        content['calculation'] = self.get_section_content(
            ['计算', '数据处理', '结果', 'Calculation', 'Result'],
            ['可接受', '合格', '标准', 'Acceptance', 'Criteria']
        )

        content['acceptance'] = self.get_section_content(
            ['可接受标准', '合格标准', 'Acceptance', 'Criteria'],
            ['职责', '注意', 'Responsibility', 'Note']
        )

        content['suitability'] = self.get_section_content(
            ['系统适用性', 'System Suitability'],
            ['可接受', '合格', 'Acceptance']
        )

        return content

    # -------------------------------------------------------------------------
    # 第四章「程序」结构化提取 - 用于分析方法文档整合
    # -------------------------------------------------------------------------

    # SOP 第四章 与 分析方法 六章节 映射
    # SOP: 实验原理 | 实验材料及配置 | 样品处理 | 操作步骤 | 数据处理和结果计算 | 可接受标准
    # 方法: 原理 | 材料和设备 | 操作步骤 | 试验成立标准 | 结果计算 | 合格标准

    def _filter_method_scope_lines(self, lines: List[str], method_name: str) -> List[str]:
        """
        方法范围过滤（改进4）：
        - 保留当前方法强相关关键词行
        - 丢弃明显属于其他方法章节的标题行
        """
        if not lines:
            return []
        method = str(method_name or '').strip()
        if not method:
            return lines

        method_kws = {
            '肽图': ['肽图', '胰蛋白酶', 'RRT', 'Peak', '214nm', '洗脱梯度'],
            '蛋白质含量': ['蛋白质含量', '蛋白含量', 'SoloVPE', 'Slope', 'A280', '朗伯-比尔'],
            '纯度（SEC-HPLC）': ['纯度', 'SEC-HPLC', '主峰%', '聚体%', '低分子量杂质%', 'TSKgel', '面积归一法'],
        }
        other_method_titles = ['肽图', '蛋白质含量', '蛋白含量', '纯度（SEC-HPLC）', '还原纯度（CE-SDS）', '非还原纯度（CE-SDS）']
        target_kws = method_kws.get(method, [method])

        out: List[str] = []
        for line in lines:
            s = str(line).strip()
            if not s:
                continue
            # 标题级串段过滤：若该行本身就是其他方法标题，则跳过
            if s in other_method_titles and s != method:
                continue
            # 方法无关短标题过滤（避免串到其它方法）
            if len(s) <= 24 and any(x in s for x in other_method_titles if x != method):
                if not any(k in s for k in target_kws):
                    continue
            out.append(s)
        return out

    def extract_procedure_chapter4(self, method_name: str = '') -> Dict[str, Any]:
        """
        从 SOP 第四章「程序」提取结构化内容，输出与分析方法六章节对应的字段。

        Returns:
            dict: {
                'principle': [],           # 1. 原理
                'materials_and_equipment': [],  # 2. 材料和设备（实验材料+试液配制）
                'sample_prep': [],          # 3. 操作步骤-样品处理
                'procedure': [],           # 3. 操作步骤-测定法
                'suitability_criteria': [], # 4. 试验成立标准（系统适用性可接受标准）
                'result_calculation': [],  # 5. 结果计算
                'acceptance_criteria': [], # 6. 合格标准（产品判定）
                'chromatography_conditions': {},  # 色谱条件表（如有）
            }
        """
        content = {
            'principle': [],
            'materials_and_equipment': [],
            'sample_prep': [],
            'procedure': [],
            'suitability_criteria': [],
            'result_calculation': [],
            'acceptance_criteria': [],
            'chromatography_conditions': {},
            'section_images': {
                'suitability_criteria': [],
                'typical_figure': [],
                'result_calculation': [],
            },
        }

        # 1. 实验原理 → 原理
        content['principle'] = self.get_section_content(
            ['实验原理', 'Experimental Principle', 'Experimental Principles', '原理', '测定原理'],
            ['实验材料', 'Experiments Material', '试剂与材料', '仪器与设备', '仪器及设备', '实验设备', '试剂及耗材', '实验样品', '操作步骤', '仪器设备Equipment', '设备', 'Equipment']
        )

        # 2. 实验材料及配置 → 材料和设备（合并实验材料 + 试液配制 + 设备表格）
        # 注意：end 仅用 样品处理，避免 溶液制备/溶液配制 等导致提前终止，丢失 实验设备 内容
        mat = self.get_section_content(
            ['实验材料', 'Experiments Material', '试剂与材料', '仪器与设备', '实验设备', '试验设备', '关键耗材'],
            ['样品处理', 'Sample Preparation']
        )
        sol = self.get_section_content(
            ['试液配制', '溶液制备', '溶液配制', 'Solution Preparation'],
            ['样品处理', 'Sample Preparation', '操作步骤', 'Operation Steps', 'Operations']
        )
        # 补充设备/材料表格内容（表格中的设备、仪器列表易被段落提取遗漏）
        equipment_table_lines = self.get_equipment_materials_tables()
        # 设备表格内容放前，确保试验设备不丢失
        content['materials_and_equipment'] = self._normalize_materials_equipment_lines(
            equipment_table_lines + mat + sol
        )

        # 3a. 样品处理 → 操作步骤(样品处理)
        content['sample_prep'] = self.get_section_content(
            ['样品处理', 'Sample Preparation', '供试品制备', '供试品溶液的制备'],
            ['操作步骤', 'Operation Steps', 'Operations', '色谱条件', '测定']
        )

        # 3b. 操作步骤 → 操作步骤(测定法)
        content['procedure'] = self.get_section_content(
            ['操作步骤', 'Operation Steps', 'Operations', '色谱条件', '平衡系统', '进样序列', '创建序列'],
            ['数据处理', 'Data Processing', '结果与计算', 'Results and Calculations', '结果计算', '计算', 'Calculate']
        )

        # 4. 系统适用性试验下的可接受标准 → 试验成立标准
        content['suitability_criteria'] = self.get_section_content(
            ['系统适用性', 'System Suitability', '系统适用性试验'],
            ['可接受标准 Acceptable Standard', 'Acceptable Standard', '色谱柱', 'Column and System', 'Column and System Cleaning', '注意事项', 'Note', '职责', 'Responsibilities']
        )
        content['section_images']['suitability_criteria'] = self.extract_section_images_with_captions(
            ['系统适用性', 'System Suitability', '系统适用性试验'],
            ['可接受标准 Acceptable Standard', 'Acceptable Standard', '色谱柱', 'Column and System', 'Column and System Cleaning', '注意事项', 'Note', '职责', 'Responsibilities']
        )
        # 典型图谱（须与模版「典型图谱」小节对齐；勿用泛词「谱图」——正文中「色谱图」等会误开节，
        # 且短标题「数据处理」会作为 end 提前 break，导致永远扫不到图在标题上方的插图。）
        content['section_images']['typical_figure'] = self.extract_section_images_with_captions(
            ['典型图谱', '典型色谱', '附图', 'Figure', 'FIGURE'],
            [
                '可接受标准 Acceptable Standard',
                'Acceptable Standard',
                '职责',
                '注意事项',
                '附录',
                '色谱柱',
                'Column and System',
            ],
        )
        # 数据处理/结果计算段内的图
        content['section_images']['result_calculation'] = self.extract_section_images_with_captions(
            ['数据处理', 'Data Processing', '结果与计算', 'Results and Calculations', '结果计算'],
            ['可接受标准 Acceptable Standard', 'Acceptable Standard', '系统适用性', '职责', '注意事项', '附录', '色谱柱']
        )
        # 过滤：若可接受标准在系统适用性试验下，通常是试验成立标准；若含 RSD/线性/分离度等则为试验成立
        filtered_suitability = []
        for line in content['suitability_criteria']:
            if any(kw in line for kw in ['RSD', '线性', '分离度', '相关系数', 'R2', '保留时间', '理论塔板', '拖尾', '≤', '≥']):
                filtered_suitability.append(line)
            elif '可接受标准' not in line and 'Acceptable' not in line and len(line) > 10:
                filtered_suitability.append(line)
        content['suitability_criteria'] = filtered_suitability or content['suitability_criteria']

        # 5. 数据处理 + 计算 → 结果计算
        data_proc = self.get_section_content(
            ['数据处理', 'Data Processing'],
            ['计算', 'Calculate', '系统适用性', 'System Suitability']
        )
        calc = self.get_section_content(
            ['计算', 'Calculate', '结果计算', '浓度计算', '数据计算'],
            ['系统适用性', 'System Suitability', '可接受标准', 'Acceptable Standard', '结果判定']
        )
        content['result_calculation'] = data_proc + calc

        # 6. 可接受标准 Acceptable Standard（产品合格判定）→ 合格标准
        content['acceptance_criteria'] = self.get_section_content(
            ['可接受标准 Acceptable Standard', '可接受标准Acceptable Standard'],
            ['注意事项', 'Note', '职责', 'Responsibilities', '色谱柱', 'Column and System']
        )
        # 排除系统适用性相关内容，只保留产品判定
        product_acceptance = []
        for line in content['acceptance_criteria']:
            if any(kw in line for kw in ['按各样品', '参照', '质量标准', '一致', 'mg/mL', 'ppm', '不低于', '不高于', '%']):
                product_acceptance.append(line)
            elif '可接受标准' not in line and 'Acceptable' not in line and len(line) < 80:
                product_acceptance.append(line)
        content['acceptance_criteria'] = product_acceptance if product_acceptance else content['acceptance_criteria']

        # 色谱条件表
        content['chromatography_conditions'] = self.get_chromatography_conditions()

        # 改进4：按方法作用域收敛，避免跨方法串段
        for key in [
            'principle', 'materials_and_equipment', 'sample_prep',
            'procedure', 'suitability_criteria', 'result_calculation',
            'acceptance_criteria'
        ]:
            content[key] = self._filter_method_scope_lines(content.get(key, []), method_name)

        return content
