# -*- coding: utf-8 -*-
"""
SOP 整合到分析方法标准文件

从 SOP 第四章「程序」提取内容，与分析方法标准模板比对整合，
输出含有内部 SOP 内容的分析方法文档。

三种模式：
  1. 默认模式 — 直接整合（SOP 原始内容写入模板）
  2. --extract-only — 仅提取，输出 JSON（再运行 refine_extracted.py）
  3. --from-json — 从 refined.json 写入 docx

映射关系：
  SOP 第四章程序                 →  分析方法六章节
  ─────────────────────────────────────────────────
  1. 实验原理                    →  1. 原理
  2. 实验材料及配置              →  2. 材料和设备
  3. 样品处理                    →  3. 操作步骤（样品处理）
  4. 操作步骤                    →  3. 操作步骤（测定法）
  5. 系统适用性可接受标准        →  4. 试验成立标准
  6. 数据处理和结果计算          →  5. 结果计算
  7. 可接受标准（产品判定）      →  6. 合格标准

使用方法:
    # 直接整合（原有行为）
    python integrate_sop_method.py -t 模板.docx -c methods.json

    # 仅提取，输出 JSON（再运行 refine_extracted.py）
    python integrate_sop_method.py -t 模板.docx -c methods.json ^
        --extract-only -r 参考模板.docx -o extracted.json

    # 从 refined.json 写入 docx
    python integrate_sop_method.py -t 模板.docx --from-json refined.json

    # 仅提供多个 SOP + 模版（自动推断章节、新项增章）
    python integrate_sop_method.py -t 模板.docx --merge-sops sop1.docx sop2.docx
"""
import os
import sys
import json
import argparse
import glob
import shutil
from datetime import datetime
from docx import Document
import re
from typing import Dict, List, Any, Optional

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)
sys.path.insert(0, os.path.join(SCRIPT_DIR, 'regulatory_core'))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)

try:
    from ra_compliance import print_skill_iron_rules
except ImportError:
    def print_skill_iron_rules(stream=None):
        pass


def _disp(text: str) -> str:
    """控制台友好显示：部分 SOP 文件名含 U+00A0，GBK 控制台打印会报错。"""
    return (text or '').replace('\u00a0', ' ')


def _default_analysis_template_path() -> Optional[str]:
    """
    未指定 -t 时，尝试使用项目 input/ 下当前标准模版：
    优先 32s42*.docx，其次文件名含「分析方法」且含「模板/模版」。
    """
    inp = os.path.join(PROJECT_ROOT, 'input')
    if not os.path.isdir(inp):
        return None
    candidates: List[str] = []
    for pat in (
        os.path.join(inp, '32s42*.docx'),
        os.path.join(inp, '*分析方法*模板*.docx'),
        os.path.join(inp, '*分析方法*模版*.docx'),
    ):
        candidates.extend(glob.glob(pat))
    seen = set()
    uniq: List[str] = []
    for p in sorted(candidates):
        if os.path.basename(p).startswith('~$'):
            continue
        key = os.path.normcase(os.path.abspath(p))
        if key in seen:
            continue
        seen.add(key)
        uniq.append(p)
    if not uniq:
        return None
    for p in uniq:
        bn = os.path.basename(p)
        if '模板' in bn or '模版' in bn:
            return p
    return uniq[0]

from regulatory_core.analysis_integrator import (
    integrate_sop_into_template,
    extract_for_refinement,
    write_refined_to_template,
    _extract_section_text_from_doc,
)
from regulatory_core.method_section_order import (
    OrderSpec,
    parse_method_order_file,
    reorder_and_insert_methods,
    refresh_toc,
    merge_order_spec_into_items,
    order_refined_methods_by_spec,
    build_auto_merge_spec_and_items,
    titles_match,
)


def load_config(path: str) -> list:
    """加载 JSON 配置文件（支持数组，或含 methods 数组的对象）。"""
    with open(path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    if isinstance(data, dict) and isinstance(data.get('methods'), list):
        return list(data['methods'])
    if isinstance(data, list):
        return data
    return [data]


def _split_nonempty_lines(text: str) -> List[str]:
    """按行拆分并去除空行。"""
    return [x.strip() for x in str(text).splitlines() if str(x).strip()]


def _build_source_pool(method_data: Dict[str, Any]) -> List[str]:
    """
    从 method 数据中提取可追溯来源文本池（sop_raw/template_existing/reference_style）。
    """
    pool: List[str] = []
    if not isinstance(method_data, dict):
        return pool

    for key in ['sop_raw', 'template_existing', 'reference_style']:
        val = method_data.get(key)
        if isinstance(val, dict):
            for v in val.values():
                if isinstance(v, list):
                    for item in v:
                        s = str(item).strip()
                        if s:
                            pool.append(s)
                elif isinstance(v, str):
                    pool.extend(_split_nonempty_lines(v))
                elif v is not None:
                    s = str(v).strip()
                    if s:
                        pool.append(s)
        elif isinstance(val, list):
            for item in val:
                s = str(item).strip()
                if s:
                    pool.append(s)
        elif isinstance(val, str):
            pool.extend(_split_nonempty_lines(val))
    return pool


def _load_extracted_context(
    from_json_path: str,
    extracted_full_path: Optional[str] = None
) -> Dict[str, Dict[str, Any]]:
    """
    自动加载同目录 extracted*.full.json / extracted.json，构建 name -> method_data 的上下文映射。
    """
    ctx_map: Dict[str, Dict[str, Any]] = {}
    base_dir = os.path.dirname(os.path.abspath(from_json_path))
    # 优先读取与 refined 同批次的 *.full.json（由 refine_extracted 写入 extracted_full_path）
    candidate_paths: List[str] = []
    if extracted_full_path and os.path.isfile(extracted_full_path):
        candidate_paths.append(os.path.abspath(extracted_full_path))
    stem, _ = os.path.splitext(os.path.abspath(from_json_path))
    if 'refined' in os.path.basename(from_json_path):
        guess = stem.replace('refined', 'extracted', 1) + '.full.json'
        candidate_paths.append(guess)
    candidate_paths.extend([
        os.path.join(base_dir, 'extracted.full.json'),
        os.path.join(base_dir, 'extracted.json'),
    ])
    seen = set()
    for extracted_path in candidate_paths:
        if not extracted_path or extracted_path in seen:
            continue
        seen.add(extracted_path)
        if not os.path.exists(extracted_path):
            continue
        try:
            with open(extracted_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            for m in data.get('methods', []):
                name = m.get('name', '')
                if name:
                    ctx_map[name] = m
            if ctx_map:
                return ctx_map
        except Exception:
            continue
    return ctx_map


def _normalize_method_name(name: str) -> str:
    """归一化方法名用于弱匹配。"""
    t = str(name or '').strip().lower()
    t = re.sub(r'（[^）]*）|\([^)]*\)', '', t)
    t = re.sub(r'[\s\u3000]', '', t)
    return t


def _find_method_context(section_name: str, ctx_map: Dict[str, Dict[str, Any]]) -> Dict[str, Any]:
    """在 extracted 上下文中按方法名做精确+弱匹配。"""
    if section_name in ctx_map:
        return ctx_map[section_name]

    target = _normalize_method_name(section_name)
    for k, v in ctx_map.items():
        nk = _normalize_method_name(k)
        if nk == target or (target and nk and (target in nk or nk in target)):
            return v

    # 常见同义：蛋白含量 <-> 蛋白质含量
    alt = target.replace('蛋白质含量', '蛋白含量').replace('蛋白含量', '蛋白质含量')
    for k, v in ctx_map.items():
        nk = _normalize_method_name(k)
        if nk == alt or (alt and nk and (alt in nk or nk in alt)):
            return v

    return {}


def _empty_sop_raw() -> Dict[str, List[str]]:
    return {
        'principle': [],
        'materials_and_equipment': [],
        'sample_prep': [],
        'procedure': [],
        'suitability_criteria': [],
        'result_calculation': [],
        'acceptance_criteria': [],
    }


def _stub_extracted_method(
    section_name: str,
    stop_keywords: List[str],
    template_path: str,
    ref_template_path: Optional[str],
) -> Dict[str, Any]:
    """无 SOP 路径时占位，便于 refined 流水线保留章节与 stop。"""
    out: Dict[str, Any] = {
        'name': section_name,
        'stop': stop_keywords,
        'sop_path': '',
        'sop_raw': _empty_sop_raw(),
        'sop_tables': [],
        'sop_section_images': {},
    }
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
        out['template_existing'] = _extract_section_text_from_doc(
            doc, section_name, stop_keywords
        )
    if ref_template_path and os.path.exists(ref_template_path):
        ref_doc = Document(ref_template_path)
        out['reference_style'] = _extract_section_text_from_doc(
            ref_doc, section_name, stop_keywords
        )
    return out


def _resolve_sop_path(sop_path: str, template_dir: str) -> str:
    """解析 SOP 路径：相对路径基于模板目录；文件名 U+00A0 与普通空格互换再尝试。"""
    if not os.path.isabs(sop_path):
        p = os.path.join(template_dir, sop_path)
    else:
        p = sop_path
    if os.path.isfile(p):
        return p
    # 双向尝试：U+00A0 <-> 普通空格
    relaxed = p.replace('\u00a0', ' ')
    if relaxed != p and os.path.isfile(relaxed):
        return relaxed
    # 反向尝试：普通空格 -> U+00A0
    unrelaxed = p.replace(' ', '\u00a0')
    if unrelaxed != p and os.path.isfile(unrelaxed):
        return unrelaxed
    return p


def _default_output_dir() -> str:
    """统一默认输出目录为项目根目录下 output/。"""
    out_dir = os.path.join(PROJECT_ROOT, 'output')
    os.makedirs(out_dir, exist_ok=True)
    return out_dir


def _resolve_json_output_path(output_arg: str, default_name: str = 'extracted.json') -> str:
    """解析 JSON 输出路径；若传入目录则自动落为默认文件名。"""
    if not output_arg:
        return os.path.join(_default_output_dir(), default_name)

    out = os.path.abspath(output_arg)
    is_dir_arg = output_arg.endswith(os.sep) or output_arg.endswith('/') or os.path.isdir(out)
    if is_dir_arg:
        os.makedirs(out, exist_ok=True)
        return os.path.join(out, default_name)

    parent = os.path.dirname(out)
    if parent:
        os.makedirs(parent, exist_ok=True)
    return out


def _resolve_docx_output_path(output_arg: str, template_path: str, ts: str, refined: bool = False) -> str:
    """
    解析 docx 输出路径：
    - 未传 output：默认 output/<模板名>[_refined]_时间戳.docx
    - output 为目录：自动补模板名前缀，避免生成 "_时间戳.docx"
    - output 为文件：在文件名末尾追加时间戳
    """
    base_name = os.path.splitext(os.path.basename(template_path))[0]
    stem = f"{base_name}_refined" if refined else base_name

    if not output_arg:
        return os.path.join(_default_output_dir(), f"{stem}_{ts}.docx")

    out_abs = os.path.abspath(output_arg)
    is_dir_arg = output_arg.endswith(os.sep) or output_arg.endswith('/') or os.path.isdir(out_abs)
    if is_dir_arg:
        os.makedirs(out_abs, exist_ok=True)
        return os.path.join(out_abs, f"{stem}_{ts}.docx")

    base, ext = os.path.splitext(out_abs)
    if not ext:
        ext = '.docx'
    if not os.path.basename(base):
        # 兼容异常输入，回退到默认目录命名
        return os.path.join(_default_output_dir(), f"{stem}_{ts}.docx")

    parent = os.path.dirname(base)
    if parent:
        os.makedirs(parent, exist_ok=True)
    return f"{base}_{ts}{ext}"


def _load_method_order_spec(args):
    """若指定 --method-order 且文件存在，返回 OrderSpec，否则 None。"""
    path = getattr(args, 'method_order', None) or ''
    path = str(path).strip()
    if not path:
        return None
    if not os.path.isfile(path):
        print(f'  [WARN] --method-order 文件不存在，已忽略: {_disp(path)}')
        return None
    try:
        return parse_method_order_file(path)
    except Exception as e:
        print(f'  [WARN] 解析 method-order 失败，已忽略: {_disp(str(e))}')
        return None


def _embedded_order_from_config(args) -> Optional[OrderSpec]:
    """从 -c 配置文件读取 order / method_order / new_sections（无 --method-order 时生效）。"""
    cfg = getattr(args, 'config', None) or ''
    cfg = str(cfg).strip()
    if not cfg or not os.path.isfile(cfg):
        return None
    try:
        with open(cfg, 'r', encoding='utf-8') as f:
            raw = json.load(f)
        if not isinstance(raw, dict):
            return None
        order = raw.get('order') or raw.get('method_order')
        ns = raw.get('new_sections')
        if not order and not ns:
            return None
        ol = [str(x).strip() for x in (order or []) if str(x).strip()]
        nd: Dict[str, Any] = {}
        if isinstance(ns, dict):
            nd = {str(k): dict(v) if isinstance(v, dict) else {} for k, v in ns.items()}
        return OrderSpec(order=ol, new_sections=nd)
    except Exception:
        return None


def _effective_order_spec(args) -> Optional[OrderSpec]:
    """
    优先 --method-order 文件，其次 --merge-sops/--infer-section 生成的内置顺序，否则 -c 内嵌。

    重要：当同时提供 --method-order 与 --merge-sops/--infer-section 时，
    需要把自动识别到的「模板不存在的新章节（new_sections）」合并进来，
    否则写回阶段无法插入骨架，会出现“章节不存在需手动添加”的假问题。
    """

    def _append_missing_titles(dst: List[str], src: List[str]) -> List[str]:
        out = list(dst or [])
        for t in src or []:
            tt = str(t).strip()
            if not tt:
                continue
            if all(not titles_match(tt, exist) for exist in out):
                out.append(tt)
        return out

    def _merge_order_specs(primary: OrderSpec, secondary: OrderSpec) -> OrderSpec:
        # primary 优先（通常来自 --method-order），secondary 补齐（通常来自 --merge-sops 自动生成）
        order = _append_missing_titles(primary.order or [], secondary.order or [])
        new_sections: Dict[str, Any] = dict(primary.new_sections or {})
        for k, v in (secondary.new_sections or {}).items():
            kk = str(k).strip()
            if not kk:
                continue
            if all(not titles_match(kk, exist) for exist in new_sections.keys()):
                new_sections[kk] = dict(v) if isinstance(v, dict) else {}
        return OrderSpec(order=order, new_sections=new_sections)

    method_spec = _load_method_order_spec(args)
    intr = getattr(args, 'intrinsic_auto_order_spec', None)

    if method_spec and (method_spec.order or method_spec.new_sections):
        if intr is not None and (intr.order or intr.new_sections):
            return _merge_order_specs(method_spec, intr)
        return method_spec

    # 若用户提供了 --method-order 但解析不到任何条目，则视为“无有效顺序文件”，
    # 允许回退到 intrinsic（--merge-sops/--infer-section 自动生成）或 config 内嵌顺序。
    # （不要在这里 return None 去阻断后续分支）

    if intr is not None and (intr.order or intr.new_sections):
        return intr

    return _embedded_order_from_config(args)


def _merge_auto_items_with_config(config_items: List[dict], auto_items: List[dict]) -> List[dict]:
    """config 优先；auto 中 config 未覆盖的 section（弱匹配）则追加。"""
    out: List[dict] = [dict(x) for x in config_items]
    for a in auto_items:
        sec_a = str(a.get('section', '') or '').strip()
        if not sec_a:
            continue
        found = False
        for o in out:
            if titles_match(sec_a, str(o.get('section', '') or '').strip()):
                found = True
                break
        if not found:
            out.append(dict(a))
    return out


def _parse_items(args) -> list:
    """从参数中解析方法配置列表"""
    if args.config:
        config_path = args.config
        if not os.path.isabs(config_path):
            config_path = os.path.join(os.getcwd(), config_path)
        return load_config(config_path)
    elif args.sop and args.section:
        return [{
            'sop': args.sop,
            'section': args.section,
            'stop': args.stop or []
        }]
    return []


def _safe_run_id(run_id: str) -> str:
    """将 run_id 规整为可用作文件名的安全字符串。"""
    return re.sub(r'[^A-Za-z0-9_.-]', '_', run_id)


def _run_manifest_path(template_path: str, run_id: str) -> str:
    """返回 run_id 对应的清单文件路径。"""
    template_dir = os.path.dirname(os.path.abspath(template_path))
    safe_id = _safe_run_id(run_id)
    return os.path.join(template_dir, f'.ra_refined_run_{safe_id}.json')


def _record_run_output(template_path: str, run_id: str, output_path: str) -> None:
    """记录本次 run 生成的 refined 输出文件。"""
    manifest_path = _run_manifest_path(template_path, run_id)
    data = {'outputs': []}

    if os.path.exists(manifest_path):
        try:
            with open(manifest_path, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
            if isinstance(loaded, dict):
                data = loaded
        except Exception:
            data = {'outputs': []}

    outputs = data.get('outputs', [])
    if not isinstance(outputs, list):
        outputs = []

    abs_output = os.path.abspath(output_path)
    if abs_output not in outputs:
        outputs.append(abs_output)
    data['outputs'] = outputs

    with open(manifest_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _cleanup_run_outputs(template_path: str, run_id: str, final_output_path: str) -> None:
    """
    清理同一 run_id 的 refined 中间文件，仅保留 final_output_path。
    仅处理清单中记录的文件，不会影响历史 run。
    """
    manifest_path = _run_manifest_path(template_path, run_id)
    if not os.path.exists(manifest_path):
        print(f'  [WARN] 未找到 run 清单，跳过清理: {manifest_path}')
        return

    try:
        with open(manifest_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        print(f'  [WARN] 读取 run 清单失败，跳过清理: {e}')
        return

    outputs = data.get('outputs', [])
    if not isinstance(outputs, list):
        outputs = []

    final_abs = os.path.abspath(final_output_path)
    deleted = 0
    skipped = 0

    for path in outputs:
        try:
            abs_path = os.path.abspath(path)
            if abs_path == final_abs:
                skipped += 1
                continue
            if os.path.exists(abs_path):
                os.remove(abs_path)
                deleted += 1
        except Exception as e:
            print(f'  [WARN] 删除失败: {path} ({e})')

    # 清理完成后删除清单，避免影响后续独立 run
    try:
        os.remove(manifest_path)
    except Exception:
        pass

    print(f'清理完成: 删除 {deleted} 个中间文件，保留最终文件 1 个（run_id={run_id}）')


# ─────────────────────────────────────────────────────────────
# 模式 1：直接整合（原有行为）
# ─────────────────────────────────────────────────────────────

def run_direct_integrate(args, items):
    """直接将 SOP 内容整合到模板"""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = _resolve_docx_output_path(args.output, args.template, ts, refined=False)

    template_dir = os.path.dirname(args.template)
    order_spec = _effective_order_spec(args)
    if order_spec:
        items = merge_order_spec_into_items(items, order_spec, args.template)
        print(f'  [INFO] 已按章节顺序规则合并配置，共 {len(items)} 个方法')

    print('=' * 60)
    print('SOP → 分析方法 直接整合')
    print('=' * 60)
    print(f'模板: {args.template}')
    print(f'输出: {output_path}')
    print(f'格式: {args.style}')
    print()

    template_for_first = args.template
    if order_spec:
        shutil.copy(args.template, output_path)
        _rdoc = Document(output_path)
        reorder_and_insert_methods(_rdoc, order_spec)
        _rdoc.save(output_path)
        template_for_first = output_path
        print('  [INFO] 已重排/插入方法块（顺序文件或 --merge-sops 自动生成）')

    first_integrate_done = False
    for item in items:
        sop_path = _resolve_sop_path(item.get('sop', ''), template_dir)
        section = item.get('section', '')
        stop = item.get('stop', [])
        if isinstance(stop, str):
            stop = [stop]
        if not str(item.get('sop', '') or '').strip() or not os.path.isfile(sop_path):
            print(f'  [SKIP] {_disp(section)}: 无有效 SOP 路径，保留模版骨架')
            continue
        try:
            tpl = output_path if first_integrate_done else template_for_first
            integrate_sop_into_template(
                template_path=tpl,
                sop_path=sop_path,
                section_name=section,
                stop_keywords=stop,
                output_path=output_path,
                style_mode=args.style
            )
            first_integrate_done = True
            print(f'  [OK] {_disp(section)} <- {_disp(os.path.basename(sop_path))}')
        except Exception as e:
            print(f'  [X] {_disp(section)}: {_disp(str(e))}')

    if order_spec:
        try:
            _tdoc = Document(output_path)
            if refresh_toc(_tdoc):
                _tdoc.save(output_path)
                print('  [INFO] 已刷新目录 TOC 域（Word 中请 F9 更新页码）')
        except Exception as e:
            print(f'  [WARN] refresh_toc: {_disp(str(e))}')

    print(f'\n输出: {output_path}')
    try:
        os.startfile(output_path)
    except Exception:
        pass


# ─────────────────────────────────────────────────────────────
# 模式 2：仅提取，输出 JSON
# ─────────────────────────────────────────────────────────────

def run_extract_only(args, items):
    """提取 SOP + 参考模板内容，输出 JSON"""
    template_dir = os.path.dirname(args.template)
    ref_path = args.ref if args.ref else None

    output_path = _resolve_json_output_path(args.output, default_name='extracted.json')

    print('=' * 60)
    print('SOP → JSON 提取（下一步: python refine_extracted.py）')
    print('=' * 60)
    print(f'模板: {args.template}')
    if ref_path:
        print(f'参考: {ref_path}')
    print(f'输出: {output_path}')
    print()

    order_spec = _effective_order_spec(args)
    if order_spec:
        items = merge_order_spec_into_items(items, order_spec, args.template)
        print(f'  [INFO] 已按章节顺序规则合并配置，共 {len(items)} 个方法')

    methods = []
    for item in items:
        sop_path = _resolve_sop_path(item.get('sop', ''), template_dir)
        section = item.get('section', '')
        stop = item.get('stop', [])
        if isinstance(stop, str):
            stop = [stop]
        try:
            if not str(item.get('sop', '') or '').strip() or not os.path.isfile(sop_path):
                methods.append(
                    _stub_extracted_method(section, stop, args.template, ref_path)
                )
                print(f'  [OK] {_disp(section)} (无 SOP，已写入占位提取项)')
                continue
            data = extract_for_refinement(
                template_path=args.template,
                sop_path=sop_path,
                section_name=section,
                stop_keywords=stop,
                ref_template_path=ref_path
            )
            methods.append(data)
            print(f'  [OK] {_disp(section)} <- {_disp(os.path.basename(sop_path))}')
        except Exception as e:
            print(f'  [X] {_disp(section)}: {_disp(str(e))}')

    result = {
        'template': os.path.basename(args.template),
        'template_path': os.path.abspath(args.template),
        'style': args.style,
        'methods': methods,
    }

    # 为 LLM 生成轻量版本：默认 extracted.json 移除 table_xml，降低上下文体积；
    # 同目录保留 extracted.full.json 供写回阶段读取样式信息。
    slim_result = json.loads(json.dumps(result))
    for method in slim_result.get('methods', []):
        for table in method.get('sop_tables', []) or []:
            if isinstance(table, dict) and 'table_xml' in table:
                table.pop('table_xml', None)

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(slim_result, f, ensure_ascii=False, indent=2)

    slim_abs = os.path.abspath(output_path)
    base, _ext = os.path.splitext(slim_abs)
    full_output_path = base + '.full.json'
    with open(full_output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f'\n已输出: {output_path}')
    print(f'已输出(保留表格样式): {full_output_path}')
    print('下一步: python refine_extracted.py 生成 refined.json')


# ─────────────────────────────────────────────────────────────
# 模式 3：从精简 JSON 写入 docx
# ─────────────────────────────────────────────────────────────

def run_from_json(args):
    """从精简后的 JSON 写入 docx"""
    json_path = args.from_json
    if not os.path.exists(json_path):
        print(f'错误: JSON 文件不存在: {json_path}')
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    style_mode = data.get('style', args.style)
    methods = data.get('methods', [])
    efp = data.get('extracted_full_path') or ''
    extracted_ctx_map = _load_extracted_context(json_path, efp if efp else None)

    # 确定模板路径
    template_path = args.template
    if not template_path and data.get('template_path'):
        template_path = data['template_path']
    if not template_path:
        print('错误: 请通过 -t 指定模板路径或在 JSON 中包含 template_path')
        sys.exit(1)

    # 写入模板必须使用用户指定的 -t（或 JSON 内 template_path）
    # 基准文档仅可作为参考，不可替代写入模板。
    if args.write_template:
        print('  [WARN] --write-template 已弃用且将被忽略；写入模板固定为 -t/JSON 的 template_path')

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = _resolve_docx_output_path(args.output, template_path, ts, refined=True)

    order_spec = _effective_order_spec(args)
    if order_spec:
        shutil.copy(template_path, output_path)
        _odoc = Document(output_path)
        reorder_and_insert_methods(_odoc, order_spec)
        _odoc.save(output_path)
        methods = order_refined_methods_by_spec(methods, order_spec, output_path)
        print(f'  [INFO] 已重排模版块并调整 JSON 方法顺序（{len(methods)} 个）')

    print('=' * 60)
    print('JSON → 分析方法 docx 写入')
    print('=' * 60)
    print(f'模板: {template_path}')
    print(f'JSON: {json_path}')
    print(f'输出: {output_path}')
    print(f'格式: {style_mode}')
    print()

    template_for_first = output_path if order_spec else template_path

    first_write_done = False
    for method in methods:
        section = method.get('name', '')
        stop = method.get('stop', [])
        if isinstance(stop, str):
            stop = [stop]
        refined = method.get('refined', {})
        if not refined:
            print(f'  [SKIP] {_disp(section)}: 无 refined 内容')
            continue
        try:
            method_ctx = _find_method_context(section, extracted_ctx_map)
            source_pool = _build_source_pool(method) + _build_source_pool(method_ctx)
            replacement_tables = method_ctx.get('sop_tables') or method.get('sop_tables') or []
            required_source = method_ctx.get('sop_raw') or {}
            replacement_section_images = method_ctx.get('sop_section_images') or method.get('sop_section_images') or {}
            source_sop_path = method_ctx.get('sop_path') or method.get('sop_path')
            tpl_in = output_path if first_write_done else template_for_first
            write_refined_to_template(
                template_path=tpl_in,
                section_name=section,
                stop_keywords=stop,
                refined_content=refined,
                output_path=output_path,
                style_mode=style_mode,
                source_pool=source_pool,
                replacement_tables=replacement_tables,
                required_source=required_source,
                replacement_section_images=replacement_section_images,
                sop_path=source_sop_path
            )
            first_write_done = True
            print(f'  [OK] {_disp(section)}')
        except Exception as e:
            print(f'  [X] {_disp(section)}: {_disp(str(e))}')

    if order_spec:
        try:
            _tdoc = Document(output_path)
            if refresh_toc(_tdoc):
                _tdoc.save(output_path)
                print('  [INFO] 已刷新目录 TOC 域（Word 中请 F9 更新页码）')
        except Exception as e:
            print(f'  [WARN] refresh_toc: {_disp(str(e))}')

    print(f'\n输出: {output_path}')

    if args.run_id:
        _record_run_output(template_path, args.run_id, output_path)
        print(f'已记录 run 输出: run_id={args.run_id}')

    if args.finalize:
        if not args.run_id:
            print('  [WARN] --finalize 需要配合 --run-id 使用，已跳过清理')
        else:
            _cleanup_run_outputs(template_path, args.run_id, output_path)

    try:
        os.startfile(output_path)
    except Exception:
        pass


def main():
    parser = argparse.ArgumentParser(
        description='SOP 第四章程序 → 分析方法六章节 整合',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
三种模式：

  1. 直接整合（默认）：
     python integrate_sop_method.py -t 模板.docx -c methods.json

  1b. 多 SOP + 模版，自动生成章节/增章：
     python integrate_sop_method.py -t 模板.docx --merge-sops "路径1.docx" "路径2.docx"

  2. 仅提取 JSON（再运行 refine_extracted.py）：
     python integrate_sop_method.py -t 模板.docx -c methods.json ^
         --extract-only -r 参考模板.docx -o extracted.json

  3. 从 refined.json 写入 docx：
     python integrate_sop_method.py -t 模板.docx --from-json refined.json

  4. 同一流程内清理 refined 中间文件（保留最终文件）：
     python integrate_sop_method.py -t 模板.docx --from-json refined.json --run-id R20260302A
     python integrate_sop_method.py -t 模板.docx --from-json refined.json --run-id R20260302A --finalize

配置示例 (methods.json):
  [
    {"sop": "纯度SOP.docx", "section": "纯度（SEC-HPLC）", "stop": ["还原纯度（CE-SDS法）"]},
    {"sop": "蛋白含量SOP.docx", "section": "蛋白质含量", "stop": ["纯度（SEC-HPLC）"]},
    {"sop": "肽图SOP.docx", "section": "肽图", "stop": ["蛋白质含量"]}
  ]
        """
    )
    parser.add_argument('-t', '--template', help='分析方法标准模板路径（可省略：自动使用 input/ 下 32s42-分析方法-模板文件.docx）')
    parser.add_argument('-s', '--sop', help='SOP 文件路径（单方法时）')
    parser.add_argument('--section', help='模板中方法章节名称')
    parser.add_argument('--stop', nargs='+', help='下一章节关键词')
    parser.add_argument('-c', '--config', help='JSON 配置文件（多方法）')
    parser.add_argument(
        '--method-order',
        metavar='PATH',
        default=None,
        help='方法章节顺序：.docx（编号列表）或 .json（order / new_sections）；优先于 -c 内嵌 order',
    )
    parser.add_argument(
        '--merge-sops',
        nargs='+',
        metavar='PATH',
        help='多个 SOP 路径：从文件名自动推断章节；模版已有则更新该节，无则末尾增章（不必写 methods.json）',
    )
    parser.add_argument(
        '--infer-section',
        action='store_true',
        help='与 -s 合用：从 SOP 文件名推断 --section，新检验项目自动增章（单文件版 --merge-sops）',
    )
    parser.add_argument('-o', '--output', help='输出路径')
    parser.add_argument('-r', '--ref', help='参考模板路径（--extract-only 模式使用）')
    parser.add_argument('--style', choices=['flat', 'hierarchical'], default='hierarchical',
                        help='flat=RA-正文扁平格式, hierarchical=RA-5/6级标题')
    parser.add_argument('--extract-only', action='store_true',
                        help='仅提取内容输出 JSON，不写入 docx')
    parser.add_argument('--from-json', metavar='PATH',
                        help='从 refined.json 写入 docx')
    parser.add_argument('--write-template',
                        help='已弃用：该参数会被忽略，写入模板固定使用 -t')
    parser.add_argument('--run-id',
                        help='流程标识（用于仅清理当次流程产生的 refined 文件）')
    parser.add_argument('--finalize', action='store_true',
                        help='标记本次输出为最终文件，并清理同 run-id 的中间 refined 文件')

    print_skill_iron_rules()
    args = parser.parse_args()
    args.intrinsic_auto_order_spec = None

    # 模式 3：从 JSON 写入
    if args.from_json:
        run_from_json(args)
        return

    # 模式 1 和 2 都需要 template 和 items
    if not args.template:
        dt = _default_analysis_template_path()
        if dt:
            args.template = os.path.abspath(dt)
            print(f'[INFO] 未指定 -t，使用默认模版: {_disp(args.template)}')
        else:
            parser.error('请指定 -t/--template（或在 input/ 放置 32s42-分析方法-模板文件.docx）')

    if args.merge_sops and (args.sop or args.section or args.infer_section):
        parser.error('不要同时使用 --merge-sops 与 -s/--section/--infer-section')
    if args.infer_section and not args.sop:
        parser.error('--infer-section 需要配合 -s/--sop 指定 SOP 文件')
    if args.infer_section and args.section:
        parser.error('--infer-section 与 --section 互斥')

    items: List[dict] = []
    if args.merge_sops:
        paths = [os.path.normpath(p.strip().strip('"').strip("'")) for p in args.merge_sops]
        spec, auto_items, logs = build_auto_merge_spec_and_items(args.template, paths)
        for line in logs:
            print(f'  {line}')
        if not auto_items:
            parser.error('--merge-sops 未得到任何有效 SOP 条目，请检查路径与文件名')
        args.intrinsic_auto_order_spec = spec
        if args.config:
            cfg_path = args.config
            if not os.path.isabs(cfg_path):
                cfg_path = os.path.join(os.getcwd(), cfg_path)
            items = _merge_auto_items_with_config(load_config(cfg_path), auto_items)
        else:
            items = auto_items
    elif args.infer_section:
        sop_one = args.sop.strip().strip('"').strip("'") if args.sop else ''
        if not sop_one:
            parser.error('-s/--sop 不能为空')
        spec, auto_items, logs = build_auto_merge_spec_and_items(
            args.template, [os.path.normpath(sop_one)]
        )
        for line in logs:
            print(f'  {line}')
        if not auto_items:
            parser.error('无法从该 SOP 推断章节，请改用 --section 显式指定')
        args.intrinsic_auto_order_spec = spec
        items = auto_items
    else:
        items = _parse_items(args)

    if not items:
        parser.error('请指定 --merge-sops、--infer-section -s、或 --sop + --section、或 --config')

    if args.extract_only:
        # 模式 2：仅提取
        run_extract_only(args, items)
    else:
        # 模式 1：直接整合
        run_direct_integrate(args, items)


if __name__ == '__main__':
    main()
